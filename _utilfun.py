import pandas as pd
import openpyxl
from openpyxl import load_workbook, Workbook
from pandas.tseries.offsets import BDay
import os
import shutil
import docx
import win32com
import win32com.client
import time
import pyinputplus as pyip
import win32api
import sys
import win32print
import subprocess
import logging


logger = logging.getLogger('server_logger')
logger.setLevel(logging.DEBUG)
     # create file handler which logs even debug messages
fh = logging.FileHandler('server.log')
fh.setLevel(logging.DEBUG)
     # create console handler with a higher log level
ch = logging.StreamHandler()
ch.setLevel(logging.ERROR)
     # create formatter and add it to the handlers
formatter = logging.Formatter('%(asctime)s - %(message)s --> %(funcName)s', datefmt='%Y-%m-%d %H:%M:%S')
ch.setFormatter(formatter)
fh.setFormatter(formatter)
     # add the handlers to logger
logger.addHandler(ch)
logger.addHandler(fh)

if (logger.hasHandlers()):
    logger.handlers.clear()

logger.addHandler(fh)

def ingreso_egreso(line,frame,legajo,exepciones,nombre,medioDia):
    try:
        exepciones = exepciones
        dia = line.split()[1]
        # hora_ingreso_limite = pd.to_datetime(('{} 6:00').format(dia))
        # hora_ingreso = pd.to_datetime(('{} 8:00').format(dia))
        # hora_ingreso_extra = pd.to_datetime(('{} 7:15').format(dia))
        # hora_salida_limite = pd.to_datetime(('{} 19:45').format(dia))
        # hora_comida_reingreso = pd.to_datetime(('{} 13:10').format(dia))
        # hora_comida_egreso = pd.to_datetime(('{} 12:22').format(dia))
        
        # salida = pd.to_datetime(('{} 16:48').format(dia))
        # comida = pd.Timedelta(seconds=2880)
        
        ingresos_egresos = []
        for x in range(2,12,1):
            try:
                ingresos_egresos.append(pd.to_datetime(('{} {}').format(dia,line.split()[x])))            
            except:
               ingresos_egresos.append(pd.to_datetime(('{} 00:00').format(dia)))
    
        lista_final = []
        lista_final.append(ingresos_egresos[0])
        for indice in range(len(ingresos_egresos) -1):
            
    #        if ingresos_egresos[indice + 1] != 0:
                if ingresos_egresos[indice + 1] - ingresos_egresos[indice] > pd.Timedelta(minutes=5):
                    lista_final.append(ingresos_egresos[indice + 1])
                else:
                    lista_final.append(pd.to_datetime(('{} 00:00').format(dia)))
                    #lista_final.append(ingresos_egresos[indice + 1])

        
        for u in range((10-len(lista_final))):
            lista_final.append(pd.to_datetime(('{} 00:00').format(dia)))         
        
        lista_final.insert(0,line.split()[0])
        lista_final.insert(1,line.split()[1])
        
        horas_trabajadas = 0
    
        for idx in range(2,11,2):
            if lista_final[idx + 1] == pd.to_datetime(('{} 00:00').format(dia)) and lista_final[idx] != pd.to_datetime(('{} 00:00').format(dia)):
                if legajo not in exepciones:
                    horas_trabajadas = 0
                    break
                else:
                    if lista_final[idx + 1] == pd.to_datetime(('{} 00:00').format(dia)) and lista_final[idx] > pd.to_datetime(('{} 19:45').format(dia)):
                        horas_trabajadas = 0
                        break
                    else:
                        horas_trabajadas += 0
                        break
                        
            if lista_final[idx] <= hora_ingreso and lista_final[idx] >= hora_ingreso_limite:
                if lista_final[idx + 1] > salida:
                    horas_trabajadas += (salida - hora_ingreso).seconds
                    continue
                else:              
                    horas_trabajadas += (lista_final[idx + 1] - hora_ingreso).seconds
                    continue
            
            if lista_final[idx] > hora_ingreso and lista_final[idx] <= salida:          
        
                if lista_final[idx + 1]  > salida:
                    horas_trabajadas += (salida - lista_final[idx]).seconds
                    continue
                    
                if lista_final[idx + 1]  == salida:
                    horas_trabajadas += (salida - lista_final[idx]).seconds
                    continue
                    
                if lista_final[idx + 1] < hora_comida_reingreso:
                    horas_trabajadas += (lista_final[idx + 1] - lista_final[idx]).seconds
                    continue
                     
                if lista_final[idx] > hora_comida_reingreso and lista_final[idx + 1] >= salida:
    
                    horas_trabajadas += (salida - lista_final[idx]).seconds
                    continue
                    
                if lista_final[idx] >= hora_comida_reingreso and lista_final[idx + 1] < salida:
    
                    horas_trabajadas += (lista_final[idx + 1] - lista_final[idx]).seconds
                    continue
                
                if lista_final[idx] < hora_comida_egreso:
                     if lista_final[idx + 1] <= salida:
                        horas_trabajadas += (lista_final[idx + 1] - lista_final[idx]).seconds
                     else:
                        horas_trabajadas += (salida - lista_final[idx]).seconds
                else:
                    if lista_final[idx + 1] <= salida:
                        horas_trabajadas += (lista_final[idx + 1] - lista_final[idx]).seconds
                    else:
                        horas_trabajadas += (salida - lista_final[idx]).seconds
                
                    
                if lista_final[idx] > salida:
                    pass
            
            if lista_final[idx] > hora_ingreso and lista_final[idx] > salida:
                horas_trabajadas += (salida - lista_final[idx]).seconds
                continue
                     
        
        comida = comida.seconds
    
        if horas_trabajadas != 0 and line.split()[1] not in medioDia:
                for ix in range(4,12,2):
                    if lista_final[ix] == pd.to_datetime(('{} 00:00').format(dia)) and ix == 4:
                        
                        if lista_final[ix - 1] >= hora_comida_reingreso:
                            
                            horas_trabajadas = round((horas_trabajadas - comida) / 3600,2)  
                                
                            break
                        else:
                            if lista_final[ix - 1] > hora_comida_egreso and lista_final[ix - 1] < hora_comida_reingreso :
                                horas_trabajadas = round((horas_trabajadas - (lista_final[ix - 1] - hora_comida_egreso).seconds) / 3600,2)
                                break
                            else:
                                horas_trabajadas = round((horas_trabajadas) / 3600,2)
                                break
                    
                    else:
                        if lista_final[ix] == hora_comida_reingreso:
                            if lista_final[ix - 1] <= hora_comida_egreso:
                                horas_trabajadas = round((horas_trabajadas) / 3600,2)
                                break
                            else:
                                if lista_final[ix - 1] > hora_comida_egreso and lista_final[ix - 1] < hora_comida_reingreso:
                                    horas_trabajadas = round(((horas_trabajadas)- (lista_final[ix + 1] - hora_comida_egreso ).seconds)/ 3600 ,2)
                                    break

                                
                        
                        if lista_final[ix] <= hora_comida_egreso:
                            if lista_final[ix + 1] >= hora_comida_reingreso:
                                horas_trabajadas = round((horas_trabajadas - comida) / 3600,2)
                                break
                            else:
                                if lista_final[ix + 1] >= hora_comida_egreso:
                                    horas_trabajadas = round(((horas_trabajadas)- (lista_final[ix + 1] - hora_comida_egreso ).seconds)/ 3600 ,2)
                                    break
#                                else:
#                                    horas_trabajadas = round(horas_trabajadas / 3600 ,2)
#                                    break
                            
                        if lista_final[ix] > hora_comida_egreso and lista_final[ix] < hora_comida_reingreso:
                            if lista_final[ix + 1] >= hora_comida_reingreso:
                                if lista_final[ix + 1] <= salida:
                                    horas_trabajadas = round(((horas_trabajadas)- (hora_comida_reingreso - lista_final[ix]).seconds)/ 3600 ,2)
                                    break
    
                        if lista_final[ix] > hora_comida_reingreso:
                            if lista_final[ix - 1] >= hora_comida_reingreso:

                                horas_trabajadas = round((horas_trabajadas - comida) / 3600,2)

                            else:
                                 horas_trabajadas = round((horas_trabajadas) / 3600,2)                               

                                                      
                        
    
        else:
            horas_trabajadas = round((horas_trabajadas) / 3600,2)

        
        lista_final.append(horas_trabajadas)
        
        hora_extra = 0
        for hora in range(2,12):
            if legajo in exepciones:
                hora_extra +=0
            else:          
                if lista_final[hora] > salida:
                    hora_extra +=((lista_final[hora] - salida).seconds)
                    
        if lista_final[2] <=  hora_ingreso_extra and lista_final[2] > hora_ingreso_limite:
            hora_extra += ((hora_ingreso - lista_final[2]).seconds)
        
        lista_final.append(round(hora_extra / 3600,2))
        
        
        
        frame = frame.append({'Empleado':legajo,'Nombre':nombre,'Dia':lista_final[0],'Fecha':lista_final[1],'Ingreso_0':lista_final[2],'Egreso_0':lista_final[3],
                              'Ingreso_1':lista_final[4],'Egreso_1':lista_final[5],'Ingreso_2':lista_final[6],'Egreso_2':lista_final[7],
                              'Ingreso_3':lista_final[8],'Egreso_3':lista_final[9],'Ingreso_4':lista_final[10],'Egreso_4':lista_final[11],
                              'H.Ext.Empleados': lista_final[13],'H.Norm.Emp':lista_final[12]},ignore_index=True)
                                
        
        return frame
    except:
        logger.exception('',exc_info=True)

def dias_fechasY_feriados():
    try:
    
#    inicio_semana = str(input('Ingrese el primer dia habil de la semana en formato DD/MM/AAAA: '))
        inicio_semana = pd.to_datetime(pyip.inputDate('Ingrese el primer dia habil de la semana en formato DD/MM/AAAA: ',formats=["%d/%m/%Y"]))
        
        Medio_dia =[]
        print('\n')
        si_medio_no = pyip.inputYesNo('Existio un dia laboral en que se haya trabajado medio dia?(Yes/No)')
        if si_medio_no =='yes':
            medio_dia = str(pyip.inputDate('Dia en formato DD/MM/AAAA que se haya trabajado medio dia: ',formats=["%d/%m/%Y"]))
            medio_dia = str(medio_dia).split()[0].replace('-','/')
            medio_dia = medio_dia.split('/')[2]+'/'+medio_dia.split('/')[1]+'/'+medio_dia.split('/')[0]
            Medio_dia.append(medio_dia)
        else:
            medio_dia=''
    
        ultimo_dia = inicio_semana + pd.Timedelta(days=4)
        inicio_semana = str(inicio_semana).split()[0].replace('-','/')
        inicio_semana = inicio_semana.split('/')[2]+'/'+inicio_semana.split('/')[1]+'/'+inicio_semana.split('/')[0]
        ultimo_dia = str(ultimo_dia).split()[0].replace('-','/')
        ultimo_dia = ultimo_dia.split('/')[2]+'/'+ultimo_dia.split('/')[1]+'/'+ultimo_dia.split('/')[0]
    
        Feriados = []
        vacaciones = []
        response = 0
        print('\n')
        si_o_no = pyip.inputYesNo('¿Desea agregar un feriado en la semana bajo analisis?(Yes/No) ')
        while si_o_no == 'yes':        
            print('Ingrese fecha del dia feriado con formato DD/MM/AAAA:')
            response = str(pyip.inputDate('Dia en formato DD/MM/AAAA que haya sido feriado: ',formats=["%d/%m/%Y"])) 
            response = str(response).split()[0].replace('-','/')
            response = response.split('/')[2]+'/'+response.split('/')[1]+'/'+response.split('/')[0]
            Feriados.append(response)
            si_o_no = pyip.inputYesNo('¿Desea agregar otro feriado en la semana bajo analisis? ')
        
        print('\n')
        vacac = pyip.inputYesNo('¿Algun operario se encontro de vacaciones/licencia durante la semana bajo analisis?(Yes/No) ')
        while vacac == 'yes':        
            response = str(pyip.inputInt('Ingrese el legajo del operario: ',min=1))
            vacaciones.append(str(response))
            vacac = pyip.inputYesNo('¿Desea agregar otro operario en la semana bajo analisis?(Yes/No) ')
        print('\n')
        print('Inicio semana: ==>',inicio_semana)
        print('Medio_dia: ======>', Medio_dia)
        print('Ultimo dia ======>', ultimo_dia)
        print('Feriados ========>', Feriados)
        print('Vacaciones ======>', vacaciones)
        print('\n')
    
        
        return inicio_semana,Feriados,Medio_dia,ultimo_dia,vacaciones
    except:
        logger.exception('',exc_info=True)

def horas_olvidados(frame,medioDia,exclusion):
    try:
        comida = pd.Timedelta(seconds=2880)
        comida = comida.seconds
        
        indices = [frame[frame['H.Norm.Emp'] == 0].index[x] for x in range(len(frame[frame['H.Norm.Emp'] == 0]))]
        
        
        exclusion = exclusion
        """ IMPORTANTE, LAS LISTAS DE EXCLUSION DEBEN ESTAR COMO VARIABLES"""
        for posicion in indices:
            horas_trabajadas = 0
            horas_extras_trabajadas = 0
            legajo = 0
            dia = frame.iloc[posicion,3]
    
            hora_ingreso_limite = pd.to_datetime(('{} 6:00').format(dia))
    
            hora_ingreso = pd.to_datetime(('{} 8:00').format(dia))
            hora_ingreso_extra = pd.to_datetime(('{} 7:15').format(dia))
            hora_salida_limite = pd.to_datetime(('{} 19:45').format(dia))
            salida = pd.to_datetime(('{} 16:48').format(dia))
            
            for i in range(4,14,2):
                legajo = str(frame.iloc[posicion,0])
                if  frame.iloc[posicion,i] >= hora_ingreso_extra and frame.iloc[posicion,i] <= hora_ingreso :         
                    if frame.iloc[posicion,i+ 1] <= salida:                
                        horas_trabajadas += (frame.iloc[posicion,i+ 1] - hora_ingreso).seconds
                        
                    else: 
                        horas_trabajadas += (salida - hora_ingreso).seconds
                        horas_extras_trabajadas += (frame.iloc[posicion,i+ 1] - salida).seconds
                        
                    
                if frame.iloc[posicion,i] > hora_ingreso:               
                    if frame.iloc[posicion,i+ 1] <= salida:                    
                        horas_trabajadas += (frame.iloc[posicion,i+ 1] - frame.iloc[posicion,i]).seconds
                        
                        
                    else:                    
                        horas_trabajadas += (salida - frame.iloc[posicion,i]).seconds
                        horas_extras_trabajadas += (frame.iloc[posicion,i+ 1] - salida).seconds
    
                            
                        
                    
                    
                if frame.iloc[posicion,i] < hora_ingreso_extra and frame.iloc[posicion,i] != pd.to_datetime(('{} 00:00').format(dia)):
                    horas_extras_trabajadas += (hora_ingreso - frame.iloc[posicion,i]).seconds
                    
                    
                    if frame.iloc[posicion,i + 1] > salida:
                        horas_extras_trabajadas += (frame.iloc[posicion,i + 1] - salida).seconds
                        horas_trabajadas += (salida - hora_ingreso).seconds
                    else:
                        horas_trabajadas += (frame.iloc[posicion,i + 1] - hora_ingreso).seconds
            
            if frame.iloc[posicion,3] not in medioDia:         
    #            frame.iloc[posicion,14] = round(((horas_extras_trabajadas/ 3600) - comida),2)
                if legajo not in exclusion:
                    frame.iloc[posicion,14] = round(((horas_extras_trabajadas)/ 3600),2)
                else:
                    frame.iloc[posicion,14] = 0
    #            frame.iloc[posicion,15] = round(((horas_trabajadas/ 3600) - comida),2)
                frame.iloc[posicion,15] = round(((horas_trabajadas - comida)/ 3600),2)
            else:
                if legajo not in exclusion:
                    frame.iloc[posicion,14] = round(((horas_extras_trabajadas)/ 3600),2)
                else:
                    frame.iloc[posicion,14] = 0
                frame.iloc[posicion,15] = round(((horas_trabajadas/ 3600)),2)
            
            
        return frame
    except:
        logger.exception('',exc_info=True)

def HorasTotales_y_Extras(frame,empleados,feriados,dias_laborables,MedioDia):
    
    """
    Esta funcion se encarga de sumar todas la horas y actualiza el diccionario
    en funcion del criterio de las 40 horas verifica si el empleado trabajo 40 semanales
    y en caso de no haberlo hecho y tener extras, resta extras para llegar a las 40"""
    try:

    
        listaDeDias = [(pd.to_datetime(dias_laborables) + pd.Timedelta(days=x)) for x in range(0,5)]
        listaDeDias = [str(x).split()[0].replace('-','/') for x in listaDeDias]
        listaDeDias = [ultimo_dia.split('/')[2]+'/'+ultimo_dia.split('/')[1]+'/'+ultimo_dia.split('/')[0] for ultimo_dia in listaDeDias]
        

        if len(feriados) == 1 :
            listaDeDias.remove(feriados[0])
        if len(feriados) > 1:
            for feria in feriados:
                listaDeDias.remove(feria)
        
        daysaLaburar = len(listaDeDias)
        halfdays = len(MedioDia)        
        horasSemanales = 8* daysaLaburar - 4.5*halfdays
        

        frame = frame
        
        
        for renglon in range(len(frame)):
            for legajo in empleados.keys():
                if str(legajo) == str(frame.iloc[renglon,0]):
                    empleados[legajo]['Hs.Totales']+= frame.iloc[renglon,15] + frame.iloc[renglon,14]
        for key in empleados.keys():
            if empleados[key]['Hs.Totales'] > horasSemanales:
                empleados[key]['Hs.Normales'] = horasSemanales
                empleados[key]['Hs.Extras'] = empleados[key]['Hs.Totales'] - horasSemanales
            else:
                empleados[key]['Hs.Normales'] = empleados[key]['Hs.Totales']
                empleados[key]['Hs.Extras'] = 0

        return empleados
    except:
        logger.exception('',exc_info=True)

def CreacionInformeExcel(path,empleados):
    """ ESta funcion va a abrir el excel en donde esta el cuadro de horas
    y va a crear el informe"""   
    
    try:
        wb = openpyxl.load_workbook(path)
        wb.create_sheet("Informe")
        wb.active=1
        hoja = wb.active

        
        fila = 1
        columnas = ['Legajo','Nombre','Hs.Normales','Hs.Extras','Hs.Totales']
        for x in range(1,len(columnas)+1):
            hoja.cell(row= fila, column = x).value = columnas[x - 1]   
        fila+= 1    
        

        
        for column in range(len(columnas)):
            if bool(empleados):
                for key in empleados.keys():
                    if key != 'None':
                        if column == 0:
                            hoja.cell(row= fila, column = column + 1).value = key
                            fila+= 1
                        else:
                            hoja.cell(row= fila, column = column+ 1).value = empleados[key][str(columnas[(column)])]
                            fila+=1
                fila = 2
        horasNormalesOperarios = 0
        horasNormalesSupervisores = 0
        horasExtrasOperarios = 0
        horasExtrasSupervisores = 0
        horas = {}
        if bool(empleados):
            for legajo in empleados.keys():
                if empleados[legajo]['Categoria'] == 'Operario':
                    horasNormalesOperarios += empleados[legajo]['Hs.Normales']
                    horasExtrasOperarios+= empleados[legajo]['Hs.Extras']
                    horas['horasNormalesOperarios'] = horasNormalesOperarios
                    horas['horasExtrasOperarios'] = horasExtrasOperarios
                else:
                    horasNormalesSupervisores += empleados[legajo]['Hs.Normales']
                    horasExtrasSupervisores+= empleados[legajo]['Hs.Extras']
                    horas['horasNormalesSupervisores'] = horasNormalesSupervisores
                    horas['horasExtrasSupervisores'] = horasExtrasSupervisores
        
        fila = hoja.max_row + 1
        lista = []
        if bool(horas):
            for key,value in horas.items():
                hoja.cell(row= fila, column = 2).value = key
                hoja.cell(row= fila, column = 3).value = value
                fila+= 1
        
            
                
        wb.save(path)
        wb.close()
        return horas,empleados
    except:
        logger.exception('',exc_info=True)
        
def informeNoFichadas(frame,empleados,dias_laborables,feriados,MedioDia,pathToFolder):
    try:
        
        diccionario = empleados
        len_noMarca = len(frame[frame['H.Norm.Emp'] == 0])
        doc = docx.Document()
        doc.add_heading(('Informe correspondiente a la semana {}').format(dias_laborables), 0)
        c = doc.add_paragraph('Informe sobre No fichajes: \n')
        for x in range(len_noMarca):
            dia = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],3]
            legajo = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],0]
            hora = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],4]
            hora_2 = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6]
            if dia not in feriados and dia in MedioDia: 
                
                if hora_2 == pd.to_datetime(('{} 00:00').format(dia)):                   
                    if hora >= pd.to_datetime(('{} 10:00').format(dia)):
                        c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Ingreso.').format(dia,diccionario[legajo]['Nombre']))
                    else:
                        c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Salida.').format(dia,diccionario[legajo]['Nombre']))
    
                else:
                    if hora_2 >= pd.to_datetime(('{} 10:00').format(dia)):
                        c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Re-ingreso.').format(dia,diccionario[legajo]['Nombre']))
                    else:   
                        c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Salida.').format(dia,diccionario[legajo]['Nombre']))
                        
            else:
                if dia not in feriados: 
                
                    if hora_2 == pd.to_datetime(('{} 00:00').format(dia)):                   
                        if hora >= pd.to_datetime(('{} 15:00').format(dia)):
                            c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Ingreso.').format(dia,diccionario[legajo]['Nombre']))
                        else:
                            c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Salida.').format(dia,diccionario[legajo]['Nombre']))
        
                    else:
                        if hora_2 >= pd.to_datetime(('{} 15:00').format(dia)):
                            c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Re-ingreso.').format(dia,diccionario[legajo]['Nombre']))
                        else:   
                            c.add_run(('\n\tEl dia {} el empleado {:10s} no ficho Salida.').format(dia,diccionario[legajo]['Nombre'])) 
        
        fini = '\\Informe No fichaje Semana {}.docx'.format(dias_laborables.replace('/','-'))
        pathToWord = pathToFolder + fini
        fini_pdf = ('\\Informe No fichaje Semana {}.pdf').format(dias_laborables.replace('/','-'))
        pathToPDF = pathToFolder + fini_pdf
        doc.save(pathToWord)
        
        
        wordFilename = pathToWord
        
        pdfFilename = pathToPDF
        
        
        wdFormatPDF = 17 # Word's numeric code for PDFs.
        wordObj = win32com.client.Dispatch('Word.Application')
        docObj = wordObj.Documents.Open(wordFilename)
        docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
        docObj.Close()
        wordObj.Quit()
        os.remove(pathToWord)
    except:
        logger.exception('',exc_info=True)

def TardanzasyFaltas(frame,empleados,dias_laborables,MedioDia,feriados,vacaciones):
    try:
        try:
            del empleados['None']
        except:
            pass
        nombreDias = {'Monday':'Lunes',
                       'Tuesday': 'Martes',
                       'Wednesday': 'Miércoles',
                       'Thursday': 'Jueves',
                       'Friday': 'Viernes'}
        for legajo in empleados.keys():
            
            listaDeDias = [(pd.to_datetime(dias_laborables,dayfirst=True) + pd.Timedelta(days=x)) for x in range(0,5)]
            listaDeDias = [str(x).split()[0].replace('-','/') for x in listaDeDias]
            listaDeDias = [ultimo_dia.split('/')[2]+'/'+ultimo_dia.split('/')[1]+'/'+ultimo_dia.split('/')[0] for ultimo_dia in listaDeDias]
    
            empleados[legajo]['Faltas'] = {}
            empleados[legajo]['Tardanzas'] = {}
            empleados[legajo]['Retiros Anticipados'] = {}
            
    
            if len(feriados) == 1 :
                listaDeDias.remove(feriados[0])
            if len(feriados) > 1:
                for feria in feriados:
                    listaDeDias.remove(feria)
            diaFaltado = listaDeDias.copy()
    
            """ACA"""
            for idx in range(len(frame)):
                dia = frame.iloc[idx,3]
                if legajo == str(frame.iloc[idx,0]):
                    if frame.iloc[idx,3] in listaDeDias:
                        diaFaltado.remove(frame.iloc[idx,3])
                        
                    if frame.iloc[idx,4] > pd.to_datetime(('{} 8:00').format(dia)) and frame.iloc[idx,4] < pd.to_datetime(('{} 15:00').format(dia)):
                        empleados[legajo]['Tardanzas'][frame.iloc[idx,3]] = round(((frame.iloc[idx,4] -  pd.to_datetime(('{} 8:00').format(dia))).seconds/ 60),2)
                    
                    
                    for ColumnaEgreso in range(5,14,2):
                        if frame.iloc[idx,ColumnaEgreso] == pd.to_datetime(('{} 00:00').format(dia)):
                            if frame.iloc[idx,3] not in MedioDia:
                                
                                if frame.iloc[idx,ColumnaEgreso- 2] >= pd.to_datetime(('{} 16:48').format(dia)):
                                    
                                    break
                                else:
                                    
                                    empleados[legajo]['Retiros Anticipados'][frame.iloc[idx,3]] = round(((pd.to_datetime(('{} 16:48').format(dia)) -frame.iloc[idx,ColumnaEgreso- 2]).seconds/ 60),2)
                                    break
                            else:
                                if frame.iloc[idx,ColumnaEgreso- 2] >= pd.to_datetime(('{} 12:30').format(dia)):
                                    
                                    break
                                else:
                                    
                                    empleados[legajo]['Retiros Anticipados'][frame.iloc[idx,3]] = round(((pd.to_datetime(('{} 12:30').format(dia)) -frame.iloc[idx,ColumnaEgreso- 2]).seconds/ 60),2)
                                    break
                    
                    
    #                if frame.iloc[idx,9] < pd.to_datetime('16:48',format = "%H:%M") and frame.iloc[idx,9] != pd.to_datetime('00:00',format = "%H:%M"):
    #                    empleados[legajo]['Retiros Anticipados'][frame.iloc[idx,3]] = round(((pd.to_datetime('16:48',format = "%H:%M") -frame.iloc[idx,9]).seconds/ 3600),2)
    #                else:
    #                    if frame.iloc[idx,7] < pd.to_datetime('16:48',format = "%H:%M") and frame.iloc[idx,7] != pd.to_datetime('00:00',format = "%H:%M"):
    #                        empleados[legajo]['Retiros Anticipados'][frame.iloc[idx,3]] = round(((pd.to_datetime('16:48',format = "%H:%M") -frame.iloc[idx,7]).seconds/ 3600),2)
    #                    else:
    #                        if frame.iloc[idx,5] < pd.to_datetime('16:48',format = "%H:%M") and frame.iloc[idx,5] != pd.to_datetime('00:00',format = "%H:%M"):
    #                            empleados[legajo]['Retiros Anticipados'][frame.iloc[idx,3]] = round(((pd.to_datetime('16:48',format = "%H:%M") -frame.iloc[idx,5]).seconds/ 3600),2)
    #                
    
            for falta in diaFaltado:
                if legajo != None:
                    if legajo not in vacaciones:
                        falta_2 = falta.split('/')[1]+'/'+falta.split('/')[0]+'/'+falta.split('/')[2]
                        empleados[legajo]['Faltas'][str(falta)] = nombreDias[pd.to_datetime(falta_2).day_name()]
#                        empleados[legajo]['Faltas'][str(falta)] = nombreDias[pd.to_datetime(falta).day_name()]
        
        return empleados
    except:
        logger.exception('',exc_info=True)

def impresionInforme(pahtToFolder,dias_laborables):
    try:
        fini = 'Informe No fichaje Semana {}.pdf'.format(dias_laborables.replace('/','-'))
        filename = pahtToFolder+'\\'+fini    
        CurrentPrinter = win32print.GetDefaultPrinter()
        win32print.SetDefaultPrinter(CurrentPrinter)
        win32api.ShellExecute (0, "print", filename, CurrentPrinter, ".", 0)
        win32print.SetDefaultPrinter(CurrentPrinter)
        os.system("TASKKILL /F /IM AcroRD32.exe")
    except:
        logger.exception('',exc_info=True)

def creacionInformeFaltasTardanzas(pathToFolder,empleados,dias_laborables,vacaciones):
    try:
        doc = docx.Document()
        doc.add_heading(('Informe correspondiente a la semana {}').format(dias_laborables), 0)
        for legajo in empleados.keys():
            if empleados[legajo]['Nombre'] == None:
                continue
            else:
                if bool(empleados[legajo]['Faltas']) or bool(empleados[legajo]['Tardanzas']) or bool(empleados[legajo]['Retiros Anticipados']):           
                    doc.add_paragraph('Informe sobre: '+empleados[legajo]['Nombre'])
                    c = doc.add_paragraph()            
                    for falta in empleados[legajo]['Faltas'].keys():
                        if legajo not in vacaciones:
                            if falta != None:
                                c.add_run(('{:15s}El empleado {:10s} falto el dia {:7s} {}.\n').format('',empleados[legajo]['Nombre'],
                                          empleados[legajo]['Faltas'][str(falta)],str(falta)))
                    for tardanza in empleados[legajo]['Tardanzas'].keys():
                        if tardanza != None:
                            c.add_run(('{:15s}El dia {} el empleado {} ingreso {} minutos tarde.\n').format('',tardanza,empleados[legajo]['Nombre'],
                                      (empleados[legajo]['Tardanzas'][tardanza])))
                    
                    for retiro in empleados[legajo]['Retiros Anticipados'].keys():
                        if retiro != None:
                            c.add_run(('{:15s}El dia {} el empleado {} se retiro {} minutos antes.\n').format('',retiro,empleados[legajo]['Nombre'],
                                      (empleados[legajo]['Retiros Anticipados'][retiro])))
                    
          
        
    
        
    
        fini = '\\'+'Informe Semana {}.docx'.format(dias_laborables.replace('/','-'))
        wordFilename = pathToFolder + fini
    
        fini_pdf = '\\'+('Informe Semana {}.pdf'.format(dias_laborables.replace('/','-')))
        pdfFilename = pathToFolder + fini_pdf
    
        doc.save(wordFilename)
    #    wordFilename = r'E:\MachinaData\Escritorio\Vulcano\Informe Semana {}.docx'.format(dias_laborables.replace('/','-'))
    #    pdfFilename = r'E:\MachinaData\Escritorio\Vulcano\Informe Semana {}.pdf'.format(dias_laborables.replace('/','-'))
        
        wdFormatPDF = 17 # Word's numeric code for PDFs.
        wordObj = win32com.client.Dispatch('Word.Application')
        docObj = wordObj.Documents.Open(wordFilename)
        docObj.SaveAs(pdfFilename, FileFormat=wdFormatPDF)
        docObj.Close()
        wordObj.Quit()
        os.remove(wordFilename)
    except:
        logger.exception('',exc_info=True)
    
def AppendExcel(pathToHorasDeposito,frame,empleados,dias_laborables,ultimoDia):
    """ ESta funcion va a abrir el excel en donde esta el Horas deposito y va a arrojar el frame va a crear el informe
    
    path: en este caso va a dirigir a la carpeta donde esta el archivo Horas Deposito 
    
    
    """   
    try:
        
        ultimoDia = ultimoDia.replace('/','-')
        dias_laborables = dias_laborables.replace('/','-')    
        inicio = dias_laborables.split('-')[0]+'-'+dias_laborables.split('-')[1]
        fin = ultimoDia.split('-')[0]+'-'+ultimoDia.split('-')[1]
        titulo = '{} al {}'.format(inicio,fin)     
        wb = openpyxl.load_workbook(pathToHorasDeposito)
        wb.create_sheet(titulo)
        try:
            wb.save(pathToHorasDeposito)
            wb.close()
        except:
            print('Tiene el archivo Horas Deposito abierto, tiene 10 segundos para cerrarlo.')
            time.sleep(10)
            wb.save(pathToHorasDeposito)
            wb.close()
        
        writer = pd.ExcelWriter(pathToHorasDeposito, engine='openpyxl')
        writer.book = load_workbook(pathToHorasDeposito)
        writer.sheets = dict((ws.title,ws) for ws in writer.book.worksheets)
        frame.to_excel(writer,sheet_name=titulo,index=False, header=False)
            
        writer.save()
        writer.close()
        
        
        wb = openpyxl.load_workbook(pathToHorasDeposito)
        for s in range(len(wb.sheetnames)):
            if wb.sheetnames[s] == titulo:
                break
        wb.active = s
        hoja = wb.active
        
        fila = hoja.max_row+ 2
        columnas = ['Legajo','Nombre','Hs.Normales','Hs.Extras','Hs.Totales']
        for x in range(1,len(columnas)+1):
            hoja.cell(row= fila, column = x).value = columnas[x - 1]   
        fila+= 1    
        traba = fila
        
        for column in range(len(columnas)):
            for key in empleados.keys():
                if key != 'None':
                    if column == 0:
                        hoja.cell(row= fila, column = column + 1).value = key
                        fila+= 1
                    else:
                        hoja.cell(row= fila, column = column+ 1).value = empleados[key][str(columnas[(column)])]
                        fila+=1
            fila = traba
            
        horasNormalesOperarios = 0
        horasNormalesSupervisores = 0
        horasExtrasOperarios = 0
        horasExtrasSupervisores = 0
        horas = {}
        for legajo in empleados.keys():
            if empleados[legajo]['Categoria'] == 'Operario':
                horasNormalesOperarios += empleados[legajo]['Hs.Normales']
                horasExtrasOperarios+= empleados[legajo]['Hs.Extras']
                horas['horasNormalesOperarios'] = horasNormalesOperarios
                horas['horasExtrasOperarios'] = horasExtrasOperarios
            else:
                horasNormalesSupervisores += empleados[legajo]['Hs.Normales']
                horasExtrasSupervisores+= empleados[legajo]['Hs.Extras']
                horas['horasNormalesSupervisores'] = horasNormalesSupervisores
                horas['horasExtrasSupervisores'] = horasExtrasSupervisores
                
        
        fila = hoja.max_row + 2
        for key,value in horas.items():
            hoja.cell(row= fila, column = 2).value = key
            hoja.cell(row= fila, column = 3).value = value
            fila+= 1   
            
                
        try:
            wb.save(pathToHorasDeposito)
            wb.close()
        except:
            print('Tiene el archivo Horas Deposito abierto, tiene 10 segundos para cerrarlo.')
            time.sleep(10)
            wb.save(pathToHorasDeposito)
            wb.close()
            
        return horas,empleados
    except:
        logger.exception('',exc_info=True)        

def escrituraExcelDeposito_EFiciencia(pathToPersonal,empleados,dias_laborables,
                                     ultimoDia,cant_horas,feriados,MedioDia,
                                     pathToIndicator,pathToMovs):
    try:
        try:
            del empleados['None']
        except:
            pass
        print('A continuacion ingrese el numero de mes en analisis Ej: Marzo = 3')
        meses = {'1':'Enero',
                 '2':'Febrero',
                 '3':'Marzo',
                 '4':'Abril',
                 '5':'Mayo',
                 '6':'Junio',
                 '7':'Julio',
                 '8':'Agosto',
                 '9':'Septiembre',
                 '10':'Octubre',
                 '11':'Noviembre',
                 '12':'Diciembre',
                 }
        mes = pyip.inputInt('Ingrese el numero de mes (formato = X): ',min=1,max=12)
        response = pyip.inputYesNo(('El mes ingresado corresponde a: {}, es correcto? (Yes/No)').format(meses[str(mes)]))    
        while response == 'no':
            mes = pyip.inputInt('Ingrese el numero de mes (formato = X): ',min=1,max=12)
            response = pyip.inputYesNo(('El mes ingresado corresponde a: {}, es el mes deseado? (Yes/No)').format(meses[str(mes)]))
        
        wb = openpyxl.load_workbook(pathToPersonal)
        hoja = wb.active
        posicionColumna_Asistencia = 0
        posicionColumna_Tardanza = 0
        
    
#        
#        for x in range(4,hoja.max_column+ 1):
#            if (hoja.cell(1,x).value) == 'FALTAS (Días)':
#                posicionColumna_Asistencia = x
#            if (hoja.cell(1,x).value) == 'RETIROS (HS)':
#                posicionColumna_Tardanza = x
        
        
        for x in range(4,hoja.max_column+ 1):
            if (hoja.cell(1,x).value) == 'ASISTENCIA':
                posicionColumna_Asistencia = x+ 13
            if (hoja.cell(1,x).value) == 'RETIROS ANTICIPADOS Y LLEGADAS TARDES (HORAS)':
                posicionColumna_Tardanza = x+ 13
        
        for legajo in empleados.keys():
    
            for i in range(3,hoja.max_row + 1):
                legajo_operario = hoja.cell(i,1).value
    
                if legajo == str(legajo_operario):
                    for indAsis in range(posicionColumna_Asistencia,posicionColumna_Asistencia+ 12):
                        if mes == (hoja.cell(2,indAsis).value).month:
                            valor = hoja.cell(i,indAsis).value
                            inasistencia = 0
                            if type(valor) == type(None):
                                valor = 0
                            if bool(empleados[legajo]['Faltas']):
                                for ina in empleados[legajo]['Faltas'].values():
                                    inasistencia+= 1                                                    
                            hoja.cell(i,indAsis).value = inasistencia + valor
                    
                    for indTard in range(posicionColumna_Tardanza,posicionColumna_Tardanza+ 12):
                        if mes == (hoja.cell(2,indTard).value).month:
                            valor_tard = hoja.cell(i,indTard).value
                            tardanza = 0
                            if type(valor_tard) == type(None):
                                valor_tard = 0
                            if bool(empleados[legajo]['Tardanzas']):
                                for tard in empleados[legajo]['Tardanzas'].values():
                                    tardanza+= tard
                            if bool(empleados[legajo]['Retiros Anticipados']):
                                for ret in empleados[legajo]['Retiros Anticipados'].values():
                                    tardanza+= ret
                            tardanza = tardanza / 60
                            
                            hoja.cell(i,indTard).value = (tardanza + valor_tard)
        try:
            wb.save(pathToPersonal)
            wb.close()
        except:
            print('Tiene el archivo 1.Personal Gral abierto, tiene 10 segundos para cerrarlo\n.')
            time.sleep(10)
            wb.save(pathToPersonal)
            wb.close()
        
        print('A continuacion se abrira el indicador de eficiencia y el archivos de movimientos.')
        subprocess.Popen([pathToIndicator],shell=True)
        subprocess.Popen([pathToMovs],shell=True)
        time.sleep(2)
    except:
        logger.exception('',exc_info=True)
        
    """ EL resto del codigo abre el indicador de eficiencia, carga los datos, actualiza la tabla dinamica y carga el resto de los datos en 
    la ultima hoja. 
    
    No se utiliza ya que sobrecarga la RAM, supongo que por la manipulacion de la tabla dinamica.
    Ademas el tiempo ahorrado no era mucho.
    """
 
#    print('A continuacion se abrira el indicador de eficiencia y el archivos de movimientos.')
#    subprocess.Popen([pathToIndicator],shell=True)
#    subprocess.Popen([pathToMovs],shell=True)
#    time.sleep(2)
#    
#    frame = pd.read_excel(pathToMovs,index_col=None)
#    frame = frame.iloc[1:,0:8]
#    frame['Dates'] = pd.to_datetime(frame['Comp. - F. Emisión'],dayfirst=True)
#    frame['Semana'] = frame['Dates'].dt.week
#    frame['Año'] = frame['Dates'].dt.year
#    frame['Comprobante'] = frame['Comp.'].astype(str)
#    frame['Comprobante'] = frame['Comprobante'].str.slice(0,2,1)
#    frame.drop(['Dates'],axis= 1,inplace=True)
#    
#    wb = openpyxl.load_workbook(pathToIndicator)
#    wb.active = 0
#    hoja = wb.active
#    ultima_fila = hoja.max_row 
#    wb.close()
#
#    writer = pd.ExcelWriter(pathToIndicator, engine='openpyxl')
#    writer.book = load_workbook(pathToIndicator)
#    writer.sheets = dict((ws.title,ws) for ws in writer.book.worksheets)
#    
#    frame.to_excel(writer,sheet_name="Registros", startrow=ultima_fila,index=False, header=False)
#    writer.save()
#    writer.close()
#    
#    wb = openpyxl.load_workbook(pathToIndicator)
#    wb.active = 1
#    hoja = wb.active
#    pivot = hoja._pivots[0] # any will do as they share the same cache
#    pivot.cache.refreshOnLoad = True
#    
#    wb.active = 2
#    hoja = wb.active
#    ultima_fila = hoja.max_row
#    orden = hoja.cell((ultima_fila - 1),1).value + 1
#    hoja.cell(ultima_fila,1).value = orden
#    hoja.cell(ultima_fila,2).value = dias_laborables
#    hoja.cell(ultima_fila,3).value = ultimoDia
#    hoja.cell(ultima_fila,4).value = cant_horas['horasNormalesOperarios']
#    hoja.cell(ultima_fila,5).value = cant_horas['horasExtrasOperarios']
#    
#
#    
#    count = 0
#    for lada in empleados.keys():
#        if empleados[lada]['Categoria'] == 'Operario' and lada != None:
#            count+=1
#    hoja.cell(ultima_fila,7).value = count
#    hoja.cell(ultima_fila,8).value = cant_horas['horasNormalesSupervisores']
#    hoja.cell(ultima_fila,9).value = cant_horas['horasExtrasSupervisores']
#    hoja.cell(ultima_fila,13).value = frame['Semana'][1]
#    
#    cant_fer_medio_dia = len(feriados) + (0.5* len(MedioDia))
#    hoja.cell(ultima_fila,14).value = cant_fer_medio_dia
#    wb.save(pathToIndicator)
#    wb.close()
    
def IngresoNoFichadas(frame,MedioDia,empleados,feriados):
    try:
        print('\n\nLos siguientes operarios no ficharon salida, ingrese la misma en formato HH:MM:SS\n')
        len_noMarca = len(frame[frame['H.Norm.Emp'] == 0])
        for x in range(len_noMarca):
            dia = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],3]
            legajo = str(frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],0])
            hora = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],4]
            hora_2 = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6]
    
            if dia not in feriados:
                
                if hora_2 == pd.to_datetime(('{} 00:00').format(dia)):
                    
                    if hora > pd.to_datetime(('{} 11:00').format(dia)) and dia in MedioDia:
                        print('\nATENCION !!!')
                        print(('\nEl dia {} es considerado como MEDIO DIA laboral. ').format(dia))
                        print(('El dia {} el empleado {:10s} no ficho INGRESO. Ingreso  a la hora: ').format(dia,empleados[legajo]['Nombre']))
                        horaEntrada = str(pyip.inputDatetime('Ingrese el horario de Ingreso en formato HH:MM: ',formats=["%H:%M"]))
                        horaEntrada = horaEntrada.split()[1]
                        horaEntrada = horaEntrada.split(':')[0]+':'+horaEntrada.split(':')[1]
                        horaEntrada = pd.to_datetime(('{} {}').format(dia,str(horaEntrada)))
                        if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] == pd.to_datetime(('{} 00:00').format(dia)):
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],4] = horaEntrada
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] = hora
     
    #            if pd.to_datetime(('{} {}').format(dia,hora)) <= pd.to_datetime(('{} 11:00').format(dia)) and dia in MedioDia:
                       
                    if hora <= pd.to_datetime(('{} 11:00').format(dia)) and dia in MedioDia:
                        print('\nATENCION !!!')
                        print(('\nEl dia {} es considerado como MEDIO DIA laboral. ').format(dia))
                        print(('El dia {} el empleado {:10s} no ficho SALIDA. Ingreso  a la hora: ').format(dia,empleados[legajo]['Nombre']))
                        horaSalida = str(pyip.inputDatetime('Ingrese el horario de salida en formato HH:MM: ',formats=["%H:%M"]))
                        horaSalida = horaSalida.split()[1]
                        horaSalida = horaSalida.split(':')[0]+':'+horaSalida.split(':')[1]
                        horaSalida = pd.to_datetime(('{} {}').format(dia,str(horaSalida)))
                        if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] == pd.to_datetime(('{} 00:00').format(dia)):
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],4] = hora
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] = horaSalida
                    
    #             if pd.to_datetime(('{} {}').format(dia,hora)) >= pd.to_datetime(('{} 15:00').format(dia)) and dia not in MedioDia:
           
                    if hora >= pd.to_datetime(('{} 15:00').format(dia)) and dia not in MedioDia:
                        print(('\nEl dia {} el empleado {:10s} no ficho INGRESO. Ingreso  a la hora: ').format(dia,empleados[legajo]['Nombre']))
                        horaEntrada = str(pyip.inputDatetime('Ingrese el horario de Ingreso en formato HH:MM: ',formats=["%H:%M"]))
                        horaEntrada = horaEntrada.split()[1]
                        horaEntrada = horaEntrada.split(':')[0]+':'+horaEntrada.split(':')[1]
                        horaEntrada = pd.to_datetime(('{} {}').format(dia,str(horaEntrada)))
                        if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] == pd.to_datetime(('{} 00:00').format(dia)):
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],4] = horaEntrada
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] = hora
                    
                    
                    
                    else :
                        if dia not in MedioDia:
                            print(('\nEl dia {} el empleado {:10s} no ficho SALIDA. Salio a la hora: ').format(dia,empleados[legajo]['Nombre']))              
                            horaSalida = str(pyip.inputDatetime('Ingrese el horario de salida en formato HH:MM: ',formats=["%H:%M"]))
                            horaSalida = horaSalida.split()[1]
                            horaSalida = horaSalida.split(':')[0]+':'+horaSalida.split(':')[1]
                            horaSalida = pd.to_datetime(('{} {}').format(dia,str(horaSalida)))
                            
                            if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] == pd.to_datetime(('{} 00:00').format(dia)):
                                frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] = horaSalida
                            else:
                                if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] > pd.to_datetime(('{} 19:45').format(dia)):
                                    horaSalidaDeposito = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] + pd.Timedelta(m=4)
                                    frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] = horaSalida
                                    frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] = horaSalidaDeposito
                    
                    print(('Ingreso: {}').format(str(frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],4]).split()[1]))
                    print(('Salida: {}').format(str(frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5]).split()[1]))
                    print('-'*80)
                    print('\n')
                else:
                    if hora_2 > pd.to_datetime(('{} 11:00').format(dia)) and dia in MedioDia:
                        print('\nATENCION !!!')
                        print(('\nEl dia {} es considerado como MEDIO DIA laboral. ').format(dia))
                        print(('El dia {} el empleado {:10s} no ficho RE-INGRESO. Re-ingreso  a la hora: ').format(dia,empleados[legajo]['Nombre']))
                        horaEntrada = str(pyip.inputDatetime('Ingrese el horario de Ingreso en formato HH:MM: ',formats=["%H:%M"]))
                        horaEntrada = horaEntrada.split()[1]
                        horaEntrada = horaEntrada.split(':')[0]+':'+horaEntrada.split(':')[1]
                        horaEntrada = pd.to_datetime(('{} {}').format(dia,str(horaEntrada)))
                        if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] == pd.to_datetime(('{} 00:00').format(dia)):
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] = horaEntrada
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] = hora_2
                            
                    if hora_2 <= pd.to_datetime(('{} 11:00').format(dia)) and dia in MedioDia:
                        print('\nATENCION !!!')
                        print(('\nEl dia {} es considerado como MEDIO DIA laboral. ').format(dia))
                        print(('El dia {} el empleado {:10s} no ficho SALIDA. Salio a la hora: ').format(dia,empleados[legajo]['Nombre']))
                        horaSalida = str(pyip.inputDatetime('Ingrese el horario de salida en formato HH:MM: ',formats=["%H:%M"]))
                        horaSalida = horaSalida.split()[1]
                        horaSalida = horaSalida.split(':')[0]+':'+horaSalida.split(':')[1]
                        horaSalida = pd.to_datetime(('{} {}').format(dia,str(horaSalida)))
                        if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] == pd.to_datetime(('{} 00:00').format(dia)):
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] = hora_2
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] = horaSalida
                            
                    if hora_2 >= pd.to_datetime(('{} 15:00').format(dia)) and dia not in MedioDia:
                        print(('\nEl dia {} el empleado {:10s} no ficho Re-INGRESO. Re-Ingreso  a la hora: ').format(dia,empleados[legajo]['Nombre']))
                        horaEntrada = str(pyip.inputDatetime('Ingrese el horario de Ingreso en formato HH:MM: ',formats=["%H:%M"]))
                        horaEntrada = horaEntrada.split()[1]
                        horaEntrada = horaEntrada.split(':')[0]+':'+horaEntrada.split(':')[1]
                        horaEntrada = pd.to_datetime(('{} {}').format(dia,str(horaEntrada)))
                        if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] == pd.to_datetime(('{} 00:00').format(dia)):
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] = horaEntrada
                            frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] = hora_2
                            
                    else :
                        if dia not in MedioDia:
                            print(('\nEl dia {} el empleado {:10s} no ficho SALIDA. Salio a la hora: ').format(dia,empleados[legajo]['Nombre']))              
                            horaSalida = str(pyip.inputDatetime('Ingrese el horario de salida en formato HH:MM: ',formats=["%H:%M"]))
                            horaSalida = horaSalida.split()[1]
                            horaSalida = horaSalida.split(':')[0]+':'+horaSalida.split(':')[1]
                            horaSalida = pd.to_datetime(('{} {}').format(dia,str(horaSalida)))
                            
                            if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] == pd.to_datetime(('{} 00:00').format(dia)):
                                frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] = horaSalida
                            else:
                                if frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] > pd.to_datetime(('{} 19:45').format(dia)):
                                    horaSalidaDeposito = frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],5] + pd.Timedelta(m=4)
                                    frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7] = horaSalida
                                    frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6] = horaSalidaDeposito
            
                    print(('Ingreso: {}').format(str(frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],6]).split()[1]))
                    print(('Salida: {}').format(str(frame.iloc[frame[frame['H.Norm.Emp'] == 0].index[x],7]).split()[1]))
                    print('-'*80)
                    print('\n')
       
        return frame
    except:
        logger.exception('',exc_info=True)             