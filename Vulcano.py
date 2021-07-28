import os
import logging.config
import traceback
import sys
import pandas as pd
import pyinputplus as pyip
import numpy as np
import win32com.client
import docx
import colorama


from paths import empleados_text,pathExcelTemporal,nombreExcelTemporal,pathExcelInforme,pathTXT,\
areas,formaDePago,rotativosInyeccion,rotativosSoplado,motivos
from analizador import Analizador,CalculadorHoras,informeNoFichadas,ingresoNoFichadas
from createDB import ManagerSQL
from queryes import (queryConsultaEmpleados,insertRegistros,selectAll,
                     selectSome,insertEmpleado,deleteEmpleado,
                     actualizarEmpleado,selectDeleteRegistro,updateRegistro,selectArea)
from Corrector import CorrectorExcel
from openpyxl import load_workbook
from time import sleep
from termcolor import colored
from datetime import timedelta
import datetime
import openpyxl
from openpyxl.styles import PatternFill
from openpyxl.worksheet.datavalidation import DataValidation

logging.config.fileConfig('logger.ini', disable_existing_loggers=False)
logger = logging.getLogger(__name__)


from paths import(nombreInformeNoFichadasWord,nombreInformeNoFichadasPDF,pathInformesNoFichadas,
pathInformesFaltasTardanzas,
nombreInformeFaltasTardanzasWord,
nombreInformeFaltasTardanzasPDF,
valoresListaDesplegable,toleranciaHoraria)

colorama.init()

def updateRegistroQuery(managerSQL,campo,fechaHora,legajo,fecha):
    """
    Method that executes a update into the REGISTROS table.

    Parameters
    ----------
    managerSQL : method
        Metodo de la clase ManagerSQL.
    campo : string
        Campo el cual se va a modificar.
    fechaHora : datetime.datetime
        fecha sobre la cual se va a hacer el cambio.
    legajo : string
        legajo del operario a cambiar.
    fecha : TYPE
        DESCRIPTION.

    Returns
    -------
    None.

    """
    
    queryInsercion = updateRegistro.format(campo,fechaHora,legajo,fecha)
    managerSQL.executeQuery(managerSQL.conexion(),queryInsercion)
    

def actualizaLinea(linea,cantidad):
    """
    

    Parameters
    ----------
    linea : TYPE
        DESCRIPTION.
    cantidad : integer
        cantidad de ingresos y egresos que posee la linea a modificar.

    Returns
    -------
    campo : TYPE
        DESCRIPTION.
    fechaHora : TYPE
        DESCRIPTION.

    """
    parametrosLista = ['ingreso0','egreso0',
                'ingreso1','egreso1',
                'ingreso2','egreso2',
                'ingreso3','egreso3',
                'ingreso4','egreso4']
    parametro = []
    for x in range(cantidad):
        parametro.append(parametrosLista[x])
        
    campo = pyip.inputMenu(parametro,prompt='¿Que campo desea actualizar ?\n',lettered=True)
    fecha = pyip.inputDate('Ingrese la fecha a actualizar DD/MM/AAAA: ',formats=["%d/%m/%Y"])
    hora = str(pyip.inputDatetime('Ingrese horario en formato HH:MM: ',formats=["%H:%M"]))                               

    hora = hora.split()[1]
    hora = hora.split(':')[0]+':'+hora.split(':')[1]
    fechaHora = pd.to_datetime(('{} {}').format(fecha,str(hora)))
    return campo,fechaHora    
    

def coloreadorRegistroModificar(linea,fecha):
    """
    Funcion que toma la linea, estblece cuantos ingresos y egresos hubo y a su vez los presenta coloreados.

    Parameters
    ----------
    linea : TYPE
        DESCRIPTION.
    fecha : datetime.datetime
        DESCRIPTION.

    Returns
    -------
    cantidad : integer
        idem above.

    """
    ingresoColor = colored('Ingreso: ','green',attrs=['bold','dark','underline']) #Se les asigna color y subrayado a las palabras ingreso-egreso
    egresoColor = colored('Egreso: ','red',attrs=['bold','dark','underline'])
    
    cero = pd.to_datetime(('{} 00:00').format(fecha))
    cantidad = 0
    for horario in range(4,12,2):#Itera sobre los registros con error y los imprime
                                
        ingreso = linea.iloc[0,horario]
        egreso = linea.iloc[0,horario +1]
                                   
        if egreso != cero:
            cantidad +=2 #suma los pares ingresos - egresos
            print(('    {} {}').format(ingresoColor,str(ingreso))+'      '+('{} {}').format(egresoColor,str(egreso))+'\n') 
        else:                                
            break
    return cantidad

def edicionRegistros(manager,query,fecha,legajo): 
    """
    Funcion que se encarga de hgacer la consulta para modificar el registro. No ejecuta la modificacion por si misma
    

    Parameters
    ----------
    manager : method
        metodo de la clase ManagerSQL.
    query : string
        Query con lo datos a corregir.
    fecha : Datetime.date
        DESCRIPTION.
    legajo : integer
        legajo del operario.

    Returns
    -------
    lineaCorregir : string
        Linea que contiene los valores a modificar.

    """
    
    columnas = ['fecha','ingreso0','egreso0',
                'ingreso1','egreso1',
                'ingreso2','egreso2',
                'ingreso3','egreso3',
                'ingreso4','egreso4']#Nombre las columnas desde la BD
    columnasReemplazo = {'legajo':'Legajo','nombre':'Nombre',
                         'dia':'Dia','fecha':'Fecha',
                         'ingreso0':'Ingreso_0','egreso0':'Egreso_0',
                         'ingreso1':'Ingreso_1','egreso1':'Egreso_1',
                         'ingreso2':'Ingreso_2','egreso2':'Egreso_2',
                         'ingreso3':'Ingreso_3','egreso3':'Egreso_3',
                         'ingreso4':'Ingreso_4','egreso4':'Egreso_4',
                         'Motivo':'Motivo','Observación':'Observación'}#Reemplapo para poder operar
    
    query = selectDeleteRegistro.format(fecha,legajo)# query select *
    lineaCorregir = pd.read_sql(query,manager.conexion(),parse_dates=columnas)
    lineaCorregir['fecha'] = pd.to_datetime(lineaCorregir['fecha']).dt.date #transforma 2020-10-01 00:00:00 a 2020-10-01    
    lineaCorregir.rename(columns = columnasReemplazo,inplace=True)
    lineaCorregir.drop(['id'],axis=1,inplace=True)
    return lineaCorregir

def agregadoListaDesplegable(pathToExcel,valores=[]):
    """
    Funcion que toma el excel creado en el ordenado registros y le incorpora la lista desplegable para establecer motivos.

    Parameters
    ----------
    pathToExcel : string
        path to the excel to modify.
    valores : list, optional
        Lista de los valores permitidos en la lista desplegable. The default is [].

    Returns
    -------
    None.

    """
    
    dv = DataValidation(type="list", formula1=valores,allow_blank=True)
    dv.error ='El motivo ingreso no esta permitido'
    dv.errorTitle = 'Motivo incorrecto'
    dv.prompt = 'Seleccione uno de los motivos siguientes'
    dv.promptTitle = 'Motivos'
    
    wb = openpyxl.load_workbook(pathToExcel)
    ws = wb.active
    ws.add_data_validation(dv)
    dv.add('O1:O1048576')
    wb.save(pathToExcel)

def fechasDeCalculo(completo=True):
    """
    Funcion que pide las fechas entre las cuales se va a operar, si el booleano esta por default pide feriados
    y medios dias.

    Parameters
    ----------
    completo : Bool, optional
        Booleano que se encarga de establecer si se va a pedir feriado o no. The default is True.

    Returns
    -------
    TYPE
        DESCRIPTION.

    """
    logger.info("Registrando fechas")
    fechaInicio = pyip.inputDate('Ingrese el primer dia habil DD/MM/AAAA: ',formats=["%d/%m/%Y"])
    print('\n')
    fechaFin = pyip.inputDate('Ingrese el ultimo dia habil  DD/MM/AAAA: ',formats=["%d/%m/%Y"])
    print('\n')
    if completo:
        feriados =[]
        feriado = None
        while feriado != '':
            feriado = pyip.inputDate(prompt='Ingrese el/los dias feriados DD/MM/AAAA: ',formats=["%d/%m/%Y"],blank=True)
            if feriado != '':
                feriados.append(feriado)
        print('\n')
        mediosDias =[]
        medioDia = None
        while medioDia != '':
            medioDia = pyip.inputDate(prompt='Ingrese el/los medios dias DD/MM/AAAA: ',formats=["%d/%m/%Y"],blank=True)
            if medioDia != '':
                mediosDias.append(medioDia)
        print('\n')
        return fechaInicio,fechaFin,feriados,mediosDias
    else:
        return fechaInicio,fechaFin


def ingreso_egreso(line,frame,legajo,nombre):
    """
    

    Parameters
    ----------
    line : str
        Linea del TXT que se esta analizando.
    frame : DataFrame
        Frame sobre el cual se estan agregando los registros una vez limpios.
    legajo: str 
        Legajo del operario bajo limpieza.
    nombre : str
        Nombre del operario bajo limpieza.

    Returns
    -------
    frame : DataFrame
        Dataframe con los datos agregados.

    """
    try:
        fecha = pd.to_datetime(line.split()[1],yearfirst=True,format='%d/%m/%Y').date()#Da formato datetime.date
        ingresos_egresos = []
        for x in range(2,12,1):#Esta parte se encarga de dar formato al horario, siempre con dia y horas.
            try:
                a = pd.to_datetime(('{} {}').format(fecha,line.split()[x].split('-')[0]))
                ingresos_egresos.append(a)            
            except:
                a = pd.to_datetime(('{} 00:00').format(fecha))
                ingresos_egresos.append(a)
    

        lista_final = []
        lista_final.append(ingresos_egresos[0])
            
        for indice in range(len(ingresos_egresos) -1):#Elimina registros duplicados,usa como limite unos 7 minutos.            
            if ingresos_egresos[indice + 1] - ingresos_egresos[indice] > pd.Timedelta(minutes=7):
               lista_final.append(ingresos_egresos[indice + 1])
           
        for u in range((10-len(lista_final))):
            lista_final.append(pd.to_datetime(('{} 00:00').format(fecha)))#Ponee en 00:00:00 los otros pares de ingreso-ergreso

        frame = frame.append({'Legajo':legajo,'Nombre':nombre,'Dia':line.split()[0],'Fecha':fecha,
                                  'Ingreso_0':lista_final[0],'Egreso_0':lista_final[1],
                                  'Ingreso_1':lista_final[2],'Egreso_1':lista_final[3],
                                  'Ingreso_2':lista_final[4],'Egreso_2':lista_final[5],
                                  'Ingreso_3':lista_final[6],'Egreso_3':lista_final[7],
                                  'Ingreso_4':lista_final[8],'Egreso_4':lista_final[9]},
                                  ignore_index=True) #agrega una fila mas al dataframe con los datos de ese dia
                                
        return frame
    except Exception:
        logger.error("excepcion desconocida: %s", traceback.format_exc())
        return None

def creacionFrameVacio():
    """

    Returns
    -------
    frame : DataFrame
        Dataframe vacio con las columnas especificadas.

    """
    columnas = ['Legajo','Nombre','Dia','Fecha','Ingreso_0','Egreso_0','Ingreso_1','Egreso_1',
                        'Ingreso_2','Egreso_2','Ingreso_3','Egreso_3','Ingreso_4','Egreso_4',
                        'Motivo','Observación']
    frame = pd.DataFrame(columns=columnas)
    return frame

def empleadosFrame():
    """
    

    Returns
    -------
    legajosNoRotativos : list
        Lista con los legajos que son no rotativos.
    frameQuerido : DataFrame
        DataFrame con los datos todos los datos de los operarios.

    """
    
    manager = ManagerSQL()
    sql_conection = manager.conexion()
    consultaEmpleados = pd.read_sql(queryConsultaEmpleados,sql_conection)

    
    legajosNoRotativos = consultaEmpleados.loc[(consultaEmpleados['AREA'] != 'INYECCION') 
                                                &(consultaEmpleados['AREA'] != 'SOPLADO') 
                                                & (consultaEmpleados['AREA'] != 'MECANIZADO')
                                                & (consultaEmpleados['AREA'] != 'ALUMINIO')] #Filtrado con las areas que NO son rotativas
    legajosNoRotativos = legajosNoRotativos['LEG'].unique()
    legajosNoRotativos = [int(x) for x in legajosNoRotativos]#los pasa de numpy.int64 a int

    frameQuerido =pd.DataFrame(consultaEmpleados.loc[:,['LEG','AREA']].drop_duplicates().values,columns=['LEG','AREA'])
    frameQuerido.set_index(['LEG'],inplace=True)
    
    return legajosNoRotativos,frameQuerido
    

def insercionBD(managerSQL,frame,query):
    """
    Funcion que se encarga de subir los datos a la BD.

    Parameters
    ----------
    managerSQL : method
        metodo de la clase ManagerSQL.
    frame : DataFrame
        Frame with the records to upload to the DB.
    query : str
        Query with all the data needed for upload 1 line.

    Returns
    -------
    None.

    """
    logger.info('Insertando registros en la BD')
    for index, row in frame.iterrows():
        queryInsercion = query.format(row['Legajo'],row['Nombre'], row['Dia'],str(row['Fecha']),
          str(row['Ingreso_0']),str(row['Egreso_0']),
          str(row['Ingreso_1']),str(row['Egreso_1']),
          str(row['Ingreso_2']),str(row['Egreso_2']),
          str(row['Ingreso_3']),str(row['Egreso_3']),
          str(row['Ingreso_4']),str(row['Egreso_4']),
          str(row['Motivo']),str(row['Observación']))
        logger.info(queryInsercion)
        managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 

def insercionBDLegajos(managerSQL,legajo,nombre,apellido,area,pago,query):

    logger.info('Insertando nuevo operario en la BD')
    queryInsercion = query.format(legajo,nombre,apellido,area,pago)
    logger.info(queryInsercion)
    managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 

def deleteBDLegajos(managerSQL,legajo,query):
    logger.info('Borrando operario en la BD')
    queryInsercion = query.format(legajo)
    logger.info(queryInsercion)
    managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 
    
def actualizaBDLegajos(managerSQL,legajo,campo,valor,query):
    msg = 'Actulizando operario legajo {} en la BD, campo actualizado {}, valor nuevo {}'.format(legajo,campo,valor)
    logger.info(msg)
    queryInsercion = query.format(campo,valor,legajo)
    logger.info(queryInsercion)
    managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 
        

def frameFichadas():
    """
    Funcion que se encarga de tomar el TXT, e ir leyendo linea a linea, la secuencia es:
        
        1.- Crea un dataFrame vacio donde va a append las lineas.
        2.- Teniendo en cuenta el pathTXT (variable), busca los archivos en esa carpeta 
        3.- Lee linea a linea el archivo y los pasa a la funcion que los limpia(ingreso_egreso).

    Returns
    -------
    frame : datFrame
       dataFrame con todos los horarios ya limpios y sin duplicados.

    """
    
    semaine = {'Lunes': 0,
               'Martes' : 1,
               'Miércoles': 2,
               'Jueves':3,
               'Viernes':4,
               'Sábado':5,
               'Domingo':6,
              }
    
    frame = creacionFrameVacio()
    MedioDia = None
    try:
        listaTXT = os.listdir(pathTXT)
    except FileNotFoundError:
        logger.warning('Esta mal el path que lleva al texto')
        listaTXT = []
        
    if len(listaTXT) == 0:        
        return frame
    elif len(listaTXT) == 1:
        archivo = listaTXT[0]
    else:
        archivo = pyip.inputMenu(listaTXT,prompt='Elija uno de los archivos:\n',lettered=True)
    
    archivo = os.path.join(os.getcwd(),pathTXT,archivo)
    print('\n')
    try:
         with open(archivo,encoding="utf-8") as file:
                    for line in file.readlines():
                        if line.startswith('Empleado'):
                            legajo = line.split()[1].replace('.',"")
                            nombre = ''
                            for x in range(3,7):
                                if line.split()[x] == 'Tarjeta':
                                    break
                                else:
                                    nombre += line.split()[x]+' '
                            nombre = nombre.upper()
                        for jour in semaine.keys():        
                            if line.startswith(jour):
                                frame = ingreso_egreso(line,frame,legajo,nombre)
    except:
        with open(archivo) as file:
            for line in file.readlines():
                        if line.startswith('Empleado'):
                            legajo = line.split()[1].replace('.',"")
                            nombre = ''
                            for x in range(3,7):
                                if line.split()[x] == 'Tarjeta':
                                    break
                                else:
                                    nombre += line.split()[x]+' '
                            nombre = nombre.upper()
                        for jour in semaine.keys():        
                            if line.startswith(jour):
                                frame = ingreso_egreso(line,frame,legajo,nombre)
   

    return frame

def logicaRotativos(frame,fechaInicio,fechaFin,legajo,area=False):
    """
    Esta funcion se encarga de limpiar los registros de los dataframes individuales de cada
    empleado rotativo, aqui dentro esta toda la logica de limpieza de esos registros.

    Parameters
    ----------
    frame : DataFrame
        Dataframe con los registros de un empleado Rotativo.

    fechaInicio : datetime.date
        Fecha de inicio para el analisis.
    fechaFin : TYPE
        DESCRIPTION.

    Returns
    -------
    frame : DataFrame
        DataFrame ya corregido. Los lugares vacios corresponden a faltas en los fichajes.

    """
    
    fechaInicioAyer = fechaInicio - timedelta(days=1)
    fechaFinAyer = fechaFin + timedelta(days=1)
    mascara = (frame['Fecha'] >= fechaInicioAyer) & (frame['Fecha'] <= fechaFinAyer) #mascara para filtrar el frame en funcion de la fecha de inicio y fin
    frameEnAnalisis = frame.loc[mascara].copy()
    limpiador = Analizador(frameEnAnalisis=frameEnAnalisis,fechaInicio = fechaInicio,fechaFin = fechaFin)
    estado = limpiador.sanityCheck()
   # print('\nIngreso: ',frame,frame.iloc[1,4],frame.iloc[1,5])
      
    if estado:
 
        newFrame = limpiador.castMascara(area=area) 
        
 
        newFrame = limpiador.limpiador(area=area)
        

        newFrame = limpiador.castMascara(area=area)
        
       
        mascaraNewFrame = (newFrame['Fecha'] >= fechaInicio) & (newFrame['Fecha'] <= fechaFin)
        newFrame = newFrame.loc[mascaraNewFrame].copy()

    else:
        msg = 'El siguiente legajo ({}) No paso las validaciones,se autocompletan los registros como 00:00'.format(legajo)
        print(msg,'\n')
        msg2 ='El siguiente legajo ({}) No paso las validaciones, se las castea en cero.'.format(legajo)
        logger.info(msg2)
        newFrame = limpiador.castMascara(area=area)
        newFrame = limpiador.limpiador(area=area)
        newFrame = limpiador.castMascara(area=area)
        print('\nNOGOOD: ',newFrame,newFrame.iloc[1,4],newFrame.iloc[1,5])   
        mascaraNewFrame = (newFrame['Fecha'] >= fechaInicio) & (newFrame['Fecha'] <= fechaFin)
        newFrame = newFrame.loc[mascaraNewFrame].copy()

   
    return newFrame 
def coloreadorExcel(pathExcel):
    """
    

    Parameters
    ----------
    pathExcel : str
        path to the excel to colored.

    Returns
    -------
    None.

    """
    try:
        wb = openpyxl.load_workbook(pathExcel)
    except PermissionError:
        print('El archivo {} esta abierto, tiene 30 segundos para cerrarlo.').format(pathExcel)
        sleep(30)
        wb = openpyxl.load_workbook(pathExcel)
    
    ws = wb.active
    my_red = openpyxl.styles.colors.Color(rgb='00FF0000')
    my_fill = openpyxl.styles.fills.PatternFill(patternType='solid', fgColor=my_red)
    my_green = openpyxl.styles.colors.Color(rgb="0000FF00")
    my_fillGreen = openpyxl.styles.fills.PatternFill(patternType='solid', fgColor=my_green)
    
    for indice,rows in enumerate(ws.iter_rows(min_row=2, min_col=17,max_col=17)): #itera sobre la columna H.Norm
        for celda in rows:#Por cada celda en la fila
            if celda.value == 0:
                valor = indice + 2 # valor de la fila desde el princio del archivo, +2 (+1 por las columnas, y +1 porque excel arranca a contar desde 1 y no 0)
                fecha = ws.cell(row=valor, column=4).value
                columnaCuatro = ws.cell(row=valor, column=5).value
                columnaCinco = ws.cell(row=valor, column=6).value
                cero = pd.to_datetime(('{} {}').format(fecha,'00:00'))

                if columnaCuatro == cero and columnaCinco == cero:
                    for idx in ws.iter_cols(min_row=valor,max_row = valor, min_col=1,max_col=19):
                        for celda in idx:
                            celda.fill = my_fillGreen
                else:
                    for idx in ws.iter_cols(min_row=valor,max_row = valor, min_col=1,max_col=19):
                        for celda in idx:
                            celda.fill = my_fill


    wb.save(pathExcel)
    wb.close()

def frameAnalisisIndividual(frame,fechaInicio,fechaFin):
    """
    Esta funcion se encarga de tomar el frame completo de todas las horas e ir armando 1 frame por cada empleado
    y si esta dentro de los rotativos encara con una logica distinta.    
    En caso de ser normal y no fichar salida o ingreso queda el espacio vacio.    
    Por ultima genera un dataframe totalizando todo

    Parameters
    ----------
    frame : datFrame
        dataFrame con todas las horas ya limpias.
    fechaInicio : datetime.date
        fecha de incio de la secuencia que se quiere limpiar.
    fechaFin : datetime.date
        fecha de finalizacion de la secuencia que se quiere limpiar.

    Returns
    -------
    frameAnalisis : dataFrame
        dataFrame con solamente los legajos de los operarios que no son rotativos.

    """
    legajosNoRotativos,frameQuerido = empleadosFrame()
    
   

    legajos = frame['Legajo'].unique()
    frameAnalisis = creacionFrameVacio()
    frameRechazados = creacionFrameVacio()

    for legajo in legajos:

        newFrame = frame[frame['Legajo']==legajo]#legajo es un STRaa
        try:
            area = frameQuerido.loc[int(legajo),'AREA']
        except:
            msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
            logger.warning(msg)
            print(msg,'\n','Se procede a obviar dicho empleado y se prosigue con el resto.\n')
            continue
        try: 
            if int(legajo) in legajosNoRotativos:
                               
                mascara = (newFrame['Fecha'] >= fechaInicio) & (newFrame['Fecha'] <= fechaFin) #mascara para filtrar el frame en funcion de la fecha de inicio y fin
                newFrame = newFrame.loc[mascara].copy()
                limpiador = Analizador(frameEnAnalisis=newFrame,fechaInicio = fechaInicio,fechaFin = fechaFin)
                newFrame = limpiador.castMascara(area)
                mascara = (newFrame['Fecha'] >= fechaInicio) & (newFrame['Fecha'] <= fechaFin)
                newFrame = newFrame.loc[mascara].copy()
                newFrame = newFrame.sort_values(by=['Legajo','Fecha'])
                frameAnalisis = frameAnalisis.append(newFrame)
            else:
                newFrame = logicaRotativos(newFrame,fechaInicio,fechaFin,legajo=legajo,area=area)
                newFrame = newFrame.sort_values(by=['Legajo','Fecha'])
                frameAnalisis = frameAnalisis.append(newFrame)
                
        except Exception as e:
             msg = 'El siguiente legajo ({}) No paso en sanitycheck, ni pudo ser casteado'.format(legajo)
             print(msg,'\n','Se procede a obviar dicho empleado y se prosigue con el resto.\n')
             frameRechazados = frameRechazados.append(frame[frame['Legajo']==legajo])
             logger.warning(msg)
             logger.error(e)



    nombreRechazados = 'Rechazados {} al {}.xlsx'.format(fechaInicio,fechaFin)
    if not frameRechazados.empty:
        pathRechazados = os.path.join(os.getcwd(),pathExcelTemporal,'Rechazados',nombreRechazados)
        frameRechazados.to_excel(pathRechazados,index=False)
    return frameAnalisis

def limpiezaDeRegistros(frame,fechaInicio,fechaFin):
    """
    

    Parameters
    ----------
    frame : DataFrame
        DataFrame con todos los registros sobre los cuales se va a trabajar.
    fechaInicio : datetime.date
        DESCRIPTION.
    fechaFin : datetime.date
        DESCRIPTION.

    Returns
    -------
    len_noMarca : int
        cantidad de registros que estan en 0 (aun poseen errores).

    """
    try:    
        logger.info('Limpiando registros')
        manager = ManagerSQL()
        legajosNoRotativos,frameQuerido = empleadosFrame()
    
        legajosFrame = frame['Legajo'].unique()

        frameAnalisis = creacionFrameVacio()
        calculador = CalculadorHoras()
        for legajo in legajosFrame:
            try:
                area = frameQuerido.loc[int(legajo),'AREA']
            except:
                msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
                logger.warning(msg)
                print(msg,'\n','Se procede a obviar dicho empleado y se prosigue con el resto.\n')
                continue
            newFrame = frame[frame['Legajo']==legajo].copy() 
            if int(legajo) in legajosNoRotativos:                           
                frameCalculado = calculador.horasTrabajadas(newFrame)        
                frameAnalisis = frameAnalisis.append(frameCalculado)
            else:
                frameCalculado = calculador.horasTrabajadasRotativos(newFrame,area)
                frameAnalisis = frameAnalisis.append(frameCalculado)

                
        frameAnalisis = frameAnalisis.reset_index(drop=True) 
        informeNoFichadas(frameAnalisis,fechaInicio,fechaFin,mediosDias=[],feriados=[])#Crea el informe de no fichadas
        campo = 'H.Norm' #campo sobre el cual se filtra para ver las filas que tienen errores en los registros. Es siempre el mismo
        len_noMarca = len(frameAnalisis[frameAnalisis[campo] == 0])
        
        if len_noMarca == 0:
            msgOk = colored('TODOS los registros se encuentran completos, se procede a escribir la Base de datos.','grey',on_color='on_green')
            print(msgOk)
            insercionBD(manager,frameAnalisis,insertRegistros)
        else:
            nombre = nombreExcelTemporal.format(str(fechaInicio).replace('/','-'),str(fechaFin).replace('/','-'))
            nombre = os.path.join(os.getcwd(),pathExcelTemporal,nombre)
            frameAnalisis['Motivo'] = ''
            frameAnalisis['Observación'] = ''
            query = selectArea
            frameConAreas = pd.read_sql(query,manager.conexion()) # genere el frame con legajo,area
            frameConAreas =  frameConAreas.rename(columns={'LEG':'Legajo'})
            frameConAreas['Legajo'] = frameConAreas['Legajo'].astype(str) # pasa ambos a str para poder mergear
            frameAnalisis['Legajo'] = frameAnalisis['Legajo'].astype(str)
            frameAnalisis = frameAnalisis.merge(frameConAreas,on='Legajo',how='left') #agregado de la columna extra al frame para ver el area.


            try:
                frameAnalisis.to_excel(nombre,index=False)
                coloreadorExcel(nombre)
                agregadoListaDesplegable(nombre,valores=valoresListaDesplegable)
            except PermissionError:
                msg = '--El archivo {} esta abierto, tiene 30 segundos para cerrarlo.'.format(nombre)
                print(msg)
                sleep(30)
                frameAnalisis.to_excel(nombre,index=False)
                coloreadorExcel(nombre)
                agregadoListaDesplegable(nombre,valores=valoresListaDesplegable)
        return len_noMarca
                
    except:
        logger.error("excepcion desconocida: %s", traceback.format_exc())
    
        
def analizadorFramesCorregidos(frame,fechaInicio,fechaFin):
    """
    Funcion que verifica los errores que puede haber en los registros, en caso de que no los halla, los sube a la BD.
    En caso de haber por lo menos 1 error devuelve el frame.

    Parameters
    ----------
    frame : DataFrame
        Dataframe con los registros a trabajar.
    fechaInicio : datetime.date
        DESCRIPTION.
    fechaFin : datetime.date
        DESCRIPTION.

    Returns
    -------
    frameConErrores : DataFrame
        Dataframe con errores en elaugnos de los registros por cada operario.

    """
    
    manager = ManagerSQL()
    
    legajosNoRotativos,frameQuerido = empleadosFrame()
    
    calculador = CalculadorHoras() 
    
    legajos = frame['Legajo'].unique()
    frameConErrores = creacionFrameVacio()
    campo = 'H.Norm'#campo sobre el cual se filtra para ver las filas que tienen errores en los registros. Es siempre el mismo
    for legajo in legajos:
        try:
            area = frameQuerido.loc[int(legajo),'AREA']
        except:
            msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
            logger.warning(msg)
            continue

        newFrame = frame[frame['Legajo']==legajo].copy() 
        if int(legajo) in legajosNoRotativos:                           
            newFrame = calculador.horasTrabajadas(newFrame)        
        else:
            newFrame = calculador.horasTrabajadasRotativos(newFrame,area)
        

        len_noMarca = len(newFrame[newFrame[campo] == 0])
        
        Ingreso_0 = list(newFrame['Ingreso_0'])
        Egreso_0 = list(newFrame['Egreso_0'])
        fecha = list(newFrame['Fecha'])

        cuenta = 0
        for x in range(len(Ingreso_0)):
            ceroHoy = pd.to_datetime(('{} 00:00').format(fecha[x]))
            if Ingreso_0[x] == ceroHoy and Egreso_0[x] == ceroHoy:
                cuenta +=1

        len_noMarca = len_noMarca - cuenta
        mascara = (frame['Fecha'] >= fechaInicio) & (frame['Fecha'] <= fechaFin)
        if len_noMarca == 0 and any(mascara): 
            newFrame = newFrame.loc[mascara].copy()
            insercionBD(manager,newFrame,insertRegistros)#Escribe la BD
        else:
            frameConErrores = frameConErrores.append(newFrame)
    return frameConErrores
            
            
    
    
    
def actualizacionRegistros(fechaInicio,fechaFin):
    """
    Funcion que se encarga de tomar el frame que se especifico antes y lo revisa, aplica las funciones para ver si existen registros con errores
    y en caso de que no borra el archivo seleccionado (ya que todo estaba bien).

    Parameters
    ----------
    fechaInicio : datetime.date
        DESCRIPTION.
    fechaFin : datetime.date
        DESCRIPTION.

    Returns
    -------
    None.

    """
    
    try:
        try:
            listaExcels = os.listdir(pathExcelTemporal)
        except FileNotFoundError:
            logger.warning('Esta mal el path que lleva al texto de la actualizacion del Excel')
            print('No existen archivos que actualizar')
            return

        
        listaExcels = os.listdir(pathExcelTemporal)
        listaCodigosClean = [x for x in listaExcels if 'Excel' in x] #solo va a incluir aquellos archivos que empiezen con Excel
        if len(listaCodigosClean) == 1:
            archivo = listaCodigosClean[0]
            msg = 'Se procede a utilizar el unico archivo en la carpeta\n{}'.format(str(listaCodigosClean[0]))
            print(msg)
        else: 
            archivo = pyip.inputMenu(listaCodigosClean,prompt='Elija uno de los archivos:\n',lettered=True)
        print('\n')
        
        
        # nombre = nombreExcelTemporal.format(str(fechaInicio).replace('/','-'),str(fechaFin).replace('/','-'))
        nombre = os.path.join(os.getcwd(),pathExcelTemporal,archivo)  
        corrector = CorrectorExcel(abs_path_excel=nombre)
        frameAnalisis = corrector.corrector_marcada()
        # frameAnalisis = pd.read_excel(nombre)
        # frameAnalisis['Fecha'] = pd.to_datetime(frameAnalisis['Fecha']).dt.date #transforma 2020-10-01 00:00:00 a 2020-10-01    

        frameConErrores = analizadorFramesCorregidos(frameAnalisis,fechaInicio,fechaFin)     

        if frameConErrores.empty:
            msgOk = colored('Todos los registros corregidos, se actualizo la base de datos.','grey',on_color='on_green')
            print(msgOk)
            os.remove(nombre)
        else:
            print('Aun persisten registros con errores, revisarlos y corregirlos para poder proceder.')
            frameConErrores.to_excel(nombre,index=False)
            coloreadorExcel(nombre)
            agregadoListaDesplegable(nombre, valores = valoresListaDesplegable)
        
    except:
        logger.error("excepcion desconocida: %s", traceback.format_exc())


def validador(eleccion):
    """
    

    Parameters
    ----------
    eleccion : str
        nombre del archivo a guardar.

    Raises
    ------
    Exception
        Excepcon para que pyinputplus arroje ese error.

    Returns
    -------
    None.

    """
    for letter in eleccion:
        if letter in ["\\","/",":","*","?","<",">","|"]:
            raise Exception('\nExisten caracteres que no pueden estar incluidos en el nombre\n\n')

def calculosExtrasRotativos(frame):
    """
    Funcion que se encarga de contnar las noche sy tardes que los oeprarios con turno nocturno han trabajdo.

    Parameters
    ----------
    frame : dataFrame
        DataFrame con todos los registros.

    Returns
    -------
    empleados : dict
        Dicciones con las noches y tardes trabajadar por lo operarios rotativos.

    """
    
    legajosNoRotativos,frameQuerido = empleadosFrame()

    legajos = frame['Legajo'].unique()
    
    empleados = {}
    for legajo in legajos:
            empleados[str(legajo)]={}
            empleados[str(legajo)]['Noches'] = 0
            empleados[str(legajo)]['Tardes'] = 0

    for legajo in legajos:
        try:
            area = frameQuerido.loc[int(legajo),'AREA']
        except:
            msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
            logger.warning(msg)
            continue
        newFrame = frame[frame['Legajo']==legajo]#legajo es un STRaa
        if int(legajo) in legajosNoRotativos:
             continue
        else:
            for fila in range(len(newFrame)):
                ingresoOperario = newFrame.iloc[fila,4]
                fecha = newFrame.iloc[fila,3]
                ayer = fecha - timedelta(days=1)
                cero = (pd.to_datetime(('{} 00:00').format(fecha)))
                
                if area in rotativosInyeccion:
                    primerIngreso = '08:00'
                    segundoIngreso = '16:00'
                    tercerIngreso = '00:00'
                    
                    turnoMañanaIngresoInyeccion = (pd.to_datetime(('{} {}').format(fecha,primerIngreso)))
                    turnoTardeIngresoInyeccion = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso)))
                    turnoNocheIngresoInyeccion = (pd.to_datetime(('{} {}').format(fecha,tercerIngreso)))
                    
                    if ingresoOperario == cero and newFrame.iloc[fila,5] == cero:
                        continue
                    
                    if  turnoMañanaIngresoInyeccion < ingresoOperario < turnoMañanaIngresoInyeccion + timedelta(hours=3) :
                        continue
                            
                    if  turnoTardeIngresoInyeccion - timedelta(hours=3) < ingresoOperario < turnoTardeIngresoInyeccion + timedelta(hours=2) :
                        empleados[str(legajo)]['Tardes'] += 1
                            
                    if  turnoNocheIngresoInyeccion - timedelta(hours=5) < ingresoOperario < turnoNocheIngresoInyeccion + timedelta(hours=3) :
                        empleados[str(legajo)]['Noches'] += 1
                    
                    
                else:
                     primerIngresoSoplado = '07:00'
                     segundoIngresoSoplado = '15:00'
                     tercerIngresoSoplado = '23:00'
                        
                     turnoMañanaIngresoSoplado = (pd.to_datetime(('{} {}').format(fecha,primerIngresoSoplado)))
                     turnoTardeIngresoSoplado = (pd.to_datetime(('{} {}').format(fecha,segundoIngresoSoplado)))
                     turnoNocheIngresoSoplado = (pd.to_datetime(('{} {}').format(ayer,tercerIngresoSoplado)))

                     if ingresoOperario == cero and newFrame.iloc[fila,5] == cero:
                        continue
                     
                     if  turnoMañanaIngresoSoplado < ingresoOperario < turnoMañanaIngresoSoplado + timedelta(hours=3) :
                        continue
                            
                     if  turnoTardeIngresoSoplado - timedelta(hours=3) < ingresoOperario < turnoTardeIngresoSoplado + timedelta(hours=2) :
                        empleados[str(legajo)]['Tardes'] += 1
                            
                     if  turnoNocheIngresoSoplado- timedelta(hours=5) < ingresoOperario < turnoNocheIngresoSoplado + timedelta(hours=3) :
                        empleados[str(legajo)]['Noches'] += 1
                    

    return empleados

def hojaTotalizadora(frame,fechaInicio,fechaFin,feriados,empleados,empleadosExtras):
    """
    Funcion que se encarga de tomar el dataframe desde la BD que ya tiene calculadas las H.Norm. H.50 y H.100
    y totaliza por cada empleado las horas trabajadas y a su vez cuenta los dias que trabajo. Los dias se calculan contanto en el dataframe.
    

    Parameters
    ----------
    frame : dataframe
        dataFrame que posee todos los datos de los empleados juntos con las horas trabajadas.
    fechaInicio : datetime.date
        fechaInicio sobre la cual se esta analizando.
    fechaFin : datetime.date
        fechaFin sobre la cual se analiza.
    feriados : list
        Lista que tiene dentro los dias entre fechaInicio y fechaFin que son feriados.

    Returns
    -------
    frameConcatenado : dataFrame
        Devuelve un dataFrame con una fila por cada empleado en donde se encarga de totalizar todas las horas y los
        dias trabajados.

    """
    frameDeTrabajo = frame.copy()
    fechaInicio = pd.to_datetime(fechaInicio).date()
    fechaFin = pd.to_datetime(fechaFin).date()
    diasLaborales = list(pd.bdate_range(fechaInicio,fechaFin))

    
    diasLaborales = len(list(pd.bdate_range(fechaInicio,fechaFin)))#calcula los dias habiles entre 2 fechas
    diasLaborales = diasLaborales - len(feriados) #Al total de dias habiles le resta los feriados
    frame = frame.groupby(["Legajo","Nombre"])#agrupa el frame entre legajo y nombre para asi poder totalizar
    frameLegajo = frame.sum() #Suma todas las columnas numericas [H.Norm,H50,H100]
    frameLegajo.reset_index(inplace=True)
    
    frameBoleano = frameDeTrabajo[(frameDeTrabajo['H.Norm'] > 0) | (frameDeTrabajo['H. 50'] > 0) | (frameDeTrabajo['H. 100'] > 0)]
    frameBoleano = frameBoleano.groupby(["Legajo","Nombre"])
    
    columnas = ['Legajo','Nombre','Hs. Noche','Hs. Tarde','Mins. Tardanzas']
    frameAdicional = pd.DataFrame(columns=columnas)


    frameDiasLaborales = frameBoleano[['Dia']].count()


    for key in empleados.keys():
        minutos = 0

        for llave,valor in empleados[str(key)]['Tardanzas'].items():
            minutos += valor[1]
        frameAdicional = frameAdicional.append({'Legajo':int(key),'Nombre': empleados[str(key)]['Nombre'],'Hs. Noche':empleadosExtras[key]['Noches'] * 0.8,
                                                'Hs. Tarde':empleadosExtras[key]['Tardes'],
                                                'Mins. Tardanzas':minutos},
                                               ignore_index=True)
        
        
    frameAdicional.reset_index(drop=True, inplace=True)
    frameConcatenado = pd.merge(frameLegajo,frameDiasLaborales,on='Legajo') #crea una nuevo Frame juntando dias como horas.
    frameConcatenado = pd.merge(frameConcatenado,frameAdicional,on='Legajo')
    frameConcatenado = frameConcatenado.rename(columns={'Dia':'Dias Trabajados','Nombre_x':'Nombre'})
    frameConcatenado = frameConcatenado.drop('Nombre_y',axis=1)
    
    return frameConcatenado
    
    
def agregadoColumnas(frame):
    frame['Tardanzas'] = 0
    frame['Retiros'] = 0
    return frame

def retTarRotativos(frame,legajo,medioDias):
    """
    Funcion que se encarga de contar los retrasos y retiros anticipados por cada registro de cada empleado

    Parameters
    ----------
    frame : DataFrame
        Frame con todos los registros necesarios.
    legajo : int
        legajo del operario.
    medioDias : list
        Lista de dias que son medios dias.

    Returns
    -------
    frame : TYPE
        DESCRIPTION.

    """
    legajosNoRotativos,frameQuerido = empleadosFrame()
    frame = frame.sort_values(by=['Legajo','Fecha'])
    

    try:
        area = frameQuerido.loc[int(legajo),'AREA']
    except:
        msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
        logger.warning(msg)
        
    toleranciaHoraria = 1

    if area in rotativosInyeccion:
                    for x in range(len(frame)):
                        legajo = frame.iloc[x,0]
                        nombre = frame.iloc[x,1]
                        dia = frame.iloc[x,2]
                        fecha = frame.iloc[x,3]
                        ayer = fecha - datetime.timedelta(days=1)
                        mañana = fecha + datetime.timedelta(days=1)
                        
                        primerIngreso = '08:00'
                        segundoIngreso = '16:00'
                        tercerIngreso = '00:00'
                        
                        primerSalida = '16:00'
                        segundaSalida = '00:00'
                        tercerSalida = '08:00'
                        
                        horaSalidaSabado = '13:00'
                        medioDia = '12:30'
                        
                        horaSalidaMedioDia = pd.to_datetime(('{} 12:30').format(fecha))

                        turnoMañanaPrimerIngreso = (pd.to_datetime(('{} {}').format(fecha,primerIngreso)))
                        turnoTardeIngreso = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso)))
                        turnoNocheIngreso = (pd.to_datetime(('{} {}').format(fecha,tercerIngreso)))
                            
                            
                        turnoMañanaPrimerSalida = (pd.to_datetime(('{} {}').format(fecha,primerSalida)))
                        turnoTardeSalida = (pd.to_datetime(('{} {}').format(mañana,segundaSalida)))
                        turnoNocheSalida = (pd.to_datetime(('{} {}').format(fecha,tercerSalida)))  
                            
                            
                        cero = pd.to_datetime(('{} 00:00').format(fecha))
                        medioDia = pd.to_datetime(('{} 12:30').format(fecha))
                        salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                        
                        tardanza = 0
                        retiro= 0
                        
                        ingresoOperario = frame.iloc[x,4]
                        for idx in range(5,13,2):
                            if frame.iloc[x,idx] == cero:
                                salida = frame.iloc[x,idx -2]
                                break
                           
                        if frame.iloc[x,2] != 'Sábado':
                            if  turnoMañanaPrimerIngreso < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoMañanaPrimerIngreso).seconds)/60),2)
                                
                            if  turnoTardeIngreso < ingresoOperario < turnoTardeIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoTardeIngreso).seconds)/60),2)
                                
                            if  turnoNocheIngreso < ingresoOperario < turnoNocheIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoNocheIngreso).seconds)/60),2)
                            
                            
                            
                            
                            
                        if fecha in medioDias:
                            if  turnoMañanaPrimerIngreso < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours=3) :
                                if salida < horaSalidaMedioDia:
                                    retiro = round((((horaSalidaMedioDia - salida).seconds)/60),2)                        
                            
                        if fecha not in medioDias:
                            if  turnoMañanaPrimerSalida - timedelta(hours=3)  < salida < turnoMañanaPrimerSalida:
                                if salida < turnoMañanaPrimerSalida and dia != 'Sábado':
                                    retiro = round((((turnoMañanaPrimerSalida - salida).seconds)/60),2)
                                        
                            if  turnoTardeSalida - timedelta(hours=3)  < salida < turnoTardeSalida:
                                if salida < turnoTardeSalida:
                                    retiro = round(((( turnoTardeSalida - salida).seconds)/60),2)
                                        
                            if  turnoNocheSalida - timedelta(hours=3)  < salida < turnoNocheSalida:
                                if salida < turnoNocheSalida:
                                    retiro = round(((( turnoNocheSalida - salida).seconds)/60),2)
                                        
                                    
            
                        if tardanza != 0:                    
                            frame.iloc[x,19] = tardanza
                        if retiro != 0:
                            frame.iloc[x,20] = retiro

    return frame
    
def seleccionInformes(fechaInicio,fechaFin,mediosDias=[],feriados=[]):  
    """
    Funcion que se encarga de realizar una query sobre la BD, traer los datos, ordenar los types
    y luego realizar el calcluo de horas trabajadas: normales,extras50 y extras100.

    Parameters
    ----------
    fechaInicio : datetime.date
        fechaInicio sobre la cual se esta analizando.
    fechaFin : datetime.date
        fechaFin sobre la cual se analiza.
    mediosDias : List, optional
        DESCRIPTION. The default is [].lista con mediosDias entre las fechas de analisis
    feriados : List, optional
        DESCRIPTION. The default is [].lista con feriados entre las fechas de analisis.

    Returns
    -------
    frameCorregido : dataFrame
        dataFrame que ya posee todas las horas calculadas por cada operario.

    """
    
    
    columnas = ['fecha','ingreso0','egreso0',
                'ingreso1','egreso1',
                'ingreso2','egreso2',
                'ingreso3','egreso3',
                'ingreso4','egreso4']
    columnasReemplazo = {'legajo':'Legajo','nombre':'Nombre',
                         'dia':'Dia','fecha':'Fecha',
                         'ingreso0':'Ingreso_0','egreso0':'Egreso_0',
                         'ingreso1':'Ingreso_1','egreso1':'Egreso_1',
                         'ingreso2':'Ingreso_2','egreso2':'Egreso_2',
                         'ingreso3':'Ingreso_3','egreso3':'Egreso_3',
                         'ingreso4':'Ingreso_4','egreso4':'Egreso_4',
                         'Motivo':'Motivo','Observación':'Observación'}
    calculador = CalculadorHoras() 
    manager = ManagerSQL()
    sql_conection = manager.conexion()
    informes = ['Todos los legajos','Algunos legajos']
    respuesta = pyip.inputMenu(informes,prompt='Seleccione alguno de los informes disponibles: \n',
                               lettered=True)
    
    
    print('\n')
    if respuesta == 'Todos los legajos':        
        query = selectAll.format(fechaInicio,fechaFin)# query select *
        frameCorregido = pd.read_sql(query,manager.conexion(),parse_dates=columnas)
        frameCorregido['fecha'] = pd.to_datetime(frameCorregido['fecha']).dt.date #transforma 2020-10-01 00:00:00 a 2020-10-01    
        frameCorregido.rename(columns = columnasReemplazo,inplace=True) #renombra las columnas
    
    elif respuesta == 'Algunos legajos':
        legajosTupla =[]
        legajo = None
        while legajo != '':
            legajo = pyip.inputInt(prompt='Ingrese el numero de legajo:\n\t',blank=True)
            legajosTupla.append(str(legajo))
        legajosTupla = tuple(legajosTupla)
        query = selectSome.format(fechaInicio,fechaFin,legajosTupla)
        frameCorregido = pd.read_sql(query,manager.conexion(),parse_dates=columnas)
        frameCorregido['fecha'] = pd.to_datetime(frameCorregido['fecha']).dt.date# same above
        frameCorregido.rename(columns = columnasReemplazo,inplace=True)#same above
       
    frameCorregido.drop(['id'],axis=1,inplace=True)  
    
    legajosNoRotativos,frameQuerido = empleadosFrame()
    
    legajos = frameCorregido['Legajo'].unique()
    frameFinalCorregido = creacionFrameVacio()
    frameFinalCorregido = agregadoColumnas(frameFinalCorregido)
    frameFinalExtras = creacionFrameVacio()
    frameFinalExtras = agregadoColumnas(frameFinalExtras)
    for legajo in legajos:
        try:
            area = frameQuerido.loc[int(legajo),'AREA']
        except:
            msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
            logger.warning(msg)
            continue
        newFrame = frameCorregido[frameCorregido['Legajo']==legajo].copy()


        if int(legajo) in legajosNoRotativos:                           
            newFrame = calculador.horasTrabajadas(newFrame,mediosDias = mediosDias)
            newFrame = agregadoColumnas(newFrame)
            newFrame = calculador.restaRetrasosTardanzas(newFrame,mediosDias = mediosDias)

            frameFinalCorregido = frameFinalCorregido.append(newFrame)

            frameExtras = calculador.horasExtrasTrabajadas(newFrame,feriados=feriados,mediosDias = mediosDias,
                                                           toleranciaHoraria = toleranciaHoraria)

            frameFinalExtras = frameFinalExtras.append(frameExtras)

        else:
            newFrame = calculador.horasTrabajadasRotativos(newFrame,area,mediosDias = mediosDias)
            newFrame = agregadoColumnas(newFrame)
            newFrame = retTarRotativos(newFrame, legajo, medioDias = mediosDias)
            frameFinalCorregido = frameFinalCorregido.append(newFrame)            
            frameExtras = calculador.horasExtrasTrabajadasRotativos(newFrame,area,feriados=feriados,mediosDias = mediosDias,
                                                                    toleranciaHoraria = toleranciaHoraria)
            frameFinalExtras = frameFinalExtras.append(frameExtras)

    if frameFinalCorregido.empty:
        msgEmpty = colored('No hay registros sobre los cuales trabajar\n','grey',on_color='on_red')
        print(msgEmpty)
        return frameFinalCorregido
    else:
        return frameFinalExtras

def cambioPorMotivos(frame,motivos):
    """
    Funcion que se encarga de mover las H.Norm a otra columna en funcion del motivo elegido. Si no hay motivo lo deja en H.Norm

    Parameters
    ----------
    frame : DataFrame
        DataFrame con los datos a modificar.
    motivos : variable
        Lista de motivos por los cuales se puede cambiar un registro.

    Returns
    -------
    frame : DataFrame
        Dataframe con los datos corregidos.

    """
    
    for fila in range(len(frame)):
        causa = frame.iloc[fila,14]
        fecha = frame.iloc[fila,3]
        legajo = frame.iloc[fila,0]
        try:
            indice = motivos[causa]
            frame.iloc[fila,21+indice] = frame.iloc[fila,16]
            frame.iloc[fila,16] = 0
        except:
            msg = 'Correcion de motivo no valida en fecha ({})  legajo: ({})'.format(fecha,legajo)
            logger.warning(msg)
            continue
    
    return frame

def calculosAdicionalesTotalizados(frameFinalExtras,fechaInicio,fechaFin,feriados,empleados,empleadosExtras):
    """
    

    Parameters
    ----------
    frameFinalExtras : TYPE
        DESCRIPTION.
    fechaInicio : TYPE
        DESCRIPTION.
    fechaFin : TYPE
        DESCRIPTION.
    feriados : TYPE
        DESCRIPTION.
    empleados : dict
        DESCRIPTION.
    empleadosExtras : TYPE
        DESCRIPTION.

    Returns
    -------
    None.

    """

    archivo = pyip.inputCustom(prompt='Ingrese el nombre que desea ponerle al EXCEL: \n',
                                customValidationFunc=validador)
    orden = ['Legajo', 'Nombre', 'Dia', 'Fecha', 'Ingreso_0', 'Egreso_0',
       'Ingreso_1', 'Egreso_1', 'Ingreso_2', 'Egreso_2', 'Ingreso_3',
       'Egreso_3', 'Ingreso_4', 'Egreso_4','Motivo', 'Observación',
       'H.Norm','H. 50', 'H. 100','Tardanzas','Retiros']
    frameFinalExtras = frameFinalExtras[orden]
    print('\n')
    nombre = os.path.join(os.getcwd(),pathExcelInforme,archivo)
    nombre = nombre+ '.xlsx'
    frameFinalExtras = frameFinalExtras.sort_values(by=['Legajo','Fecha'])
        
    for key in motivos.keys():
        frameFinalExtras[key] = 0
            
    frameFinalExtras = cambioPorMotivos(frameFinalExtras,motivos=motivos)
        
        
    frameFinalExtras.to_excel(nombre,sheet_name='Registros',index=False)
        
    book = load_workbook(nombre)
    writer = pd.ExcelWriter(nombre, engine = 'openpyxl') #writer para escribir 2 hojas en el excel
    writer.book = book      
        
    frameTotalizado = hojaTotalizadora(frameFinalExtras, fechaInicio, fechaFin,feriados,empleados,empleadosExtras)
    frameTotalizado.to_excel(writer, sheet_name = 'Totalizado',index=False)
    writer.save()
    writer.close()     

    

def informeFaltasTardanzas(frame,fechaInicio,fechaFin,medioDias=[],feriados=[]):
    """
    Funcion que se encarga de crear el informe en pdf de los retrasos, faltas  tardanzas, para eso crea un dict
    e itera sobre el dataframe para ir contando

    Parameters
    ----------
    frame : dataFrame
        dataFrame con las horas calculadas.
    fechaInicio : datetime.date
        fechaInicio sobre la cual se esta analizando.
    fechaFin : datetime.date
        fechaFin sobre la cual se analiza.
    medioDias : TYPE, optional
        DESCRIPTION. The default is [].
    feriados : TYPE, optional
        DESCRIPTION. The default is [].

    Returns
    -------
    None.

    """
    
    legajosNoRotativos,frameQuerido = empleadosFrame()
    frame = frame.sort_values(by=['Legajo','Fecha'])
    
    try:    
        fechaInicio = pd.to_datetime(fechaInicio).date()
        fechaFin = pd.to_datetime(fechaFin).date()
        diasLaborales = list(pd.bdate_range(fechaInicio,fechaFin))
        diasLaborales = [x.date() for x in diasLaborales]
 
        legajosFrame = frame['Legajo'].unique()
        if len(feriados) >= 1 :
            for feria in feriados:
                try:
                    diasLaborales.remove(feria)
                except:
                    continue
            
            
        empleados = {}
        
        for legajo in legajosFrame:
            empleados[str(legajo)]={}
            empleados[str(legajo)]['Tardanzas'] = {}
            empleados[str(legajo)]['Retiros'] = {}
            empleados[str(legajo)]['Nombre'] = ''
            empleados[str(legajo)]['Faltas'] = {}
            empleados[str(legajo)]['Area'] = ''

        for legajo in legajosFrame: 
            try:
                area = frameQuerido.loc[int(legajo),'AREA']
                empleados[str(legajo)]['Area'] = area
            except:
                msg = 'El siguiente legajo ({}) no se encuentra en la BD'.format(legajo)
                logger.warning(msg)
                continue
            
            faltas = []
            newFrame = frame[frame['Legajo']==int(legajo)].copy()
            diasTrabajados = list(newFrame['Fecha'])
            

            for fila in range(len(newFrame)):
                fecha = newFrame.iloc[fila,3]
                if newFrame.iloc[fila,4] == pd.to_datetime(('{} 00:00').format(fecha)) and newFrame.iloc[fila,5] == pd.to_datetime(('{} 00:00').format(fecha)):
                    empleados[str(legajo)]['Faltas'][str(fecha)] = {}
                    empleados[str(legajo)]['Faltas'][str(fecha)]['Motivo'] = newFrame.iloc[fila,14]
                    empleados[str(legajo)]['Faltas'][str(fecha)]['Observacion'] = newFrame.iloc[fila,15]
            # diasFrame = list(newFrame['Ingreso_0'])
            # diasTrabajados = list(set(diasTrabajados + diasFrame))         
       
            # Ingreso_0 = newFrame['Ingreso_0']
            # Egreso_0 = newFrame['Egreso_0']
            # horaIngresoCero = list(Ingreso_0.dt.hour)
            # horaEgresoCero = list(Egreso_0.dt.hour)
    
            # for x in range(len(horaIngresoCero)):
            #     if horaIngresoCero[x] == 0 and horaEgresoCero[x] == 0 and diasTrabajados[x] not in feriados:
            #         faltas.append(diasTrabajados[x])
            # empleados[str(legajo)]['Faltas'] = faltas
            
            toleranciaHoraria = 1
            
            if int(legajo) in legajosNoRotativos:                         
           
                for x in range(len(newFrame)):
                    legajo = newFrame.iloc[x,0]
                    nombre = newFrame.iloc[x,1]
                    dia = newFrame.iloc[x,2]
                    fecha = newFrame.iloc[x,3]
                    horaIngreso = pd.to_datetime(('{} 08:00').format(fecha))
                    horaSalida = pd.to_datetime(('{} 16:48').format(fecha))
                    horaSalidaMedioDia = pd.to_datetime(('{} 12:30').format(fecha))
                    cero = pd.to_datetime(('{} 00:00').format(fecha))
                    tardanza = 0
                    retiro= 0
                    
                    for idx in range(5,13,2):
                        if newFrame.iloc[x,idx] == cero:
                            salida = newFrame.iloc[x,idx -2]
                            break

                    if newFrame.iloc[x,4] > horaIngreso and newFrame.iloc[x,2] != 'Sábado':
                        tardanza = round((((newFrame.iloc[x,4] - horaIngreso).seconds)/60),2)               
                    
                    if fecha in medioDias:
                        if salida < horaSalidaMedioDia:
                            retiro = round((((horaSalidaMedioDia - salida).seconds)/60),2)                        
                    elif fecha not in medioDias:
                        if salida < horaSalida and dia != 'Sábado':
                            retiro = round((((horaSalida - salida).seconds)/60),2)  
    
                    if tardanza != 0:                    
                        empleados[str(legajo)]['Tardanzas'][str(fecha)] = (dia,tardanza)
                    if retiro != 0:
                        empleados[str(legajo)]['Retiros'][str(fecha)] = (dia,retiro)
                    
                    empleados[str(legajo)]['Nombre']= nombre
                
            else:
                if area in rotativosInyeccion:
                    for x in range(len(newFrame)):
                        legajo = newFrame.iloc[x,0]
                        nombre = newFrame.iloc[x,1]
                        dia = newFrame.iloc[x,2]
                        fecha = newFrame.iloc[x,3]
                        ayer = fecha - datetime.timedelta(days=1)
                        mañana = fecha + datetime.timedelta(days=1)
                        
                        primerIngreso = '08:00'
                        segundoIngreso = '16:00'
                        tercerIngreso = '00:00'
                        
                        primerSalida = '16:00'
                        segundaSalida = '00:00'
                        tercerSalida = '08:00'
                        
                        horaSalidaSabado = '13:00'
                        medioDia = '12:30'

                        turnoMañanaPrimerIngreso = (pd.to_datetime(('{} {}').format(fecha,primerIngreso)))
                        turnoTardeIngreso = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso)))
                        turnoNocheIngreso = (pd.to_datetime(('{} {}').format(fecha,tercerIngreso)))
                            
                            
                        turnoMañanaPrimerSalida = (pd.to_datetime(('{} {}').format(fecha,primerSalida)))
                        turnoTardeSalida = (pd.to_datetime(('{} {}').format(mañana,segundaSalida)))
                        turnoNocheSalida = (pd.to_datetime(('{} {}').format(fecha,tercerSalida)))  
                            
                            
                        cero = pd.to_datetime(('{} 00:00').format(fecha))
                        medioDia = pd.to_datetime(('{} 12:30').format(fecha))
                        salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                        
                        tardanza = 0
                        retiro= 0
                        
                        ingresoOperario = newFrame.iloc[x,4]
                        for idx in range(5,13,2):
                            if newFrame.iloc[x,idx] == cero:
                                salida = newFrame.iloc[x,idx -2]
                                break
                           
                        if newFrame.iloc[x,2] != 'Sábado':
                            if  turnoMañanaPrimerIngreso < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoMañanaPrimerIngreso).seconds)/60),2)
                                
                            if  turnoTardeIngreso < ingresoOperario < turnoTardeIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoTardeIngreso).seconds)/60),2)
                                
                            if  turnoNocheIngreso < ingresoOperario < turnoNocheIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoNocheIngreso).seconds)/60),2)
                            
                            
                            
                            
                            
                        if fecha in medioDias:
                            if  turnoMañanaPrimerIngreso < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours=3) :
                                if salida < horaSalidaMedioDia:
                                    retiro = round((((horaSalidaMedioDia - salida).seconds)/60),2)                        
                            
                        if fecha not in medioDias:
                            if  turnoMañanaPrimerSalida - timedelta(hours=3)  < salida < turnoMañanaPrimerSalida:
                                if salida < turnoMañanaPrimerSalida and dia != 'Sábado':
                                    retiro = round((((turnoMañanaPrimerSalida - salida).seconds)/60),2)
                                        
                            if  turnoTardeSalida - timedelta(hours=3)  < salida < turnoTardeSalida:
                                if salida < turnoTardeSalida:
                                    retiro = round(((( turnoTardeSalida - salida).seconds)/60),2)
                                        
                            if  turnoNocheSalida - timedelta(hours=3)  < salida < turnoNocheSalida:
                                if salida < turnoNocheSalida:
                                    retiro = round(((( turnoNocheSalida - salida).seconds)/60),2)
                                        
                                    
            
                        if tardanza != 0:                    
                            empleados[str(legajo)]['Tardanzas'][str(fecha)] = (dia,tardanza)
                        if retiro != 0:
                            empleados[str(legajo)]['Retiros'][str(fecha)] = (dia,retiro)
                            
                        empleados[str(legajo)]['Nombre']= nombre
                else:
                    for x in range(len(newFrame)):
                        legajo = newFrame.iloc[x,0]
                        nombre = newFrame.iloc[x,1]
                        dia = newFrame.iloc[x,2]
                        fecha = newFrame.iloc[x,3]
                        ayer = fecha - datetime.timedelta(days=1)
                        mañana = fecha + datetime.timedelta(days=1)
                        
                        primerIngreso = '07:00'
                        segundoIngreso = '15:00'
                        tercerIngreso = '23:00'
                        
                        primerSalida = '15:00'
                        segundaSalida = '23:00'
                        tercerSalida = '07:00'
                        
                        horaSalidaSabado = '13:00'
                        medioDia = '12:30'
                        
                        turnoMañanaPrimerIngreso = (pd.to_datetime(('{} {}').format(fecha,primerIngreso)))
                        turnoTardeIngreso = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso)))
                        turnoNocheIngreso = (pd.to_datetime(('{} {}').format(fecha,tercerIngreso)))
                            
                            
                        turnoMañanaPrimerSalida = (pd.to_datetime(('{} {}').format(fecha,primerSalida)))
                        turnoTardeSalida = (pd.to_datetime(('{} {}').format(mañana,segundaSalida)))
                        turnoNocheSalida = (pd.to_datetime(('{} {}').format(fecha,tercerSalida)))             
                            
                            
                        cero = pd.to_datetime(('{} 00:00').format(fecha))
                        medioDia = pd.to_datetime(('{} 12:30').format(fecha))
                        salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                        
                        tardanza = 0
                        retiro= 0
                        
                        ingresoOperario = newFrame.iloc[x,4]
                        for idx in range(5,13,2):
                            if newFrame.iloc[x,idx] == cero:
                                salida = newFrame.iloc[x,idx -2]
                                break
                            
                        if newFrame.iloc[x,2] != 'Sábado':
                            if  turnoMañanaPrimerIngreso < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoMañanaPrimerIngreso).seconds)/60),2)
                                
                            if  turnoTardeIngreso < ingresoOperario < turnoTardeIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoTardeIngreso).seconds)/60),2)
                                
                            if  turnoNocheIngreso < ingresoOperario < turnoNocheIngreso + timedelta(hours=3) :
                                tardanza = round((((ingresoOperario - turnoNocheIngreso).seconds)/60),2)
                            
                            
                            
                            
                            
                        if fecha in medioDias:
                            if  turnoMañanaPrimerIngreso < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours=3) :
                                if salida < horaSalidaMedioDia:
                                    retiro = round((((horaSalidaMedioDia - salida).seconds)/60),2)                        
                            
                        if fecha not in medioDias:
                            if  turnoMañanaPrimerSalida - timedelta(hours=3)  < salida < turnoMañanaPrimerSalida:
                                if salida < turnoMañanaPrimerSalida and dia != 'Sábado':
                                    retiro = round((((turnoMañanaPrimerSalida - salida).seconds)/60),2)
                                        
                            if  turnoTardeSalida - timedelta(hours=3)  < salida < turnoTardeSalida:
                                if salida < turnoTardeSalida:
                                    retiro = round(((( turnoTardeSalida - salida).seconds)/60),2)
                                        
                            if  turnoNocheSalida - timedelta(hours=3)  < salida < turnoNocheSalida:
                                if salida < turnoNocheSalida:
                                    retiro = round(((( turnoNocheSalida - salida).seconds)/60),2)
                                        
                                    
            
                        if tardanza != 0:                    
                            empleados[str(legajo)]['Tardanzas'][str(fecha)] = (dia,tardanza)
                        if retiro != 0:
                            empleados[str(legajo)]['Retiros'][str(fecha)] = (dia,retiro)
                            
                        empleados[str(legajo)]['Nombre']= nombre
        
        return empleados       

    except:
        print('\nExistio un problema en la creacion del word/pdf.\nProbablemente el archivo en word/pdf con ese nombre esta abierto.')
        print('\nCierrelo y vuelva a ingresar los datos.')
        logger.error("excepcion desconocida: %s", traceback.format_exc())

def escritorInformeFaltasTardanzas(empleados,fechaInicio,fechaFin):
    """
    Funcionq ue se encarga de escribir el informe de faltas,tardanzas y retrasos en funcion de los datos pasados en el diccionario.

    Parameters
    ----------
    empleados : dict
        DIccionario con todos los valores sobre los que se va a iterar en funcion del legajo.
    fechaInicio : datetime.date
        DESCRIPTION.
    fechaFin : datetime.date
        DESCRIPTION.

    Returns
    -------
    None.

    """
    
    fechaInicio = pd.to_datetime(fechaInicio).date()
    fechaFin = pd.to_datetime(fechaFin).date()

    faltasWord = '\tFalta registrada el dia {}. Motivo: {}. Observación: {}. \n'
    tardanzasWord = '\tTardanza de {} minutos registrada el dia {} ({}).\n'
    retirosWord = '\tRetiro anticipado de {} minutos registrado el dia {} ({}).\n'
    
    try:
        doc = docx.Document()
        doc.add_heading(('Faltas, tardanzas y retiros entre {} y {}').format(fechaInicio,fechaFin), 0)
        for key in empleados.keys():
            
            elBoleano = False
            for llave in empleados[key].keys():
                if llave != 'Nombre':
                    elBoleano |= bool(empleados[key][llave])
            if elBoleano == False:
                continue
                    
            
            doc.add_heading(('Informe sobre {} ({}) ({}):').format(empleados[key]['Nombre'],key,empleados[key]['Area']),level=1)
            
            
            doc.add_heading(('Faltas:'),level=2)
            primerParrafo = doc.add_paragraph()
            # for falta in empleados[key]['Faltas']:#es una lista el resultado 
            #     primerParrafo.add_run((faltasWord).format(falta))
            
            for falta in empleados[key]['Faltas'].keys():#es una lista el resultado
                motivo = empleados[key]['Faltas'][falta]['Motivo']
                observacion = empleados[key]['Faltas'][falta]['Observacion']
                if observacion == 'nan':
                    observacion = 'Sin información'
                if motivo == 'nan':
                    motivo = 'Sin especificar'
                primerParrafo.add_run((faltasWord).format(falta,motivo,observacion))
                
            
            doc.add_heading(('Tardanzas:'),level=2)
            segundoParrafo = doc.add_paragraph()
            for llave,valor in empleados[key]['Tardanzas'].items():
                segundoParrafo.add_run((tardanzasWord).format(valor[1],llave,valor[0]))
            
            
            doc.add_heading(('Retiros anticipados:'),level=2)
            tercerParrafoInterno = doc.add_paragraph()
            for llave,valor in empleados[key]['Retiros'].items():
                tercerParrafoInterno.add_run((retirosWord).format(valor[1],llave,valor[0]))
        
        archivo = pyip.inputCustom(prompt='Ingrese el nombre que desea ponerle al informe PDF: \n',
                            customValidationFunc=validador)

        nombre = os.path.join(os.getcwd(),pathInformesFaltasTardanzas,archivo)
        word = nombre+ '.docx'
        pdf = nombre + '.pdf'
        
        pathToWord = word
        pathToPDF = pdf
        doc.save(pathToWord)
        logger.info('Creacion del Word de manera correcta')
        sleep(0.5)
        wdFormatPDF = 17 # Word's numeric code for PDFs.
        wordObj = win32com.client.Dispatch('Word.Application')
        docObj = wordObj.Documents.Open(pathToWord)
        docObj.SaveAs(pathToPDF, FileFormat=wdFormatPDF)
        docObj.Close()
        wordObj.Quit()
        logger.info('Creacion del PDF de manera correcta')
        os.remove(pathToWord)
    except:
        print('\nExistio un problema en la creacion del word/pdf.\nProbablemente el archivo en word/pdf con ese nombre esta abierto.')
        print('\nCierrelo y vuelva a ingresar los datos.')
        logger.error("excepcion desconocida: %s", traceback.format_exc())
    


def datosOperario(areas,formaDePago):
    
    legajo = pyip.inputInt(prompt = 'Ingrese el LEGAJO del empleado:\n',min=0)
    print('\n')
    nombre = pyip.inputStr(prompt = 'Ingrese el NOMBRE del empleado:\n').upper()
    print('\n')
    apellido = pyip.inputStr(prompt = 'Ingrese el APELLIDO del empleado:\n').upper()
    print('\n')
    area = pyip.inputMenu(areas,prompt='Ingrese una de las AREAS posibles:\n',lettered=True).upper()
    print('\n')
    pago = pyip.inputMenu(formaDePago,prompt='Ingrese el tipo de pago:\n',lettered=True)
    print('\n')
    
    return legajo,nombre,apellido,area,pago

def repreguntar():
    decision = pyip.inputYesNo(prompt='¿Los datos ingresados son correctos? (SI/NO)  ',yesVal='SI',noVal='NO')
    print('\n')    
    if decision == 'SI':
        return True
    else:
        return False

def actualizarValor(campo):
    
    if campo == 'AREA':
        valor = pyip.inputMenu(areas,prompt='Ingrese una de las AREAS posibles:\n',lettered=True)
    elif campo == 'TIPO_DE_PAGO':
        valor = pyip.inputMenu(formaDePago,prompt='Ingrese el tipo de pago:\n',lettered=True)
    elif campo == 'LEG':
        valor = pyip.inputInt(prompt = 'Ingrese el LEGAJO del empleado:\n',min=0)
    elif campo == 'NOMBRE':
        valor = pyip.inputStr(prompt = 'Ingrese el NOMBRE del empleado:\n').upper()
    elif campo == 'APELLIDO':
        valor = pyip.inputStr(prompt = 'Ingrese el APELLIDO del empleado:\n').upper()
        
    return valor
        
    
    
# class Motor:
#     def mainLoop(self):
#         print('-'*100)
#         print(' '*50,'Bienvenido',' '*50)
#         print('-'*100)
    
    
#         tareas = ['Ordenado de registros','Creación de informes','Gestion de Base de datos','Salir']
#         tareasOrdenado = ['Limpieza de registros','Actualización de registros','Volver','Salir']
#         tareasInformes = ['Ingreso de fechas','Volver','Salir']
#         tareasBD = ['Agregar empleado','Actualizar empleado','Eliminar empleado','Modificar Registro','Descargar','Volver','Salir']
        
#         continuar = True
        
        
#         while continuar:
#             respuesta = pyip.inputMenu(tareas,prompt='¿Que desea hacer?\n',lettered=True)
#             print('\n')
#             continuarOrdenado = True
#             continuarCreacion = True
#             continuarGestionBD = True
            
#             if respuesta == 'Ordenado de registros':
#                 while continuarOrdenado:
#                     ordenadoRespuesta = pyip.inputMenu(tareasOrdenado,prompt='\n¿Que desea hacer?\n',lettered=True)
#                     print('\n')                
#                     if ordenadoRespuesta == 'Limpieza de registros':
                        
#                         decision = False
#                         while not decision:
#                             fechaInicio,fechaFin = fechasDeCalculo(completo=False)
#                             decision = repreguntar()
#                         frame = frameFichadas()
                        
#                         if frame.empty:
#                             noArchivos = colored('No existen archivos que limpiar\n','grey',on_color='on_red')
#                             print(noArchivos)
#                         else:
#                             legajos = frameAnalisisIndividual(frame,fechaInicio,fechaFin)
#                             if legajos.empty:
#                                 msgRechazos = colored('No existen archivos que limpiar, archivo guardado en Rechazos.\n','grey',on_color='on_red')
#                                 print(msgRechazos)
#                                 len_noMarca = limpiezaDeRegistros(legajos, fechaInicio, fechaFin) 
#                             else:
#                                 msgPersistenErrores = colored('\nRegistros erroneos y duplicados eliminados, excel a completar creado.\n','grey',on_color='on_red')
#                                 len_noMarca = limpiezaDeRegistros(legajos, fechaInicio, fechaFin)
#                                 if len_noMarca != 0:
#                                     print(msgPersistenErrores)
                            
#                     elif ordenadoRespuesta == 'Actualización de registros':
                        
#                         decision = False
#                         while not decision:
#                             fechaInicio,fechaFin= fechasDeCalculo(completo=False)
#                             decision = repreguntar()
#                         actualizacionRegistros(fechaInicio,fechaFin)
                        
#                     elif ordenadoRespuesta == 'Volver':
#                         print('Volviendo al PRIMER MENU')
#                         continuarOrdenado = False 
                        
#                     elif ordenadoRespuesta == 'Salir':
#                         continuarOrdenado = False
#                         continuar = False
            
            
#             elif respuesta == 'Creación de informes':
#                 while continuarCreacion:
#                     informesRespuesta = pyip.inputMenu(tareasInformes,prompt='\n¿Que desea hacer?\n',lettered=True)
#                     print('\n')
#                     if informesRespuesta == 'Ingreso de fechas':
#                         decision = False
#                         while not decision:
#                             fechaInicio,fechaFin,feriados,mediosDias = fechasDeCalculo()
#                             decision = repreguntar()
#                         frameCorregido = seleccionInformes(fechaInicio, fechaFin,feriados = feriados,mediosDias= mediosDias)
#                         if frameCorregido.empty:
#                             pass #Internamente ya hay un msj
#                         else:
#                             empleados = informeFaltasTardanzas(frameCorregido,fechaInicio,fechaFin,
#                                     feriados=feriados,medioDias = mediosDias)
                            
#                             empleadosExtras = calculosExtrasRotativos(frameCorregido)
                            
#                             calculosAdicionalesTotalizados(frameCorregido, fechaInicio, fechaFin, feriados, empleados,empleadosExtras)
                            
#                             escritorInformeFaltasTardanzas(empleados, fechaInicio, fechaFin)
                            
#                         continuarCreacion = True                
                    
#                     elif informesRespuesta == 'Volver':
#                         print('Volviendo al PRIMER MENU')
#                         continuarCreacion = False
                        
#                     elif informesRespuesta == 'Salir':
#                         continuarCreacion = False
#                         continuar = False
    
            
            
#             elif respuesta == 'Gestion de Base de datos':
#                 while continuarGestionBD:
                    
#                     baseDeDatosRespuesta = pyip.inputMenu(tareasBD,prompt='\n¿Que desea hacer?\n',lettered=True)
#                     print('\n')
#                     ['Agregar empleado','Actualizar empleado','Eliminar empleado','Modificar registro','Descargar','Volver','Salir']
#                     if baseDeDatosRespuesta == 'Agregar empleado':
                        
#                         decision = False
#                         while not decision:
#                             legajo,nombre,apellido,area,pago = datosOperario(areas, formaDePago)
#                             decision = repreguntar()

#                         managerSQL = ManagerSQL()
#                         insercionBDLegajos(managerSQL, legajo, nombre, apellido, area, pago, insertEmpleado)
#                         continuarGestionBD = True                
                    
#                     elif baseDeDatosRespuesta == 'Actualizar empleado':
                        
#                         campos = ['LEG','APELLIDO','NOMBRE','AREA','TIPO_DE_PAGO']
#                         decision = False
#                         while not decision:
#                             legajo = pyip.inputInt(prompt='Ingrese el LEGAJO del empleado a actualizar:\n',min=0)
#                             campo = pyip.inputMenu(campos,prompt='Elija que campo va a actualizar:\n',lettered=True)
#                             valor = actualizarValor(campo)
#                             decision = repreguntar()
#                         managerSQL = ManagerSQL()
#                         actualizaBDLegajos(managerSQL, legajo, campo, valor, actualizarEmpleado)
#                         print('\n')
#                         continuarGestionBD = True
                    
#                     elif baseDeDatosRespuesta == 'Eliminar empleado':
                        
#                         decision = False
#                         while not decision:
#                             legajo = pyip.inputInt(prompt='Ingrese el LEGAJO del empleado a eliminar:\n',min=0)
#                             decision = repreguntar()
#                         print('\n')
#                         managerSQL = ManagerSQL()
#                         deleteBDLegajos(managerSQL, legajo, deleteEmpleado)     
#                         continuarGestionBD = True
                        
#                     elif baseDeDatosRespuesta == 'Modificar Registro':
                        
#                         decision = False
#                         while not decision:
#                             legajo = pyip.inputInt(prompt='Ingrese el LEGAJO:\n',min=0)
#                             fecha = pyip.inputDate('Ingrese la fecha a actualizar DD/MM/AAAA: ',formats=["%d/%m/%Y"])
#                             decision = repreguntar()
#                         print('\n')
#                         managerSQL = ManagerSQL()
#                         registro = edicionRegistros(managerSQL,selectDeleteRegistro, fecha, legajo)
#                         if registro.empty:
#                             print('\n No existen registros para esa fecha y legajo\n')
#                         else:
#                             cantidad = coloreadorRegistroModificar(registro, fecha)
#                             campo,fechaHora = actualizaLinea(registro, cantidad)
#                             updateRegistroQuery(managerSQL, campo, fechaHora, legajo, fecha)
#                             registro = edicionRegistros(managerSQL,selectDeleteRegistro, fecha, legajo)
#                             print('Nuevo registro: \n')
#                             cantidad = coloreadorRegistroModificar(registro, fecha)                           


#                         continuarGestionBD = True
                    
                        
#                     elif baseDeDatosRespuesta == 'Descargar':
                        
#                         manager = ManagerSQL()
#                         sql_conection = manager.conexion()
#                         consultaEmpleados = pd.read_sql(queryConsultaEmpleados,sql_conection)
#                         archivo = pyip.inputCustom(prompt='Ingrese el nombre que desea ponerle al EXCEL: \n',
#                                 customValidationFunc=validador)
#                         print('\n')        
#                         nombre = os.path.join(os.getcwd(),pathExcelInforme,archivo)
#                         nombre = nombre+ '.xlsx'
#                         consultaEmpleados = consultaEmpleados.sort_values(by=['LEG'])
#                         consultaEmpleados.to_excel(nombre,sheet_name='Registros',index=False)                        
#                         continuarGestionBD = True
                    
#                     elif baseDeDatosRespuesta == 'Volver':
#                         print('Volviendo al PRIMER MENU')
#                         continuarGestionBD = False
                        
#                     elif baseDeDatosRespuesta == 'Salir':
#                         continuarGestionBD = False
#                         continuar = False
    
            
#             elif respuesta == 'Salir':
#                 continuar= False


        
    
    
    
#if __name__ == '__main__':
    # frameParaVer = None
    # try:
    #     motor = Motor()
    #     motor.mainLoop()
    #     sys.exit()
    # except Exception:
    #     logger.error("excepcion desconocida: %s", traceback.format_exc())
    #     sleep(5)
if __name__ == '__main__':
    pass