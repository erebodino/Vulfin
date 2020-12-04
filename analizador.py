import docx
import pandas as pd
import os
import win32com.client
import datetime
import pyinputplus as pyip
from termcolor import colored
from paths import(nombreInformeNoFichadasWord,nombreInformeNoFichadasPDF,pathInformesNoFichadas,
pathInformesFaltasTardanzas,
nombreInformeFaltasTardanzasWord,
nombreInformeFaltasTardanzasPDF,
rotativosInyeccion,
rotativosSoplado)
from datetime import timedelta

import logging.config
import traceback

logging.config.fileConfig('logger.ini', disable_existing_loggers=False)
logger = logging.getLogger(__name__)

class Analizador:
    """
    Clase que viene a hacer de limpieza para el dataframe, se encarga de leer linea por linea
    e ir acomodando el dataFrame.
    
    El criterio actual de turnos es el siguiente:
        -1er Turno del dia de 00:00 a 08:00
        -2do turno del dia de 08:00 a 16:00 o 16:48 dependiendo del operario (rotativo o no)
        -3er turno del dia de 16:00 a 00:00


    """    
    
    def __init__(self,frameEnAnalisis,fechaInicio,fechaFin):
        self.frameEnAnalisis = frameEnAnalisis
        self.fechaInicio = fechaInicio
        self.fechaFin = fechaFin
        
    def sanityCheck(self):
        
        diaAnterior = self.fechaInicio - timedelta(days=1)
        diaPosterior = self.fechaFin + timedelta(days=1)
        self.frameEnAnalisis['Fecha'] = pd.to_datetime(self.frameEnAnalisis['Fecha']).dt.date
        
        estado = diaAnterior in self.frameEnAnalisis['Fecha'].values and diaPosterior in self.frameEnAnalisis['Fecha'].values
        estado &=  self.fechaInicio in self.frameEnAnalisis['Fecha'].values and self.fechaFin in self.frameEnAnalisis['Fecha'].values
        return estado
    
    def creacionDiasCeros(self,legajo,nombre):
        pass
    
    def castMascara(self):
        
        semaine = {'Monday': 'Lunes',
               'Tuesday' : 'Martes',
               'Wednesday': 'Miércoles',
               'Thursday':'Jueves',
               'Friday':'Viernes',
               'Saturday':'Sabado',
               'Sunday':'Domingo',
              }
        
        legajo = self.frameEnAnalisis.iloc[0,0]
        nombre = self.frameEnAnalisis.iloc[0,1]
        diasOperario = list(self.frameEnAnalisis['Fecha'])

        diasLaborales = list(pd.bdate_range(self.fechaInicio,(self.fechaFin + timedelta(days=1))))

        for fecha in diasLaborales:
            lista_final = []
            if fecha not in diasOperario:
                for u in range(10):
                    lista_final.append(pd.to_datetime(('{} 00:00').format(fecha)))

                self.frameEnAnalisis = self.frameEnAnalisis.append({'Empleado':legajo,'Nombre':nombre,
                                                                    'Dia':semaine[fecha.day_name()],'Fecha':fecha,
                                                                    'Ingreso_0':lista_final[0],'Egreso_0':lista_final[1],
                                                                    'Ingreso_1':lista_final[2],'Egreso_1':lista_final[3],
                                                                    'Ingreso_2':lista_final[4],'Egreso_2':lista_final[5],
                                                                    'Ingreso_3':lista_final[6],'Egreso_3':lista_final[7],
                                                                    'Ingreso_4':lista_final[8],'Egreso_4':lista_final[9]},
                                                                   ignore_index=True)
        return self.frameEnAnalisis
        
    def limpiador(self,area):
        """
        Funcion que itera por cada uno de las lineas y limpia.
        Se necesita que los registros comienzen un dia antes del dia de analisis.

        Returns
        -------
        TYPE
            DESCRIPTION.

        """
        legajo = self.frameEnAnalisis.iloc[0,0]

        
        import pandas as pd
        import datetime
        self.frameEnAnalisis = self.frameEnAnalisis.reset_index(drop=True)
        fechas = self.frameEnAnalisis['Fecha']
        inicioIndex = fechas[fechas == self.fechaInicio]
        finIndex = fechas[fechas == self.fechaFin]
        inicioIndex = fechas[fechas == self.fechaInicio].index[0]
        finIndex = fechas[fechas == self.fechaFin].index[0]

       
        for renglon in range(inicioIndex,finIndex+ 1):
            if area in rotativosInyeccion:
                dia = self.frameEnAnalisis.iloc[renglon,2]
                fecha = self.frameEnAnalisis.iloc[renglon,3]
                ayer = fecha - datetime.timedelta(days=1)
                mañana = fecha + datetime.timedelta(days=1)
                
                turnoMañanaIngreso = pd.to_datetime(('{} 8:00').format(fecha))
                turnoMañanaSalida = pd.to_datetime(('{} 16:48').format(fecha))
                turnoMañanaSalidaSinComer = pd.to_datetime(('{} 16:00').format(fecha))
                turnoTardeIngreso = pd.to_datetime(('{} 16:00').format(fecha))
                turnoTardeEgreso = pd.to_datetime(('{} 00:00').format(mañana))
                turnoNocheIngreso = pd.to_datetime(('{} 00:00').format(mañana))
                turnoNocheSalida = pd.to_datetime(('{} 8:00').format(mañana))
                #turnoMañanaIngreso = pd.to_datetime(('{} 8:00').format(dia))
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                medioDia = pd.to_datetime(('{} 12:00').format(fecha))
                
                #--------------------Turnos de ayer----------------------
                turnoMañanaIngresoAyer = pd.to_datetime(('{} 8:00').format(ayer))
                turnoMañanaSalidaAyer = pd.to_datetime(('{} 16:48').format(ayer))
                turnoMañanaSalidaSinComerAyer = pd.to_datetime(('{} 16:00').format(ayer))
                turnoTardeIngresoAyer = pd.to_datetime(('{} 16:00').format(ayer))
                turnoNocheIngresoAyer = pd.to_datetime(('{} 00:00').format(ayer))
                turnoNocheSalidaAyer = pd.to_datetime(('{} 8:00').format(ayer))
                #turnoMañanaIngreso = pd.to_datetime(('{} 8:00').format(dia))
                ceroAyer = pd.to_datetime(('{} 00:00').format(ayer))
                medioDiaAyer = pd.to_datetime(('{} 12:00').format(ayer))
                
                turnoTardeIngresoTomorrow = pd.to_datetime(('{} 16:00').format(mañana))
                turnoMañanaIngresoTomorrow = pd.to_datetime(('{} 8:00').format(mañana))
                turnoNocheIngresoTomorrow = pd.to_datetime(('{} 00:00').format(mañana))
                ceroMañana = pd.to_datetime(('{} 00:00').format(mañana))
                

                
                for posicion in range(4,12,2):# iteracion sobre las columnas del dataFrame, arrancando con el primer ingreso.
                    #print(self.frameEnAnalisis.iloc[renglon,0],dia,' ',self.frameOriginal.iloc[renglon+ 2,posicion],self.frameEnAnalisis.iloc[renglon,posicion]) 
                    
                    if self.frameEnAnalisis.iloc[renglon,posicion+ 1] == cero and self.frameEnAnalisis.iloc[renglon,posicion] == cero:
                        break
                    
                    if self.frameEnAnalisis.iloc[renglon,posicion+ 1] == cero and dia != 'Sábado':

                            if posicion <6:
                                if self.frameEnAnalisis.iloc[renglon +1,3] == fecha +timedelta(days=1):# Verifica que el proximo dia en donde se va a sacar el datos sea un dia siguiente.-
                                    
                                    if (turnoNocheIngreso - timedelta(hours=5)) <= self.frameEnAnalisis.iloc[renglon,posicion] <= (turnoNocheIngreso + timedelta(hours=5)) :
                                        if self.frameEnAnalisis.iloc[renglon +1,posicion] >= turnoMañanaIngresoTomorrow:
                                            self.frameEnAnalisis.iloc[renglon+ 1,posicion+ 2] = self.frameEnAnalisis.iloc[renglon+ 1,posicion+ 1]
                                            self.frameEnAnalisis.iloc[renglon+ 1,posicion+ 1] = self.frameEnAnalisis.iloc[renglon+ 1,posicion]
                                            self.frameEnAnalisis.iloc[renglon+ 1,posicion] = self.frameEnAnalisis.iloc[renglon,posicion]
                                            self.frameEnAnalisis.iloc[renglon,posicion] = cero
                                            
                                                
                                            break
                                    
                                    elif (turnoTardeIngreso - timedelta(hours=2)) <= self.frameEnAnalisis.iloc[renglon,posicion] <= (turnoTardeIngreso + timedelta(hours=2)):
                                       
                                        if  (turnoTardeEgreso - timedelta(hours=2)) <= self.frameEnAnalisis.iloc[renglon+ 1,posicion] <= (turnoTardeEgreso + timedelta(hours=2)):
                                            self.frameEnAnalisis.iloc[renglon,posicion +1] = self.frameEnAnalisis.iloc[renglon+ 1,posicion]
                                            self.frameEnAnalisis.iloc[renglon+ 1,posicion] = self.frameEnAnalisis.iloc[renglon+ 1,posicion + 1]
                                            self.frameEnAnalisis.iloc[renglon+ 1,posicion +1] = self.frameEnAnalisis.iloc[renglon+ 1,posicion + 2]
                                            self.frameEnAnalisis.iloc[renglon+ 1,posicion + 2] = ceroMañana
                                            
                                        else:
                                            break
                                        
                                    else:
                                        break
                            else:
                                break                               
        
                    else: 
                        if posicion <6:
                            if self.frameEnAnalisis.iloc[renglon- 1,posicion +2] == ceroAyer:
                                if self.frameEnAnalisis.iloc[renglon,posicion] > turnoMañanaIngreso and \
                                    (self.frameEnAnalisis.iloc[renglon,posicion +1] > (turnoTardeIngreso + timedelta(hours=3)) or \
                                     self.frameEnAnalisis.iloc[renglon,posicion +1] == cero):  
                                        
                                        self.frameEnAnalisis.iloc[renglon,posicion+ 2] = self.frameEnAnalisis.iloc[renglon,posicion+ 1]
                                        self.frameEnAnalisis.iloc[renglon,posicion+ 1] = self.frameEnAnalisis.iloc[renglon,posicion]
                                        self.frameEnAnalisis.iloc[renglon,posicion] = self.frameEnAnalisis.iloc[renglon- 1,posicion +1]
                                        self.frameEnAnalisis.iloc[renglon- 1,posicion +2]= cero
                                        break
                            
                            else:
                                if self.frameEnAnalisis.iloc[renglon,posicion] > turnoMañanaIngreso and \
                                    (self.frameEnAnalisis.iloc[renglon,posicion +1] > (turnoTardeIngreso + timedelta(hours=3)) or \
                                     self.frameEnAnalisis.iloc[renglon,posicion +1] == cero):  
                                        
                                        self.frameEnAnalisis.iloc[renglon,posicion+ 2] =  self.frameEnAnalisis.iloc[renglon,posicion+ 1]
                                        self.frameEnAnalisis.iloc[renglon,posicion +1] = self.frameEnAnalisis.iloc[renglon,posicion]
                                        self.frameEnAnalisis.iloc[renglon,posicion] = self.frameEnAnalisis.iloc[renglon- 1,posicion+ 2]
                                        self.frameEnAnalisis.iloc[renglon- 1,posicion+ 2] = ceroAyer
                                        
                                        if str(self.fechaFin) == str(fecha):
                                            self.frameEnAnalisis.iloc[renglon,posicion+ 2] = cero
                                        
                                        break
                                

                                        
                  
                            
                    
                   
            elif area in rotativosSoplado:               
                    legajo =  self.frameEnAnalisis.iloc[renglon,0]
                    dia = self.frameEnAnalisis.iloc[renglon,2]
                    fecha = self.frameEnAnalisis.iloc[renglon,3]
                    ayer = fecha - datetime.timedelta(days=1)
                    mañana = fecha + datetime.timedelta(days=1)
                    margen = datetime.timedelta(minutes=5)
                    
                    turnoMañanaIngreso = pd.to_datetime(('{} 07:00').format(fecha))
                    turnoMañanaSalida = pd.to_datetime(('{} 15:00').format(fecha))
                    turnoMañanaSalidaSinComer = pd.to_datetime(('{} 15:00').format(fecha))
                    turnoTardeIngreso = pd.to_datetime(('{} 15:00').format(fecha))
                    turnoTardeEgreso = pd.to_datetime(('{} 23:00').format(fecha))
                    turnoNocheIngreso = pd.to_datetime(('{} 23:00').format(fecha))
                    turnoNocheSalida = pd.to_datetime(('{} 7:00').format(mañana))
                    #turnoMañanaIngreso = pd.to_datetime(('{} 8:00').format(dia))
                    cero = pd.to_datetime(('{} 00:00').format(fecha))
                    medioDia = pd.to_datetime(('{} 12:00').format(fecha))
                    
                    #--------------------Turnos de ayer----------------------
                    turnoMañanaIngresoAyer = pd.to_datetime(('{} 7:00').format(ayer))
                    turnoMañanaSalidaAyer = pd.to_datetime(('{} 15:48').format(ayer))
                    turnoMañanaSalidaSinComerAyer = pd.to_datetime(('{} 15:00').format(ayer))
                    turnoTardeIngresoAyer = pd.to_datetime(('{} 15:00').format(ayer))
                    turnoNocheIngresoAyer = pd.to_datetime(('{} 23:00').format(ayer))
                    turnoNocheSalidaAyer = pd.to_datetime(('{} 7:00').format(ayer))
                    #turnoMañanaIngreso = pd.to_datetime(('{} 8:00').format(dia))
                    ceroAyer = pd.to_datetime(('{} 00:00').format(ayer))
                    medioDiaAyer = pd.to_datetime(('{} 12:00').format(ayer))
                    
                    #--------------------Turnos de mañana----------------------
                    turnoMañanaIngresoTomorrow = pd.to_datetime(('{} 7:00').format(mañana))
                    turnoMañanaSalidaTomorrow = pd.to_datetime(('{} 15:48').format(mañana))
                    turnoMañanaSalidaSinComerTomorrow = pd.to_datetime(('{} 15:00').format(mañana))
                    turnoTardeIngresoTomorrow = pd.to_datetime(('{} 15:00').format(mañana))
                    turnoNocheIngresoTomorrow = pd.to_datetime(('{} 23:00').format(mañana))
                    turnoNocheSalidaTomorrow = pd.to_datetime(('{} 7:00').format(mañana))
                    #turnoMañanaIngreso = pd.to_datetime(('{} 8:00').format(dia))
                    ceroMañana = pd.to_datetime(('{} 00:00').format(mañana))
                    medioDiaTomorrow = pd.to_datetime(('{} 12:00').format(mañana))
                    
                    
                    for posicion in range(4,12,2):# iteracion sobre las columnas del dataFrame, arrancando con el primer ingreso.
                    #print(self.frameEnAnalisis.iloc[renglon,0],dia,' ',self.frameOriginal.iloc[renglon+ 2,posicion],self.frameEnAnalisis.iloc[renglon,posicion]) 
                    
                        if self.frameEnAnalisis.iloc[renglon,posicion+ 1] == cero and self.frameEnAnalisis.iloc[renglon,posicion] == cero:
                            break
                        
                        if self.frameEnAnalisis.iloc[renglon,posicion+ 1] == cero and dia != 'Sábado':
    
                                if posicion <6:
                                    if self.frameEnAnalisis.iloc[renglon +1,3] == fecha +timedelta(days=1):# Verifica que el proximo dia en donde se va a sacar el datos sea un dia siguiente.-
                                        
                                        if (turnoNocheIngreso - timedelta(hours=5)) <= self.frameEnAnalisis.iloc[renglon,posicion] <= (turnoNocheIngreso + timedelta(hours=5)) :
                                            if self.frameEnAnalisis.iloc[renglon +1,posicion] >= turnoMañanaIngresoTomorrow:
                                                self.frameEnAnalisis.iloc[renglon+ 1,posicion+ 2] = self.frameEnAnalisis.iloc[renglon+ 1,posicion+ 1]
                                                self.frameEnAnalisis.iloc[renglon+ 1,posicion+ 1] = self.frameEnAnalisis.iloc[renglon+ 1,posicion]
                                                self.frameEnAnalisis.iloc[renglon+ 1,posicion] = self.frameEnAnalisis.iloc[renglon,posicion]
                                                self.frameEnAnalisis.iloc[renglon,posicion] = cero
                                                
                                                    
                                                break
                                        
                                        elif (turnoTardeIngreso - timedelta(hours=2)) <= self.frameEnAnalisis.iloc[renglon,posicion] <= (turnoTardeIngreso + timedelta(hours=2)):
                                           
                                            if  (turnoTardeEgreso - timedelta(hours=2)) <= self.frameEnAnalisis.iloc[renglon+ 1,posicion] <= (turnoTardeEgreso + timedelta(hours=2)):
                                                self.frameEnAnalisis.iloc[renglon,posicion +1] = self.frameEnAnalisis.iloc[renglon+ 1,posicion]
                                                self.frameEnAnalisis.iloc[renglon+ 1,posicion] = self.frameEnAnalisis.iloc[renglon+ 1,posicion + 1]
                                                self.frameEnAnalisis.iloc[renglon+ 1,posicion +1] = self.frameEnAnalisis.iloc[renglon+ 1,posicion + 2]
                                                self.frameEnAnalisis.iloc[renglon+ 1,posicion + 2] = ceroMañana
                                                
                                            else:
                                                break
                                            
                                        else:
                                            break
                                else:
                                    break                               
            
                        else: 
                            if posicion <6:
                                if self.frameEnAnalisis.iloc[renglon- 1,posicion +2] == ceroAyer:
                                    if  (turnoMañanaIngreso - timedelta(hours=1) < self.frameEnAnalisis.iloc[renglon,posicion] < turnoTardeIngreso - timedelta(hours=3)) and \
                                        (self.frameEnAnalisis.iloc[renglon,posicion +1] > (turnoTardeIngreso + timedelta(hours=3)) or \
                                         self.frameEnAnalisis.iloc[renglon,posicion +1] == cero):  
                                            
                                            self.frameEnAnalisis.iloc[renglon,posicion+ 2] = self.frameEnAnalisis.iloc[renglon,posicion+ 1]
                                            self.frameEnAnalisis.iloc[renglon,posicion+ 1] = self.frameEnAnalisis.iloc[renglon,posicion]
                                            self.frameEnAnalisis.iloc[renglon,posicion] = self.frameEnAnalisis.iloc[renglon- 1,posicion +1]
                                            self.frameEnAnalisis.iloc[renglon- 1,posicion +2]= cero
                                            break
                                
                                else:
                                    if (turnoMañanaIngreso - timedelta(hours=1) < self.frameEnAnalisis.iloc[renglon,posicion] < turnoTardeIngreso - timedelta(hours=3)) and \
                                        (self.frameEnAnalisis.iloc[renglon,posicion +1] > (turnoTardeIngreso + timedelta(hours=3)) or \
                                         self.frameEnAnalisis.iloc[renglon,posicion +1] == cero):  
                                            
                                            self.frameEnAnalisis.iloc[renglon,posicion+ 2] =  self.frameEnAnalisis.iloc[renglon,posicion+ 1]
                                            self.frameEnAnalisis.iloc[renglon,posicion +1] = self.frameEnAnalisis.iloc[renglon,posicion]
                                            self.frameEnAnalisis.iloc[renglon,posicion] = self.frameEnAnalisis.iloc[renglon- 1,posicion+ 2]
                                            self.frameEnAnalisis.iloc[renglon- 1,posicion+ 2] = ceroAyer
                                            
                                            if str(self.fechaFin) == str(fecha):
                                                self.frameEnAnalisis.iloc[renglon,posicion+ 2] = cero
                                            
                                            break
        
        self.frameEnAnalisis = self.borradoCeros()
        return self.frameEnAnalisis
    
    def borradoCeros(self):
        
        columnas = [ 'Empleado','Nombre','Dia','Fecha',
                    'Ingreso_0','Egreso_0','Ingreso_1','Egreso_1',
                    'Ingreso_2','Egreso_2','Ingreso_3','Egreso_3',
                    'Ingreso_4','Egreso_4']
        frame = self.frameEnAnalisis.copy()
        frame['Fecha'] = pd.to_datetime(frame['Fecha']).dt.date
        frame.set_index('Fecha',inplace=True)
        toErase = []
        for renglon in range(len(frame)):
            for posicion in range(3,5,2):
                legajo = frame.iloc[renglon,1]
                fecha = frame.index[renglon]
                cero = pd.to_datetime(('{} 00:00').format(fecha))  
                if frame.iloc[renglon,posicion] == cero:# and frame.iloc[renglon,posicion+ 1] == cero:
                    toErase.append(fecha)
        if len(toErase) != 0:
            for fecha in toErase:
                frame.drop([fecha],inplace=True)
        frame.reset_index(inplace=True)
        frame = frame[columnas]
        return frame
    
    def borradoRegistroIndividual(self,frame):
        fecha = frame.iloc[-1,3]
        nombre = frame.iloc[-1,1]
        
        horaAnalizada = frame.iloc[-1,6]
        horaAnalizadaAnterior = frame.iloc[-1,5]
        horaAnalizadaSiguiente = frame.iloc[-1,7]
        

        
        if horaAnalizada > pd.to_datetime(('{} 19:30').format(fecha)) and horaAnalizadaAnterior < pd.to_datetime(('{} 10:00').format(fecha)) and \
            horaAnalizadaSiguiente == pd.to_datetime(('{} 00:00').format(fecha)):
            frame.iloc[-1,6] = pd.to_datetime(('{} 00:00').format(fecha))
        return frame

class CalculadorHoras:
    
    def horasTurnoRotativo(self,frame,fila,fecha,dia,area,ingreso='00:00',egreso='08:00'):
  
        horas_trabajadas = 0
        if area in rotativosInyeccion:
            horaIngreso = pd.to_datetime(('{} {}').format(fecha,ingreso))
            if egreso != '00:00':
                horaEgreso = pd.to_datetime(('{} {}').format(fecha,egreso))
            else:
                horaEgreso = pd.to_datetime(('{} {}').format(fecha,egreso))
                horaEgreso = horaEgreso + timedelta(days=1)
        
        else:
            horaEgreso = pd.to_datetime(('{} {}').format(fecha,egreso))
            if ingreso != '23:00':
                horaIngreso = pd.to_datetime(('{} {}').format(fecha,ingreso))
            else:
                horaIngreso = pd.to_datetime(('{} {}').format(fecha,ingreso))
                horaIngreso = horaIngreso - timedelta(days=1) 
                
        legajo = frame.iloc[fila,0]

        
        for idx in range(4,14,2):
                if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                   horas_trabajadas = 0
                   break
                elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                   break
                else:# Aca comienza a contar las horas trabajadas reales.
               
                    if horaIngreso - timedelta(hours=5) <= frame.iloc[fila,idx] <= horaIngreso + timedelta(hours=1): #Primer Turno

                        if frame.iloc[fila,idx] <= horaIngreso:
                            if frame.iloc[fila,idx +1] >= horaEgreso:

                                horas_trabajadas += (horaEgreso - horaIngreso).seconds

                            else:

                                horas_trabajadas += (frame.iloc[fila,idx +1] - horaIngreso).seconds
                        else:                                                               
                            if frame.iloc[fila,idx +1] >= horaEgreso:
                               
                                horas_trabajadas += (horaEgreso - frame.iloc[fila,idx]).seconds
                            else:
  
                                horas_trabajadas += (frame.iloc[fila,idx +1] - frame.iloc[fila,idx]).seconds
                                
                    else:

                        if frame.iloc[fila,idx +1] >= horaEgreso:
                           
                            horas_trabajadas += ( horaEgreso - frame.iloc[fila,idx]).seconds
                        else:
                            
                            horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds
                            
                # msg = 'Contabilizando horas en los Medio Dia, se trabajo {}'.format(horas_trabajadas)
                # logger.info(msg)

        return horas_trabajadas
    
    def horasTrabajadasRotativos(self,frame,area,mediosDias=[]):
        
        frame['H.Norm'] = 0
        frame['H. 50'] = 0
        frame['H. 100'] = 0  
        

        for fila in range(len(frame)):
            if area in rotativosInyeccion:       
                    horas_trabajadas = 0
                    fecha = frame.iloc[fila,3]
                    dia = frame.iloc[fila,2]
                    ayer = fecha - datetime.timedelta(days=1)
                    mañana = fecha + datetime.timedelta(days=1)
                    
                    primerIngreso = '08:00'
                    segundoIngreso = '16:00'
                    tercerIngreso = '00:00'
                    
                    primerSalida = '16:00'
                    segundaSalida = '00:00'
                    tercerSalida = '08:00'
                    
                    horaSalidaSabado = '13:00'
                    medioDia = '12:30'
                    
                    
                        
                    turnoMañanaPrimerIngreso = pd.to_datetime(('{} {}').format(fecha,primerIngreso))
                    turnoTardeIngreso = pd.to_datetime(('{} {}').format(fecha,segundoIngreso))
                    turnoNocheIngreso = pd.to_datetime(('{} {}').format(fecha,tercerIngreso))
                    
                    
                    turnoMañanaPrimerSalida = pd.to_datetime(('{} {}').format(fecha,primerSalida))
                    turnoTardeSalida = pd.to_datetime(('{} {}').format(mañana,segundaSalida))
                    turnoNocheSalida = pd.to_datetime(('{} {}').format(fecha,tercerSalida))                   
                    
                    
                    cero = pd.to_datetime(('{} 00:00').format(fecha))
                    medioDia = pd.to_datetime(('{} 12:00').format(fecha))
                    salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
    
            
                    for idx in range(4,14,2):
                        if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                            horas_trabajadas = 0
                            break
                        elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                            break
                        else:# Aca comienza a contar las horas trabajadas reales.
                            
                            if dia == 'Sábado':
                                horas_trabajadas = self.horasTrabajadasRotativosSabado(frame, fila, fecha, area)
                                #horas_trabajadas = self.horasTrabajadasSabado(frame, fila, fecha, dia)
                            elif dia in mediosDias:
                                pass
                                #horas_trabajadas = self.horasTrabajadasMedioDia(frame, fila, fecha, dia)
                            else:  

                                #ACA ESTA VIENDO QUE TIPO DE HORARIO ARRANCA Y DIRECCIONA AL TIPO DE CALCULADOR DE HORAS.
                                if frame.iloc[fila,idx] > turnoMañanaPrimerIngreso - timedelta(hours=4) and \
                                    frame.iloc[fila,idx] < turnoMañanaPrimerIngreso + timedelta(hours=4): #Primer Turno

                                        horas_trabajadas += self.horasTurnoRotativo(frame, fila, fecha, dia,area = area, ingreso='08:00',egreso='16:00')
                                        break
                                elif frame.iloc[fila,idx] > turnoTardeIngreso - timedelta(hours=4) and \
                                    frame.iloc[fila,idx] < turnoTardeIngreso + timedelta(hours=4): #Segundo Turno

                                        horas_trabajadas += self.horasTurnoRotativo(frame, fila, fecha, dia,area = area, ingreso='16:00',egreso='00:00')
                                        break
                                
                                elif frame.iloc[fila,idx] > turnoNocheIngreso - timedelta(hours=4,minutes=30) and \
                                    frame.iloc[fila,idx] < turnoNocheIngreso + timedelta(hours=4,minutes=30): #Segundo Turno

                                        horas_trabajadas += self.horasTurnoRotativo(frame, fila, fecha, dia,area = area, ingreso='00:00',egreso='08:00')
                                        break
                    
                    horas_trabajadas = round(horas_trabajadas /3600,2)
                    
                     
                    frame.iloc[fila,14] = horas_trabajadas
            else:
                
                    horas_trabajadas = 0
                    fecha = frame.iloc[fila,3]
                    dia = frame.iloc[fila,2]
                    ayer = fecha - datetime.timedelta(days=1)
                    mañana = fecha + datetime.timedelta(days=1)
                    
                    primerIngreso = '07:00'
                    segundoIngreso = '15:00'
                    tercerIngreso = '23:00'
                    
                    primerSalida = '15:00'
                    segundaSalida = '23:00'
                    tercerSalida = '07:00'
                    
                    horaSalidaSabado = '13:00'
                    medioDia = '12:30'
                    
                    
                        
                    turnoMañanaPrimerIngreso = pd.to_datetime(('{} {}').format(fecha,primerIngreso))
                    turnoTardeIngreso = pd.to_datetime(('{} {}').format(fecha,segundoIngreso))
                    turnoNocheIngreso = pd.to_datetime(('{} {}').format(ayer,tercerIngreso))
                    
                    
                    turnoMañanaPrimerSalida = pd.to_datetime(('{} {}').format(fecha,primerSalida))
                    turnoTardeSalida = pd.to_datetime(('{} {}').format(fecha,segundaSalida))
                    turnoNocheSalida = pd.to_datetime(('{} {}').format(fecha,tercerSalida))                   
                    
                    
                    cero = pd.to_datetime(('{} 00:00').format(fecha))
                    medioDia = pd.to_datetime(('{} 12:00').format(fecha))
                    salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
    
            
                    for idx in range(4,14,2):
                        if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                            horas_trabajadas = 0
                            break
                        elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                            break
                        else:# Aca comienza a contar las horas trabajadas reales.
                            
                            if dia == 'Sábado':
                                horas_trabajadas = self.horasTrabajadasRotativosSabado(frame, fila, fecha, area)
                                #horas_trabajadas = self.horasTrabajadasSabado(frame, fila, fecha, dia)
                            elif dia in mediosDias:
                                pass
                                #horas_trabajadas = self.horasTrabajadasMedioDia(frame, fila, fecha, dia)
                            else:  
                                
                                #ACA ESTA VIENDO QUE TIPO DE HORARIO ARRANCA Y DIRECCIONA AL TIPO DE CALCULADOR DE HORAS.
                                if frame.iloc[fila,idx] > turnoMañanaPrimerIngreso - timedelta(hours=4) and \
                                    frame.iloc[fila,idx] < turnoMañanaPrimerIngreso + timedelta(hours=4): #Primer Turno
                                          
                                        horas_trabajadas += self.horasTurnoRotativo(frame, fila, fecha, dia,area = area, ingreso='07:00',egreso='15:00')
                                        break
                                elif frame.iloc[fila,idx] > turnoTardeIngreso - timedelta(hours=4) and \
                                    frame.iloc[fila,idx] < turnoTardeIngreso + timedelta(hours=4): #Segundo Turno
                                        
                                        horas_trabajadas += self.horasTurnoRotativo(frame, fila, fecha, dia,area = area, ingreso='15:00',egreso='23:00')
                                        break
                                
                                elif frame.iloc[fila,idx] > turnoNocheIngreso - timedelta(hours=4) and \
                                    frame.iloc[fila,idx] < turnoNocheIngreso + timedelta(hours=4): #Segundo Turno
                                        
                                        horas_trabajadas += self.horasTurnoRotativo(frame, fila, fecha, dia,area = area, ingreso='23:00',egreso='07:00')
                                        break
                    horas_trabajadas = round(horas_trabajadas /3600,2)
                    
                     
                    frame.iloc[fila,14] = horas_trabajadas
                
                
                    
                    # msg = 'Contabilizando horas dias normales, se trabajo {}'.format(horas_trabajadas)
                    # logger.info(msg)
        return frame
    
    def horasExtrasTrabajadasRotativos(self,frame,area,feriados=[],mediosDias=[],toleranciaHoraria=1):
        """
        
        Parameters
        ----------
        frame : dataFrame
            Frame con todos los registros (datos del operarios e ingresos/egresos por dia).
        feriados : List, optional
            DESCRIPTION. The default is []. Opcional, una lista con los dias feriados dentro de los dias a analizar
        toleranciaHoraria : int, optional
            DESCRIPTION. The default is 1. Opcional, un numero que representa el valor minimo de minutos que tiene que pasar desde
            la hora de ingreso/egreso para que se empiezen a computar minutos extras.

        Returns
        -------
        frame : dataFrame
            dataFrame con todos los campos calculados Horas extras al 50% y al 100%.

        """

        ingreso = '08:00'
        salida = '16:48'
        horaSalidaSabado = '13:00'
        horaSalidaMedioDia = '12:30'
        for fila in range(len(frame)):
            if area in rotativosInyeccion:
                
                horas_trabajadas = 0
                minutosExtras100 = 0
                minutosExtras50 = 0 
                
                fecha = frame.iloc[fila,3]
                dia = frame.iloc[fila,2]
                ayer = fecha - datetime.timedelta(days=1)
                mañana = fecha + datetime.timedelta(days=1)
                
                primerIngreso = '08:00'
                segundoIngreso = '16:00'
                tercerIngreso = '00:00'
                
                primerSalida = '16:00'
                segundaSalida = '00:00'
                tercerSalida = '08:00'
                
                horaSalidaSabado = '13:00'
                medioDia = '12:30'
                    
                    
                        
                turnoMañanaPrimerIngreso = (pd.to_datetime(('{} {}').format(fecha,primerIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                turnoTardeIngreso = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                turnoNocheIngreso = (pd.to_datetime(('{} {}').format(fecha,tercerIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                
                
                turnoMañanaPrimerSalida = (pd.to_datetime(('{} {}').format(fecha,primerSalida))+ datetime.timedelta(minutes=toleranciaHoraria))
                turnoTardeSalida = (pd.to_datetime(('{} {}').format(mañana,segundaSalida))+ datetime.timedelta(minutes=toleranciaHoraria))
                turnoNocheSalida = (pd.to_datetime(('{} {}').format(fecha,tercerSalida)) + datetime.timedelta(minutes=toleranciaHoraria))                 
                
                
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                medioDia = pd.to_datetime(('{} 12:00').format(fecha))
                salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                
            
                       
                salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                salidaMedioDia = pd.to_datetime(('{} {}').format(fecha,horaSalidaMedioDia))
                ceroHoy = pd.to_datetime(('{} 00:00').format(fecha))#Media noche del dia en analisis
            
            
                ingresoOperario = frame.iloc[fila,4]
                salidaOperario = 0
                fechaIngreso = ingresoOperario.date()
            
                for x in range(5,15,2):
                    if frame.iloc[fila,x] ==  ceroHoy:
                        salidaOperario = frame.iloc[fila,x -2]
                        break
                fechaSalidaOperarios = salidaOperario.date()
                # print(feriados,'     ',fechaIngreso,'   ',fechaSalidaOperarios, fechaIngreso in feriados )
                # print('--->',fechaSalidaOperarios in feriados or fechaIngreso in feriados or fecha in feriados)
                if dia == 'Sábado':
                    if ingresoOperario < turnoMañanaPrimerIngreso - timedelta(hours=4):
                        if ingresoOperario < turnoNocheIngreso:                                                      
                            minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                        elif salidaOperario > turnoNocheSalida:
                            minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600
                    else:
                        if ingresoOperario <= turnoMañanaPrimerIngreso:
                            minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                        minutosExtras50 += frame.iloc[fila,14]
                        frame.iloc[fila,14] = 0
                    
                    if salidaOperario > salidaSabado: #Checkea si es sabado pasadas las 13
                        minutosExtras100 += ((salidaOperario - salidaSabado).seconds)/3600
                        frame.iloc[fila,16] = round(minutosExtras100,2) #Asigna las horas extras al 100%
                        
                        
                
                elif fechaSalidaOperarios in feriados or fechaIngreso in feriados:
                    
                    if fechaIngreso in feriados and fechaSalidaOperarios not in feriados: 
                        
                        
                        if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:
                            if ingresoOperario < turnoNocheIngreso:                                                     
                                minutosExtras100 += ((cero - ingresoOperario).seconds)/3600
                                frame.iloc[fila,16] = round(minutosExtras100,2)
                                
                            if salidaOperario > turnoNocheSalida:
                                minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600 
                        
                        
                        if  turnoTardeIngreso- timedelta(hours=4) < ingresoOperario < turnoTardeIngreso + timedelta(hours=2):  
                            if salidaOperario > turnoTardeSalida:
                                minutosExtras100 += ((turnoTardeSalida - ingresoOperario).seconds)/3600
                                frame.iloc[fila,16] = round(minutosExtras100,2)
                                frame.iloc[fila,14] = 0
                                
                            else:
                                minutosExtras100 += ((salidaOperario - ingresoOperario).seconds)/3600
                                frame.iloc[fila,16] = round(minutosExtras100,2)
                                frame.iloc[fila,14] = 0
                                
                                
                                
                        
                    elif fechaSalidaOperarios in feriados and fechaIngreso  not in feriados:
                        if salidaOperario > cero:
                            minutosExtras100 += ((salidaOperario - cero).seconds)/3600
                            frame.iloc[fila,16] = round(minutosExtras100,2)
                            frame.iloc[fila,14] = 0
                        if ingresoOperario < turnoNocheIngreso:
                            minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600 
                    
                    elif fechaSalidaOperarios in feriados and fechaIngreso  in feriados:
                        minutosExtras100 += ((salidaOperario - ingresoOperario).seconds)/3600
                        frame.iloc[fila,14] = 0
                        frame.iloc[fila,16] = round(minutosExtras100,2)
                        
                        
                
                elif fechaSalidaOperarios in mediosDias or fechaIngreso in mediosDias:
                    
                    if fechaIngreso in mediosDias and fechaSalidaOperarios not in mediosDias:
                        
                        
                        if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:
                            if ingresoOperario < turnoNocheIngreso:                                                     
                                minutosExtras100 += ((cero - ingresoOperario).seconds)/3600
                                frame.iloc[fila,16] = round(minutosExtras100,2)
                                
                            if salidaOperario > turnoNocheSalida:
                                minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600 
                                
                             
                        
                    elif fechaSalidaOperarios in mediosDias and fechaIngreso  not in mediosDias:
                        if salidaOperario > salidaMedioDia:
                            minutosExtras100 += ((salidaOperario - salidaMedioDia).seconds)/3600
                            frame.iloc[fila,16] = round(minutosExtras100,2)
                        
                        if ingresoOperario < turnoNocheIngreso:
                            minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                        
                        if turnoNocheSalida < salidaOperario < salidaMedioDia:
                            minutosExtras50 += (( salidaOperario - turnoNocheSalida).seconds)/3600
                    
                    
                    elif fechaSalidaOperarios in mediosDias and fechaIngreso  in mediosDias:
                        
                         if salidaOperario > salidaMedioDia and ingresoOperario > salidaMedioDia: 
                                minutosExtras100 += ((salidaOperario - ingresoOperario).seconds)/3600
                                frame.iloc[fila,16] = round(minutosExtras100,2)
                                frame.iloc[fila,14] = 0
                         else:
                            if turnoMañanaPrimerIngreso- timedelta(hours=2) < ingresoOperario < turnoMañanaPrimerIngreso + timedelta(hours = 1):
                                if ingresoOperario < turnoMañanaPrimerIngreso:                                                     
                                    minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                                    
                                    if salidaOperario > salidaMedioDia:
                                        minutosExtras100 += ((salidaOperario - salidaMedioDia).seconds)/3600 
                                        frame.iloc[fila,16] = round(minutosExtras100,2)
                                        frame.iloc[fila,14] = round(((salidaMedioDia - turnoMañanaPrimerIngreso).seconds)/3600,2)
                                    else:
                                        frame.iloc[fila,14] = round(((salidaOperario - turnoMañanaPrimerIngreso).seconds)/3600,2)
                                
                                else:
                                    if salidaOperario > salidaMedioDia:
                                        minutosExtras100 += ((salidaOperario - salidaMedioDia).seconds)/3600 
                                        frame.iloc[fila,14] = round(((salidaMedioDia - ingresoOperario).seconds)/3600,2)
                                    else:
                                        frame.iloc[fila,14] = round(((salidaOperario - ingresoOperario).seconds)/3600,2)
                                        
                            if turnoNocheIngreso- timedelta(hours=2) < ingresoOperario < turnoNocheIngreso + timedelta(hours = 1):
                                if ingresoOperario < turnoNocheIngreso:                                                     
                                   pass
                                
                                else:
                                    if salidaOperario > turnoNocheSalida:
                                        minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600 
                                        frame.iloc[fila,14] = round(((turnoNocheSalida - ingresoOperario).seconds)/3600,2)
                                    else:
                                        frame.iloc[fila,14] = round(((salidaOperario - ingresoOperario).seconds)/3600,2)
                                    
                                
                                    
                        
                        # if salidaOperario > salidaMedioDia: 
                        #     if ingresoOperario > salidaMedioDia:
                        #         minutosExtras100 += ((salidaOperario - ingresoOperario).seconds)/3600
                        #         frame.iloc[fila,16] = round(minutosExtras100,2)
                        #         frame.iloc[fila,14] = 0
                        #     else:
                        #         if turnoMañanaPrimerIngreso- timedelta(hours=2) < ingresoOperario < turnoMañanaPrimerIngreso:
                        #             minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                    
                    
                    
                    # if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:
                    #     if ingresoOperario < turnoNocheIngreso:                                                      
                    #         minutosExtras100 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                    #         frame.iloc[fila,16] = round(minutosExtras100,2)
                    #     elif salidaOperario > turnoNocheSalida:
                    #         if salidaOperario > horaSalidaMedioDia:
                    #             minutosExtras100 += ((salidaOperario - horaSalidaMedioDia).seconds)/3600
                    #         # else:
                    #         #     minutosExtras100 += ((salidaOperario - turnoNocheSalida).seconds)/3600
                    # else:                                
                    #     if salidaOperario > salidaMedioDia:
                    #         minutosExtras100 += round(((salidaOperario - salidaMedioDia).seconds)/3600,2)
                    #         frame.iloc[fila,16] = minutosExtras100
                            
                    #     if turnoMañanaPrimerIngreso - timedelta(hours=4) < ingresoOperario < turnoMañanaPrimerIngreso:                            
                    #         minutosExtras50 += round(((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600,2)
                    #         frame.iloc[fila,14] = round(((salidaMedioDia - turnoMañanaPrimerIngreso ).seconds)/3600,2)
                    #     else:
                    #         frame.iloc[fila,14] = round(((salidaMedioDia - ingresoOperario ).seconds)/3600,2)
                            
    
                
                else:

                    #Ingreso-Egreso 3er Turno-----------------
                    if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:

                        if ingresoOperario < turnoNocheIngreso:
                                                     
                            minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                    if turnoNocheSalida - timedelta(hours=2) < salidaOperario < turnoNocheSalida + timedelta(hours=4):

                            if salidaOperario > turnoNocheSalida:
                                minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600
                    
                    #Ingreso-Egreso 1er Turno-----------------
                    if turnoMañanaPrimerIngreso - timedelta(hours=4) < ingresoOperario < turnoMañanaPrimerIngreso: 
                                                   
                        minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                    if turnoMañanaPrimerSalida - timedelta(hours=2) < salidaOperario < turnoMañanaPrimerSalida  + timedelta(hours=4):
                            if salidaOperario > turnoMañanaPrimerSalida:
 
                                minutosExtras50 += ((salidaOperario - turnoMañanaPrimerSalida).seconds)/3600
                            
                   
                    #Ingreso-Egreso 2do Turno
                    if turnoTardeIngreso - timedelta(hours=4) < ingresoOperario < turnoTardeIngreso: 
                                                     
                        minutosExtras50 += ((turnoTardeIngreso - ingresoOperario).seconds)/3600
                    if turnoTardeSalida - timedelta(hours=2) < salidaOperario < turnoTardeSalida  + timedelta(hours=4):
                        if salidaOperario > turnoTardeSalida:
                             minutosExtras50 += ((salidaOperario - turnoTardeSalida).seconds)/3600

                frame.iloc[fila,15] = round(minutosExtras50,2) #Asigna las horas extras al 50%
        
            else: 
                                      
                horas_trabajadas = 0
                minutosExtras100 = 0
                minutosExtras50 = 0 
                
                fecha = frame.iloc[fila,3]
                dia = frame.iloc[fila,2]
                ayer = fecha - datetime.timedelta(days=1)
                mañana = fecha + datetime.timedelta(days=1)
                
                primerIngreso = '07:00'
                segundoIngreso = '15:00'
                tercerIngreso = '23:00'
                
                primerSalida = '15:00'
                segundaSalida = '23:00'
                tercerSalida = '07:00'
                
                horaSalidaSabado = '13:00'
                medioDia = '12:30'
                
                
                
                turnoMañanaPrimerIngreso = (pd.to_datetime(('{} {}').format(fecha,primerIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                turnoTardeIngreso = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                turnoNocheIngreso = (pd.to_datetime(('{} {}').format(ayer,tercerIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                
                
                turnoMañanaPrimerSalida = (pd.to_datetime(('{} {}').format(fecha,primerSalida))+ datetime.timedelta(minutes=toleranciaHoraria))
                turnoTardeSalida = (pd.to_datetime(('{} {}').format(mañana,segundaSalida))+ datetime.timedelta(minutes=toleranciaHoraria))
                turnoNocheSalida = (pd.to_datetime(('{} {}').format(fecha,tercerSalida)) + datetime.timedelta(minutes=toleranciaHoraria))                 
                
                    
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                medioDia = pd.to_datetime(('{} 12:00').format(fecha))
                salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                    
                
                           
                salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                salidaMedioDia = pd.to_datetime(('{} {}').format(fecha,horaSalidaMedioDia))
                ceroHoy = pd.to_datetime(('{} 00:00').format(fecha))#Media noche del dia en analisis
                
                
                
                #----------------------------------------------------
                
                ingresoOperario = frame.iloc[fila,4]
                fechaIngreso = ingresoOperario.date()
                salidaOperario = 0
                
                for x in range(5,15,2):
                    if frame.iloc[fila,x] ==  ceroHoy:
                        salidaOperario = frame.iloc[fila,x -2]
                        break
                fechaSalidaOperarios = salidaOperario.date()
                
                if dia == 'Sábado':
                    if ingresoOperario < turnoMañanaPrimerIngreso - timedelta(hours=4):
                        if ingresoOperario < turnoNocheIngreso:                                                      
                                minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                        if salidaOperario > turnoNocheSalida:
                                minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600
                        else:
                            if ingresoOperario <= turnoMañanaPrimerIngreso:
                                minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                            minutosExtras50 += frame.iloc[fila,14]
                            frame.iloc[fila,14] = 0
                        
                    if salidaOperario > salidaSabado: #Checkea si es sabado pasadas las 13
                       minutosExtras100 += ((salidaOperario - salidaSabado).seconds)/3600
                       frame.iloc[fila,16] = round(minutosExtras100,2) #Asigna las horas extras al 100%
                       
                       
        
                elif fechaSalidaOperarios in feriados or fechaIngreso in feriados:
                    
                    if fechaIngreso in feriados and fechaSalidaOperarios not in feriados:                    
                        if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:
                            if ingresoOperario < turnoNocheIngreso:                                                     
                                minutosExtras100 += ((cero - ingresoOperario).seconds)/3600
                                frame.iloc[fila,16] = round(minutosExtras100,2)
                                
                            if salidaOperario > turnoNocheSalida:
                                minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600
                                frame.iloc[fila,14] = round(((turnoNocheSalida - cero).seconds)/3600,2)
                            else:
                                frame.iloc[fila,14] = round(((salidaOperario - cero).seconds)/3600,2)                         
     
                    elif fechaSalidaOperarios in feriados and fechaIngreso  not in feriados:
                        if salidaOperario > cero:
                            minutosExtras100 += ((salidaOperario - cero).seconds)/3600
                            frame.iloc[fila,16] = round(minutosExtras100,2)
                            frame.iloc[fila,14] = 0
                        if ingresoOperario < turnoNocheIngreso:
                            minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600 
                    
                    elif fechaSalidaOperarios in feriados and fechaIngreso  in feriados:
                        minutosExtras100 += ((salidaOperario - ingresoOperario).seconds)/3600
                        frame.iloc[fila,14] = 0
                        frame.iloc[fila,16] = round(minutosExtras100,2)
                        
                        
                        
                    
                elif fechaSalidaOperarios in mediosDias or fechaIngreso in mediosDias:
                    
                    if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:
                        if ingresoOperario < turnoNocheIngreso:                                                      
                            minutosExtras100 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                            frame.iloc[fila,16] = round(minutosExtras100,2)
                        elif salidaOperario > turnoNocheSalida:
                            if salidaOperario > horaSalidaMedioDia:
                                minutosExtras100 += ((salidaOperario - horaSalidaMedioDia).seconds)/3600
                            else:
                                minutosExtras100 += ((salidaOperario - turnoNocheSalida).seconds)/3600
                    else:
                        
                        if turnoMañanaPrimerIngreso - timedelta(hours=4) < ingresoOperario < turnoMañanaPrimerIngreso:
                            if ingresoOperario < turnoMañanaPrimerIngreso:
                                minutosExtras50 += round(((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600,2)
                                frame.iloc[fila,14] = round(((salidaMedioDia - turnoMañanaPrimerIngreso ).seconds)/3600,2)
                            else:
                                frame.iloc[fila,14] = round(((salidaMedioDia - ingresoOperario).seconds)/3600,2)
                            
                            if salidaOperario > salidaMedioDia:
                                minutosExtras100 += round(((salidaOperario - salidaMedioDia).seconds)/3600,2)
                                frame.iloc[fila,16] = minutosExtras100
                       
                        if salidaMedioDia < ingresoOperario < turnoTardeIngreso:
                            
                            frame.iloc[fila,16] = round(((salidaOperario - ingresoOperario ).seconds)/3600,2)
                            frame.iloc[fila,14] = 0
                            
                        
        
                    
                else:

                        #Ingreso-Egreso 3er Turno-----------------
                        if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso:

                            if ingresoOperario < turnoNocheIngreso:
                                                      
                                minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600
                            if salidaOperario > turnoNocheSalida:

                                minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600

                        #Ingreso-Egreso 1er Turno-----------------
                        if turnoMañanaPrimerIngreso - timedelta(hours=4) < ingresoOperario < turnoMañanaPrimerIngreso: 
                                               
                            minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                        if turnoMañanaPrimerSalida - timedelta(hours=2) < salidaOperario < turnoMañanaPrimerSalida + timedelta(hours=3):
                            if salidaOperario > turnoMañanaPrimerSalida:

                                minutosExtras50 += ((salidaOperario - turnoMañanaPrimerSalida).seconds)/3600
                                
                        #Ingreso-Egreso 2do Turno
                        if turnoTardeIngreso - timedelta(hours=4) < ingresoOperario < turnoTardeIngreso: 
                                                   
                            minutosExtras50 += ((turnoTardeIngreso - ingresoOperario).seconds)/3600
                        if turnoTardeSalida - timedelta(hours=2) < salidaOperario < turnoTardeSalida + timedelta(hours=3):
                            if salidaOperario > turnoTardeSalida:

                                minutosExtras50 += ((salidaOperario - turnoTardeSalida).seconds)/3600
              
                frame.iloc[fila,15] = round(minutosExtras50,2) #Asigna las horas extras al 50%
            
        # msg = 'Contabilizando horas al 50 y 100'
        # logger.info(msg)
        return frame

    def horasTrabajadasRotativosSabado(self,frame,fila,fecha,area):
  
        horas_trabajadas = 0

        if area in rotativosInyeccion: 
            ingreso = '08:00'
            egreso = '08:00'
            horaSalidaSabado = '13:00'
            ceroSabado = '00:00'
            #---------Turno Inyeccion-Mecanizado
            horaIngresoTercer = pd.to_datetime(('{} {}').format(fecha,ceroSabado))
            horaEgresoTercer = pd.to_datetime(('{} {}').format(fecha,egreso))
            
            horaIngresoPrimer = pd.to_datetime(('{} {}').format(fecha,ingreso))
            horaEgresoPrimer = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))            
            
            for idx in range(4,14,2):
                    if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                       horas_trabajadas = 0
                       break
                    elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                       break
                    else:# Aca comienza a contar las horas trabajadas reales.
                        if horaIngresoTercer - timedelta(hours=4) < frame.iloc[fila,idx] < horaIngresoTercer:
                            if horaEgresoTercer < frame.iloc[fila,idx +1] :
                                horas_trabajadas += (horaEgresoTercer - horaIngresoTercer).seconds                        
                            else:
                                horas_trabajadas += (frame.iloc[fila,idx + 1] - horaIngresoTercer).seconds
                        else:
                            if horaIngresoTercer < frame.iloc[fila,idx] < horaIngresoTercer + timedelta(hours=2):
                                if frame.iloc[fila,idx +1] < horaEgresoTercer  :
                                    horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds                        
                                else:
                                    horas_trabajadas += (horaEgresoTercer  - horaIngresoTercer).seconds
                        
                        
                        
                        #Ingreso de los sabados pero a partir de las 8
                        if horaIngresoPrimer - timedelta(hours=2) < frame.iloc[fila,idx] < horaIngresoPrimer :
                            if horaEgresoPrimer < frame.iloc[fila,idx +1] :
                                horas_trabajadas += (horaEgresoPrimer - horaIngresoPrimer).seconds
                            else:
                                horas_trabajadas += (frame.iloc[fila,idx +1] - horaIngresoPrimer).seconds
                        else:
                            if horaIngresoPrimer < frame.iloc[fila,idx] < horaIngresoPrimer + timedelta(hours=2):
                                if horaEgresoPrimer < frame.iloc[fila,idx +1] :
                                    horas_trabajadas += (horaEgresoPrimer - frame.iloc[fila,idx]).seconds
                                else:
                                    horas_trabajadas += (frame.iloc[fila,idx +1] - frame.iloc[fila,idx]).seconds
        # else:
        #     ingreso = '08:00'
        #     egreso = '07:00'
        #     horaSalidaSabado = '13:00'
        #     ceroSabado = '00:00'
        #     #---------Turno Inyeccion-Mecanizado
        #     horaIngresoTercer = pd.to_datetime(('{} {}').format(fecha,ceroSabado)) - timedelta(hours=1)
        #     horaEgresoTercer = pd.to_datetime(('{} {}').format(fecha,egreso))
            
        #     horaIngresoPrimer = pd.to_datetime(('{} {}').format(fecha,ingreso))
        #     horaEgresoPrimer = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))            
            
        #     for idx in range(4,14,2):
        #             if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
        #                horas_trabajadas = 0
        #                break
        #             elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
        #                break
        #             else:# Aca comienza a contar las horas trabajadas reales.
        #                 if horaIngresoTercer - timedelta(hours=4) < frame.iloc[fila,idx] :
        #                     if horaEgresoTercer < frame.iloc[fila,idx +1] :
        #                         horas_trabajadas += (horaEgresoTercer - horaIngresoTercer).seconds                        
        #                     else:
        #                         horas_trabajadas += (frame.iloc[fila,idx + 1] - horaIngresoTercer).seconds
        #                 else:
        #                     if  horaIngresoTercer < frame.iloc[fila,idx] < horaIngresoTercer + timedelta(hours=2):
        #                         if frame.iloc[fila,idx +1] < horaEgresoTercer  :
        #                             horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds                        
        #                         else:
        #                             horas_trabajadas += (horaEgresoTercer  - horaIngresoTercer).seconds
                        
        #                 if horaIngresoPrimer - timedelta(hours=2) < frame.iloc[fila,idx]:
        #                     if horaEgresoPrimer < frame.iloc[fila,idx +1] :
        #                         horas_trabajadas += (horaEgresoPrimer - horaIngresoPrimer).seconds
        #                     else:
        #                         horas_trabajadas += (frame.iloc[fila,idx +1] - horaIngresoPrimer).seconds
        #                 else:
        #                     if horaEgresoPrimer < frame.iloc[fila,idx +1] :
        #                         horas_trabajadas += (horaEgresoPrimer - frame.iloc[fila,idx]).seconds
        #                     else:
        #                         horas_trabajadas += (frame.iloc[fila,idx +1] - frame.iloc[fila,idx]).seconds
                

                        # msg ='Contabilizando horas del sabado, se trabajo {}'.format(horas_trabajadas)
                        # logger.info(msg)
        return horas_trabajadas
    
    # def horasTrabajadasMedioDiaRotativos(self,frame,fila,area,fecha,dia,toleranciaHoraria=1):
        
    #     horas_trabajadas = 0

    #     ingreso = '08:00'
    #     salida = '16:48'
    #     horaSalidaSabado = '13:00'
    #     horaSalidaMedioDia = '12:30'
    #     for fila in range(len(frame)):
    #         if area in ['INYECCION','MECANIZADO']:
                
    #             horas_trabajadas = 0
    #             minutosExtras100 = 0
    #             minutosExtras50 = 0 
                
    #             fecha = frame.iloc[fila,3]
    #             dia = frame.iloc[fila,2]
    #             ayer = fecha - datetime.timedelta(days=1)
    #             mañana = fecha + datetime.timedelta(days=1)
                
    #             primerIngreso = '08:00'
    #             segundoIngreso = '16:00'
    #             tercerIngreso = '00:00'
                
    #             primerSalida = '16:00'
    #             segundaSalida = '00:00'
    #             tercerSalida = '08:00'
                
    #             horaSalidaSabado = '13:00'
    #             medioDia = '12:30'
                    
                    
                        
    #             turnoMañanaPrimerIngreso = (pd.to_datetime(('{} {}').format(fecha,primerIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
    #             turnoTardeIngreso = (pd.to_datetime(('{} {}').format(fecha,segundoIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
    #             turnoNocheIngreso = (pd.to_datetime(('{} {}').format(fecha,tercerIngreso))- datetime.timedelta(minutes=toleranciaHoraria))
                
                
    #             turnoMañanaPrimerSalida = (pd.to_datetime(('{} {}').format(fecha,primerSalida))+ datetime.timedelta(minutes=toleranciaHoraria))
    #             turnoTardeSalida = (pd.to_datetime(('{} {}').format(mañana,segundaSalida))+ datetime.timedelta(minutes=toleranciaHoraria))
    #             turnoNocheSalida = (pd.to_datetime(('{} {}').format(fecha,tercerSalida)) + datetime.timedelta(minutes=toleranciaHoraria))                 
                
                
    #             cero = pd.to_datetime(('{} 00:00').format(fecha))
    #             medioDia = pd.to_datetime(('{} 12:00').format(fecha))
    #             salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                
            
                       
    #             salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
    #             salidaMedioDia = pd.to_datetime(('{} {}').format(fecha,horaSalidaMedioDia))
    #             ceroHoy = pd.to_datetime(('{} 00:00').format(fecha))#Media noche del dia en analisis
            
            
    #             ingresoOperario = frame.iloc[fila,4]
    #             salidaOperario = 0
    #             fechaIngreso = ingresoOperario.date()
            
    #             for x in range(5,15,2):
    #                 if frame.iloc[fila,x] ==  ceroHoy:
    #                     salidaOperario = frame.iloc[fila,x -2]
    #                     break
    #             fechaSalidaOperarios = salidaOperario.date()
                

                    
    #             if fechaIngreso in dia and fechaSalidaOperarios not in dia: 
                        
          
    #                     if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso + timedelta(hours=1):
    #                         if ingresoOperario < turnoNocheIngreso: 
    #                                 if salidaOperario > turnoNocheSalida:
    #                                     horas_trabajadas = ((cero - turnoNocheSalida).seconds)/3600
    #                                     minutosExtras50 += ((salidaOperario - turnoNocheSalida).seconds)/3600 
    #                                 else:
    #                                     horas_trabajadas = ((cero - salidaOperario).seconds)/3600  

   
    #             elif fechaSalidaOperarios in dia and fechaIngreso  not in dia:
    #                     if  turnoNocheIngreso- timedelta(hours=5) < ingresoOperario < turnoNocheIngreso + timedelta(hours=1):
    #                         if ingresoOperario < turnoNocheIngreso: 
    #                                 if salidaOperario > salidaMedioDia:
    #                                     horas_trabajadas = ((cero - turnoNocheSalida).seconds)/3600 
    #                                 else:
    #                                     horas_trabajadas = ((cero - salidaOperario).seconds)/3600            
                        
    #                     if ingresoOperario < turnoNocheIngreso:
    #                         minutosExtras50 += ((turnoNocheIngreso - ingresoOperario).seconds)/3600 
                    
                    
    #                 elif fechaSalidaOperarios in mediosDias and fechaIngreso  in mediosDias:
    #                     if salidaOperario > salidaMedioDia: 
    #                         if ingresoOperario > salidaMedioDia:
    #                             minutosExtras100 += ((salidaOperario - ingresoOperario).seconds)/3600
    #                             frame.iloc[fila,16] = round(minutosExtras100,2)
    #                             frame.iloc[fila,14] = 0
    #                         else:
    #                             if turnoMañanaPrimerIngreso- timedelta(hours=2) < ingresoOperario < turnoMañanaPrimerIngreso:
    #                                 minutosExtras50 += ((turnoMañanaPrimerIngreso - ingresoOperario).seconds)/3600
                
                
    #     return horas_trabajadas
   
    def horasTrabajadasSabado(self,frame,fila,fecha,dia):
        
        ingreso = '08:00'
        horaSalidaSabado = '13:00'       
        
        horas_trabajadas = 0
        salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
        horaIngreso = pd.to_datetime(('{} {}').format(fecha,ingreso))

        
        for idx in range(4,14,2):
                if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                   horas_trabajadas = 0
                   break
                elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                   break
                else:# Aca comienza a contar las horas trabajadas reales.
                    if frame.iloc[fila,idx + 1] <= salidaSabado:

                        if frame.iloc[fila,idx] <= horaIngreso:

                            horas_trabajadas += (frame.iloc[fila,idx + 1] - horaIngreso).seconds                        
                        else:

                            horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds
                    else:

                        if frame.iloc[fila,idx] <= horaIngreso:

                            horas_trabajadas += (salidaSabado - horaIngreso).seconds  
                        else:

                            horas_trabajadas += (salidaSabado - frame.iloc[fila,idx]).seconds
                    # msg ='Contabilizando horas del sabado, se trabajo {}'.format(horas_trabajadas)
                    # logger.info(msg)

        return horas_trabajadas
    
    def horasTrabajadasMedioDia(self,frame,fila,fecha,dia):
        
        ingreso = '08:00'
        medioDia = '12:30'     
        
        horas_trabajadas = 0
        salidaMedioDia = pd.to_datetime(('{} {}').format(fecha,medioDia))
        horaIngreso = pd.to_datetime(('{} {}').format(fecha,ingreso))

        
        for idx in range(4,14,2):
                if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                   horas_trabajadas = 0
                   break
                elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                   break
                else:# Aca comienza a contar las horas trabajadas reales.
               
                    if frame.iloc[fila,idx + 1] <= salidaMedioDia:

                        if frame.iloc[fila,idx] <= horaIngreso: 

                            horas_trabajadas += (frame.iloc[fila,idx + 1] - horaIngreso).seconds                        
                        else:

                            horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds
                    else:

                        if frame.iloc[fila,idx] <= horaIngreso: 

                            horas_trabajadas += (salidaMedioDia - horaIngreso).seconds                        
                        else:

                            horas_trabajadas += (salidaMedioDia - frame.iloc[fila,idx]).seconds
                # msg = 'Contabilizando horas en los Medio Dia, se trabajo {}'.format(horas_trabajadas)
                # logger.info(msg)
        return horas_trabajadas
        
    
    def horasTrabajadas(self,frame,mediosDias=[]):
        import pandas as pd


        frame['H.Norm'] = 0
        frame['H. 50'] = 0
        frame['H. 100'] = 0
        
        ingreso = '08:00'
        salida = '16:48'
        horaSalidaSabado = '13:00'
        medioDia = '12:30'
        
        
        for fila in range(len(frame)):
                horas_trabajadas = 0
                fecha = frame.iloc[fila,3]
                dia = frame.iloc[fila,2]
                salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
                ceroHoy = pd.to_datetime(('{} 00:00').format(fecha))
                horaIngreso = pd.to_datetime(('{} {}').format(fecha,ingreso))
                horaSalida =pd.to_datetime(('{} {}').format(fecha,salida))
                horaSalidaMedioDia = pd.to_datetime(('{} {}').format(fecha,medioDia))

        
                for idx in range(4,14,2):
                    if frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] != pd.to_datetime(('{} 00:00').format(fecha)):
                        horas_trabajadas = 0
                        break
                    elif frame.iloc[fila,idx + 1] == pd.to_datetime(('{} 00:00').format(fecha)) and frame.iloc[fila,idx] == pd.to_datetime(('{} 00:00').format(fecha)): 
                        break
                    else:# Aca comienza a contar las horas trabajadas reales.

                        if dia == 'Sábado':
                            horas_trabajadas = self.horasTrabajadasSabado(frame, fila, fecha, dia)
                        elif dia in mediosDias:
                            horas_trabajadas = self.horasTrabajadasMedioDia(frame, fila, fecha, dia)
                        else:  
                            
                            if  horaIngreso - timedelta(hours = 3) < frame.iloc[fila,idx] <= horaIngreso:
                                if frame.iloc[fila,idx +1] <= horaSalida:
                                    if frame.iloc[fila,idx +1] > horaIngreso:
                                        horas_trabajadas += (frame.iloc[fila,idx + 1] - horaIngreso).seconds 
                                    else:
                                        continue
                                else:
                                    horas_trabajadas += (horaSalida - horaIngreso).seconds
                            
                            elif frame.iloc[fila,idx] > horaIngreso:
                                if frame.iloc[fila,idx +1] <= horaSalida:
                                    horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds 
                                else:
                                    horas_trabajadas += (horaSalida - frame.iloc[fila,idx]).seconds
                            
                            
                            
                            # if frame.iloc[fila,idx + 1] <= horaSalida and frame.iloc[fila,idx + 1] > horaIngreso:
                            #     if frame.iloc[fila,idx] <= horaIngreso: 
                            #         horas_trabajadas += (frame.iloc[fila,idx + 1] - horaIngreso).seconds 
                            #     else:
                            #         horas_trabajadas += (frame.iloc[fila,idx + 1] - frame.iloc[fila,idx]).seconds
                            # else:
                            #     if frame.iloc[fila,idx] < horaIngreso and frame.iloc[fila,idx+ 1] < horaIngreso :
                            #         continue
                            #     else:
                            #         horas_trabajadas += (horaSalida - frame.iloc[fila,idx]).seconds
                
                horas_trabajadas = round(horas_trabajadas /3600,2)

                 
                frame.iloc[fila,14] = horas_trabajadas
                
                # msg = 'Contabilizando horas dias normales, se trabajo {}'.format(horas_trabajadas)
                # logger.info(msg)
        return frame
    

    def horasExtrasTrabajadas(self,frame,feriados=[],mediosDias=[],toleranciaHoraria=1):
        """
        
        Parameters
        ----------
        frame : dataFrame
            Frame con todos los registros (datos del operarios e ingresos/egresos por dia).
        feriados : List, optional
            DESCRIPTION. The default is []. Opcional, una lista con los dias feriados dentro de los dias a analizar
        toleranciaHoraria : int, optional
            DESCRIPTION. The default is 1. Opcional, un numero que representa el valor minimo de minutos que tiene que pasar desde
            la hora de ingreso/egreso para que se empiezen a computar minutos extras.

        Returns
        -------
        frame : dataFrame
            dataFrame con todos los campos calculados Horas extras al 50% y al 100%.

        """
        
        #indices = [frame[frame['H.Norm.'] != 0].index[x] for x in range(len(frame[frame['H.Norm.Emp'] != 0]))]
        ingreso = '08:00'
        salida = '16:48'
        horaSalidaSabado = '13:00'
        horaSalidaMedioDia = '12:30'
        for fila in range(len(frame)):
            fecha = frame.iloc[fila,3]
            dia = frame.iloc[fila,2]
            minutosExtras100 = 0
            minutosExtras50 = 0            
            salidaSabado = pd.to_datetime(('{} {}').format(fecha,horaSalidaSabado))
            salidaMedioDia = pd.to_datetime(('{} {}').format(fecha,horaSalidaMedioDia))
            ceroHoy = pd.to_datetime(('{} 00:00').format(fecha))#Media noche del dia en analisis
            horaIngreso = pd.to_datetime(pd.to_datetime(('{} {}').format(fecha,ingreso)) - datetime.timedelta(minutes=toleranciaHoraria))
            #Hora con el limite de tiempo ya asiganado, en el ingreso se resta, en la salida se suma.
            horaSalida = pd.to_datetime(pd.to_datetime(('{} {}').format(fecha,salida)) + datetime.timedelta(minutes=toleranciaHoraria))#Same above.
            
            ingresoOperario = frame.iloc[fila,4]
            salidaOperario = 0
            
            for x in range(5,15,2):
                if frame.iloc[fila,x] ==  ceroHoy:
                    salidaOperario = frame.iloc[fila,x -2]
                    break

            if dia == 'Sábado':
                if frame.iloc[fila,4] <= horaIngreso:
                    minutosExtras50 += ((horaIngreso - frame.iloc[fila,4]).seconds)/3600
                

                minutosExtras50 += frame.iloc[fila,14]

                frame.iloc[fila,14] = 0
                
                if salidaOperario > salidaSabado: #Checkea si es sabado pasadas las 13
                    minutosExtras100 += ((salidaOperario - salidaSabado).seconds)/3600
                    frame.iloc[fila,16] = round(minutosExtras100,2) #Asigna las horas extras al 100%

            elif fecha in feriados:
                for posicion in range(4,12,2):
                    if frame.iloc[fila,posicion+ 1] == ceroHoy and frame.iloc[fila,posicion] == ceroHoy:
                        break
                    else:                        
                        minutosExtras100 += round(((frame.iloc[fila,posicion+ 1] - frame.iloc[fila,posicion]).seconds)/3600,2)
                frame.iloc[fila,14] = 0
                frame.iloc[fila,16] = round(minutosExtras100,2)
            
            elif fecha in mediosDias:
                if salidaOperario > salidaMedioDia:
                    minutosExtras100 += round(((salidaOperario - salidaMedioDia).seconds)/3600,2)
                    frame.iloc[fila,16] = round(minutosExtras100,2)
                if ingresoOperario < horaIngreso:
                    minutosExtras50 += ((horaIngreso - ingresoOperario).seconds)/3600

            
            else:
                if salidaOperario > horaSalida:
                    minutosExtras50 += ((salidaOperario - horaSalida).seconds)/3600
                if ingresoOperario < horaIngreso:
                    minutosExtras50 += ((horaIngreso - ingresoOperario).seconds)/3600

            
            frame.iloc[fila,15] = round(minutosExtras50,2) #Asigna las horas extras al 50%
        # msg = 'Contabilizando horas al 50 y 100'
        # logger.info(msg)
        return frame
    
    def restaRetrasosTardanzas(self,frame,mediosDias=[]):
        
        logger.info('Iniciando resta de minutos tardes')
        for x in range(len(frame)):
                legajo = frame.iloc[x,0]
                nombre = frame.iloc[x,1]
                dia = frame.iloc[x,2]
                fecha = frame.iloc[x,3]
                horaIngreso = pd.to_datetime(('{} 08:00').format(fecha))
                horaSalida = pd.to_datetime(('{} 16:48').format(fecha))
                horaSalidaMedioDia = pd.to_datetime(('{} 12:30').format(fecha))
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                limiteSalida = pd.to_datetime(('{} 14:00').format(fecha))
                tardanza = 0
                retiro= 0
                sinAlmuerzo = False
                
                for idx in range(5,13,2):
                    if frame.iloc[x,idx] == cero:
                        salida = frame.iloc[x,idx -2]
                        if idx == 7:
                            sinAlmuerzo = True
                        break
                
                if frame.iloc[x,4] > horaIngreso:
                    tardanza = round((((frame.iloc[x,4] - horaIngreso).seconds)/3600),2)               
                
                if fecha in mediosDias:
                    if salida < horaSalidaMedioDia:
                        retiro = round((((horaSalidaMedioDia - salida).seconds)/3600),2)                        
                elif fecha not in mediosDias:
                    if salida < horaSalida and dia != 'Sábado':
                        retiro = round((((horaSalida - salida).seconds)/3600),2)
                
                horas_trabajadas = frame.iloc[x,14]
                if sinAlmuerzo and dia != 'Sábado':
                    if salida > limiteSalida:
                        horas_trabajadas -= 0.5
                
                if horas_trabajadas > 8:
                    horas_trabajadas = 8 
                
                if dia != 'Sábado':
                    if horas_trabajadas >= 8 and (tardanza > 0 or retiro >0):
                        horas_trabajadas = horas_trabajadas- tardanza - retiro #Si entro tarde resta esos minutos
                    # elif horas_trabajadas < 8 and tardanza > 0:
                    #     horas_trabajadas = horas_trabajadas + (8 - horas_trabajadas) - tardanza
                
                               
                frame.iloc[x,14] = horas_trabajadas
            
        return frame

def repreguntar():
    decision = pyip.inputYesNo(prompt='¿Los datos ingresados son correctos? (SI/NO)  ',yesVal='SI',noVal='NO')
    print('\n')     
    if decision == 'SI':
        return True
    else:
        return False
    
def informeNoFichadas(frame,fechaInicio,fechaFin,mediosDias=[],feriados=[]):    
        """
        Parameters
        ----------
        frame : dataFrame
            Frame con todos los registros (datos del operarios e ingresos/egresos por dia).
        fechaInicio : datetime.date
            Fecha de inicio de analisis de los registros.
        fechaFin : datetime.date
            Fecha de fin de analisis de los registros.
        mediosDias : List, optional
            DESCRIPTION. The default is []. Opcional, una lista con los medios dias en los
            dias bajo analisis.
        feriados : List, optional
            DESCRIPTION. The default is [].Opcional, una lista con los feriados en los
            dias bajo analisis.
    
        Returns
        -------
        None.
        
        La forma de analisis es en forma de pares ingreso-egreso, en funcion de eso se aplica la logica.
    
        """

        try:
            logger.info('Comenzando escritura en el word')
            campo = 'H.Norm' #campo sobre el cual se filtra para ver las filas que tienen errores en los registros. Es siempre el mismo
            len_noMarca = len(frame[frame[campo] == 0]) #Devuelve la cantidad de filas a las que no se le calculo
            #horas trabajasdas debido a un error en los registros (faltan datos)
            doc = docx.Document()
            doc.add_heading(('Olvidos de fichaje entre {} y {}').format(fechaInicio,fechaFin), 0)
            c = doc.add_paragraph('Personal que no ha fichado: \n')
            msgWord = '\n\tEl dia {} ({}) el empleado {:10s} no ficho {}, {} a las {}.'
            ingreso = 'ingreso'
            reIngreso = 're-ingreso'
            salida = 'salida'
            almuerzo = salida+' o '+reIngreso
            
            for x in range(len_noMarca):
                legajo = frame.iloc[frame[frame[campo] == 0].index[x],0]
                nombre = frame.iloc[frame[frame[campo] == 0].index[x],1]
                dia = frame.iloc[frame[frame[campo] == 0].index[x],2]
                fecha = frame.iloc[frame[frame[campo] == 0].index[x],3]            
                #----------------- limites horarios -------------------
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                salidaMedioDia = pd.to_datetime(('{} 12:30').format(fecha))
                mitadMañana = pd.to_datetime(('{} 10:30').format(fecha))
                treceHoras = pd.to_datetime(('{} 13:00').format(fecha))
                
                
                if dia in mediosDias: #Checkea que sea medio dia. Analisis diferente para dias completos y medios dias. 
                    for posicion in range(4,14,2): #Itera sobre los ingresos/egresos
                        hora_ingreso = frame.iloc[frame[frame[campo] == 0].index[x],posicion]
                        hora_egreso = frame.iloc[frame[frame[campo] == 0].index[x],(posicion +1)]
                        
                        hora_ingresoHoras = '{:02d}:{:02d}'.format(hora_ingreso.hour,hora_ingreso.minute)#String con hora de ingreso
                        hora_salidaHoras = '{:02d}:{:02d}'.format(hora_egreso.hour,hora_egreso.minute)#String con hora de salida
                    
                        if (hora_ingreso and hora_egreso) != cero: #Si el par esta completo, sigue.
                            continue
                        
                        elif hora_ingreso != cero and hora_egreso == cero:
                            if posicion == 4 and hora_ingreso > mitadMañana:# condicion para ver si no ficho INGRESO (solo 1 par ingreso-egreso)
                                c.add_run((msgWord).format(dia,fecha,nombre,ingreso,salida,hora_ingresoHoras))
                                break
                            elif posicion == 4 and hora_ingreso < mitadMañana:# condicion para ver si no ficho EGRESO (solo 1 par ingreso-egreso)
                                c.add_run((msgWord).format(dia,fecha,nombre,salida,ingreso,hora_ingresoHoras))
                                break                           
                                
                            elif hora_ingreso < salidaMedioDia:# condicion para ver si no ficho INGRESO (MAS de 1 par de ingresos-egresos)
                                c.add_run((msgWord).format(dia,fecha,nombre,salida,reIngreso,hora_ingresoHoras))
                                break
                                
                            elif hora_ingreso > salidaMedioDia:# condicion para ver si no ficho EGRESO (MAS de 1 par de ingresos-egresos)
                                c.add_run((msgWord).format(dia,fecha,nombre,reIngreso,salida,hora_ingresoHoras))
                                break
                else:
                    for posicion in range(4,14,2):#analisis de las columnas en pares de 2 (ingreso,egreso)
                        hora_ingreso = frame.iloc[frame[frame[campo] == 0].index[x],posicion]#(ingreso)
                        hora_egreso = frame.iloc[frame[frame[campo] == 0].index[x],(posicion +1)]#(egreso)
                        
                        hora_ingresoHoras = '{:02d}:{:02d}'.format(hora_ingreso.hour,hora_ingreso.minute)#String con hora de ingreso
                        hora_salidaHoras = '{:02d}:{:02d}'.format(hora_egreso.hour,hora_egreso.minute)#String con hora de salida
                        
    
                        if posicion > 4:#Si hay mas de 1 pas de ingreso-egreso, carga el egreso inmediato anterior
                            hora_egreso_anterior = frame.iloc[frame[frame[campo] == 0].index[x],(posicion -1)]                        
    
                        
                        if (hora_ingreso and hora_egreso) != cero: #si el par esta completo sigue
                            continue
                        
    
                        if hora_ingreso != cero and hora_egreso == cero: #condicion de par INCOMPLETO, con solo 1 par de ingreso-egreso                    
                            if posicion == 4 and hora_ingreso > mitadMañana: #condicion para ver si no ficho INGRESO 
                                c.add_run((msgWord).format(dia,fecha,nombre,ingreso,salida,hora_ingresoHoras))
                                break
                            
                            elif posicion == 4 and hora_ingreso < mitadMañana:#condicion para ver si no ficho EGRESO
                                c.add_run((msgWord).format(dia,fecha,nombre,salida,ingreso,hora_ingresoHoras))
                                break
                        
                        # A partir de aca, estas condiciones implican mas de 1 par de ingreso-egreso.
                        if ((hora_ingreso - hora_egreso_anterior).seconds)/3600 > 1.0: #Verifica si la diferencia entre el ultimo ingreso y la ultima salida registrada
                        # es mayor a 1 hora, en ese caso considera que no se ficho el INGRESO o EGRESO del ALMUERZO.
                            c.add_run((msgWord).format(dia,fecha,nombre,almuerzo,salida,hora_ingresoHoras))
                            break
                        
                        elif ((hora_ingreso - hora_egreso_anterior).seconds)/3600 < 1.0:#Verifica si la diferencia entre el ultimo ingreso y la ultima salida registrada
                        # es menor a 1 hora, en ese caso considera que no se ficho EGRESO del dia.
                            c.add_run((msgWord).format(dia,fecha,nombre,salida,salida,hora_ingresoHoras))
                            break
                            
                            
            
    
            word = nombreInformeNoFichadasWord.format(str(fechaInicio).replace('/','-'),str(fechaFin).replace('/','-'))
            pathToWord = os.path.join(os.getcwd(),pathInformesNoFichadas,word)
            pdf = nombreInformeNoFichadasPDF.format(str(fechaInicio).replace('/','-'),str(fechaFin).replace('/','-'))
            pathToPDF = os.path.join(os.getcwd(),pathInformesNoFichadas,pdf)
            doc.save(pathToWord)
     
            wdFormatPDF = 17 # Word's numeric code for PDFs.
            wordObj = win32com.client.Dispatch('Word.Application')
            docObj = wordObj.Documents.Open(pathToWord)
            docObj.SaveAs(pathToPDF, FileFormat=wdFormatPDF)
            docObj.Close()
            wordObj.Quit()
            os.remove(pathToWord)
        except Exception:
            logger.error("excepcion desconocida: %s", traceback.format_exc())
        finally:
            pass

    
        
def ingresoNoFichadas(frame,MedioDia=[],feriados=[]):
        """
        
    
        Parameters
        ----------
        frame : dataFrame
            Frame con todos los registros (datos del operarios e ingresos/egresos por dia).
        MedioDia : List, optional
            DESCRIPTION. The default is []. Opcional, una lista con los medios dias en los
            dias bajo analisis.
        feriados : List, optional
            DESCRIPTION. The default is [].Opcional, una lista con los feriados en los
            dias bajo analisis.
    
        Returns
        -------
        frame : dataFrame
            dataFrame corregido ya con todos los registros de manera correcta, a partir de este se puede guardar en la
            BD o hacer calculos.    
        """
        try:
            logger.info('Comenzando inform de no fichadas')
            print('\n\nLos siguientes operarios no ficharon, ingrese las horas en formato HH:MM\n')
            campo = 'H.Norm' #campo sobre el cual se filtra para ver las filas que tienen errores en los registros. Es siempre el mismo
            len_noMarca = len(frame[frame[campo] == 0]) #Devuelve la cantidad de filas a las que no se le calculo
    
            msgWord = '\n\tEl dia {} ({}) el empleado {:10s} no ficho {}, {} a las {}.'
            ingreso = 'ingreso'
            reIngreso = 're-ingreso'
            salida = 'salida'
            almuerzo = colored(salida+' o '+reIngreso,attrs=['underline']) #Se les asigna color y subrayado a las palabras ingreso-egreso        
            ingresoColor = colored('Ingreso: ','green',attrs=['bold','dark','underline']) #Se les asigna color y subrayado a las palabras ingreso-egreso
            egresoColor = colored('Egreso: ','red',attrs=['bold','dark','underline']) #Se les asigna color y subrayado a las palabras ingreso-egreso
            for x in range(len_noMarca): #Itera sobre las filas que tienen algun error,
    
                legajo = frame.iloc[frame[frame[campo] == 0].index[x],0] #con x y el filtro establece la fila con error.
                nombre = frame.iloc[frame[frame[campo] == 0].index[x],1]
                nombre = colored(nombre,'grey',on_color='on_red')
                dia = frame.iloc[frame[frame[campo] == 0].index[x],2]
                fecha = frame.iloc[frame[frame[campo] == 0].index[x],3]
                
                #-------- limites horarios ----------
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                salidaMedioDia = pd.to_datetime(('{} 12:30').format(fecha))
                mitadMañana = pd.to_datetime(('{} 10:30').format(fecha))
                treceHoras = pd.to_datetime(('{} 13:00').format(fecha))
                           
                for posicion in range(4,14,2):#analisis de las columnas en pares de 2 (ingreso,egreso)
                        hora_ingreso = frame.iloc[frame[frame[campo] == 0].index[x],posicion]#(ingreso)
                        hora_egreso = frame.iloc[frame[frame[campo] == 0].index[x],(posicion +1)]#(egreso)
                        
                        hora_ingresoHoras = '{}:{}'.format(hora_ingreso.hour,hora_ingreso.minute)#String con hora de ingreso
                        hora_salidaHoras = '{}:{}'.format(hora_egreso.hour,hora_egreso.minute)#String con hora de salida
                        
    
                        if posicion > 4:
                            hora_egreso_anterior = frame.iloc[frame[frame[campo] == 0].index[x],(posicion -1)]                        
    
                        
                        if (hora_ingreso and hora_egreso) != cero: #si el par esta completo sigue
                            continue
                        
    
                        if hora_ingreso != cero and hora_egreso == cero:                        
                            if posicion == 4 and hora_ingreso > mitadMañana: # Verifica si solo hay 1 par de ingreso-egreso y si no se ficho INGRESO.
                                print(('El dia {} el operario {:10s} no ficho INGRESO. Ingreso  a la hora: ').format(fecha,nombre))
                                horaEntrada = str(pyip.inputDatetime('Ingrese el horario de Ingreso en formato HH:MM: ',formats=["%H:%M"]))
                                print('\n') 
                                horaEntrada = horaEntrada.split()[1] #Rompe la str ya que inputDatetime devuelve '1900-01-01 HH:MM:SS' y se queda
                                # con la segunda parte HH:MM:SS
                                horaEntrada = horaEntrada.split(':')[0]+':'+horaEntrada.split(':')[1] #Vuelve a romper y se queda con HH y MM
                                horaEntrada = pd.to_datetime(('{} {}').format(fecha,str(horaEntrada))) #Da el formato correcto con 'fecha HH:MM'
                                
                                horaSalida = frame.iloc[frame[frame[campo] == 0].index[x],4] #guarda la hora salida para no borrarla.
                                frame.iloc[frame[frame[campo] == 0].index[x],4] = horaEntrada #Donde esta la hora de salida, pone el ingreso
                                frame.iloc[frame[frame[campo] == 0].index[x],5] = horaSalida #Donde esta la hora de ingreso, pone el salida.
                                break
                            
                            elif posicion == 4 and hora_ingreso < mitadMañana:# Verifica si solo hay 1 par de ingreso-egreso y si no se ficho EGRESO.
                                print(('El dia {} el operario {:10s} no ficho SALIDA. Salio  a la hora: ').format(fecha,nombre))
                                horaSalida = str(pyip.inputDatetime('Ingrese el horario de SALIDA en formato HH:MM: ',formats=["%H:%M"]))
                                print('\n') 
                                horaSalida = horaSalida.split()[1]
                                horaSalida = horaSalida.split(':')[0]+':'+horaSalida.split(':')[1]
                                horaSalida = pd.to_datetime(('{} {}').format(fecha,str(horaSalida)))
                                
                                horaEntrada = frame.iloc[frame[frame[campo] == 0].index[x],4]
                                frame.iloc[frame[frame[campo] == 0].index[x],4] = horaEntrada
                                frame.iloc[frame[frame[campo] == 0].index[x],5] = horaSalida
                                break
                        
                        # A partir de aca, estas condiciones implican mas de 1 par de ingreso-egreso.
                        if ((hora_ingreso - hora_egreso_anterior).seconds)/3600 > 1.0: #Verifica si la diferencia entre el ultimo ingreso y la ultima salida registrada
                        # es mayor a 1 hora, en ese caso considera que no se ficho el INGRESO o EGRESO del ALMUERZO.
                            print('='*80)
                            print(('El dia {} el operario {:10s} no ficho Salida o re-ingreso (Almuerzo). Re-ingrese las horas:\n').format(fecha,nombre))
                            print('Entradas y salidas del dia:\n')
                            cantidad = 0
                            for horario in range(4,12,2):#Itera sobre los registros con error y los imprime
                                
                                ingreso = frame.iloc[frame[frame[campo] == 0].index[x],horario]
                                egreso = frame.iloc[frame[frame[campo] == 0].index[x],horario +1]
                                
                                if egreso != cero:
                                    cantidad +=2
                                    print(('{} {}').format(ingresoColor,str(ingreso))+'      '+('{} {}').format(egresoColor,str(egreso))+'\n')
    
                                else:#en caso de llegar a un registro 00:00 corta el ciclo
                                    cantidad+= 1
                                    if ingreso != cero:
                                        print(('{} {}').format(ingresoColor,str(ingreso))+'\n')                                    
                                    break
                            desicion = False    
                            while not desicion:
                            #En esta variable se almacenan TODOS los horarios NUEVAMENTE, mas el horario que faltaba.
                                horarios = [str(pyip.inputDatetime('Ingrese horario en formato HH:MM: ',formats=["%H:%M"])) for x in range(cantidad +1)]
                                print('\n')                                
                                horasLimpio = []
                                
                                for y in range(len(horarios)):
                                    hora = horarios[y]
                                    hora = hora.split()[1]
                                    hora = hora.split(':')[0]+':'+hora.split(':')[1]
                                    hora = pd.to_datetime(('{} {}').format(fecha,str(hora)))
                                    horasLimpio.append(hora)
                                
                                for idx in range(0,len(horasLimpio),2):#Iterera sobre los nuevos horarios que van a quedar y los imprime
                                    ingreso = horasLimpio[idx]
                                    egreso = horasLimpio[idx +1]
        
                                    print(('{} {}').format(ingresoColor,str(ingreso))+'      '+('{} {}').format(egresoColor,str(egreso))+'\n')
                                desicion = repreguntar()
                            
                            for idx in range(0,len(horasLimpio),2):
                                ingreso = horasLimpio[idx]
                                egreso = horasLimpio[idx +1]
                                frame.iloc[frame[frame[campo] == 0].index[x],4 +idx] = ingreso
                                frame.iloc[frame[frame[campo] == 0].index[x],4 +(idx +1)] = egreso
                            print('='*80)
                            print('\n') 
                            break
                        
                        elif ((hora_ingreso - hora_egreso_anterior).seconds)/3600 < 1.0:#Verifica si la diferencia entre el ultimo ingreso y la ultima salida registrada
                        # es menor a 1 hora, en ese caso considera que no se ficho EGRESO del dia.
                            print('='*80)
                            print(('El dia {} el operario {:10s} no ficho SALIDA. Ingrese la hora:\n').format(fecha,nombre))
                            cantidad = 0
                            for horario in range(4,12,1):
                                hora = frame.iloc[frame[frame[campo] == 0].index[x],horario]
                                if hora == cero:
                                    break
                                else:
                                    cantidad +=1
                            hora = str(pyip.inputDatetime('Ingrese horario en formato HH:MM: ',formats=["%H:%M"]))
                            print('\n') 
                            hora = hora.split()[1]
                            hora = hora.split(':')[0]+':'+hora.split(':')[1]
                            hora = pd.to_datetime(('{} {}').format(fecha,str(hora)))
                            frame.iloc[frame[frame[campo] == 0].index[x],4 +cantidad] = hora
                            print('='*80)
                            print('\n') 
                            break
        
           
            return frame
        except Exception:
            print('Hubo un error en la actualizacion, por favor vuelva a intentarlo')
            logger.error("excepcion desconocida: %s", traceback.format_exc())
        finally:
            pass

if __name__ == '__main__':
    pass