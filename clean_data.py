import os
import logging.config
import traceback
import sys
import pandas as pd
import pyinputplus as pyip
import numpy as np
import win32com.client
import docx
import colorama
colorama.init()
from paths import empleados_text,pathExcelTemporal,nombreExcelTemporal,pathExcelInforme,pathTXT,areas,formaDePago
from analizador import Analizador,CalculadorHoras,informeNoFichadas,ingresoNoFichadas
from createDB import ManagerSQL
from queryes import queryConsultaEmpleados,insertRegistros,selectAll,selectSome,insertEmpleado,deleteEmpleado
from openpyxl import load_workbook
from time import sleep
from termcolor import colored

logging.config.fileConfig('logger.ini', disable_existing_loggers=False)
logger = logging.getLogger(__name__)


from paths import(nombreInformeNoFichadasWord,nombreInformeNoFichadasPDF,pathInformesNoFichadas,
pathInformesFaltasTardanzas,
nombreInformeFaltasTardanzasWord,
nombreInformeFaltasTardanzasPDF)


def fechasDeCalculo():
    logger.info("Registrando fechas")
    fechaInicio = pyip.inputDate('Ingrese el primer dia habil DD/MM/AAAA: ',formats=["%d/%m/%Y"])
    print('\n')
    fechaFin = pyip.inputDate('Ingrese el ultimo dia habil  DD/MM/AAAA: ',formats=["%d/%m/%Y"])
    print('\n')
    feriados =[]
    feriado = None
    while feriado != '':
        feriado = pyip.inputDate(prompt='Ingrese el/los dias feriados DD/MM/AAAA: ',formats=["%d/%m/%Y"],blank=True)
        if feriado != '':
            feriados.append(feriado)
    print('\n')
    mediosDias =[]
    medioDia = None
    while medioDia != '':
        medioDia = pyip.inputDate(prompt='Ingrese el/los medios dias DD/MM/AAAA: ',formats=["%d/%m/%Y"],blank=True)
        if medioDia != '':
            mediosDias.append(medioDia)
    print('\n')
    return fechaInicio,fechaFin,feriados,mediosDias


def ingreso_egreso(line,frame,legajo,nombre):
    try:
        fecha = pd.to_datetime(line.split()[1],yearfirst=True,format='%d/%m/%Y').date()#Da formato datetime.date
        ingresos_egresos = []
        for x in range(2,12,1):#Esta parte se encarga de dar formato al horario, siempre con dia y horas.
            try:
                ingresos_egresos.append(pd.to_datetime(('{} {}').format(fecha,line.split()[x])))            
            except:
                ingresos_egresos.append(pd.to_datetime(('{} 00:00').format(fecha)))
    

        lista_final = []
        lista_final.append(ingresos_egresos[0])
            
        for indice in range(len(ingresos_egresos) -1):#Elimina registros duplicados,usa como limite unos 5 minutos.            
            if ingresos_egresos[indice + 1] - ingresos_egresos[indice] > pd.Timedelta(minutes=5):
               lista_final.append(ingresos_egresos[indice + 1])
           
        for u in range((10-len(lista_final))):
            lista_final.append(pd.to_datetime(('{} 00:00').format(fecha)))

            # turno = pyip.inputCustom(customValidationFunc=validador,prompt='Ingrese el turno del operario:\n1. Mañana\n2. Tarde\n3. Noche',
            #             postValidateApplyFunc=postValidacion)
            # turno = input('Seleccione un turno:\n1.Mañana\n2.Noche\n')
        frame = frame.append({'Empleado':legajo,'Nombre':nombre,'Dia':line.split()[0],'Fecha':fecha,
                                  'Ingreso_0':lista_final[0],'Egreso_0':lista_final[1],
                                  'Ingreso_1':lista_final[2],'Egreso_1':lista_final[3],
                                  'Ingreso_2':lista_final[4],'Egreso_2':lista_final[5],
                                  'Ingreso_3':lista_final[6],'Egreso_3':lista_final[7],
                                  'Ingreso_4':lista_final[8],'Egreso_4':lista_final[9]},
                                  ignore_index=True) #agrega una fila mas al dataframe con los datos de ese dia
                                
        return frame
    except Exception as e:
        logger.error('',exc_info=True)
        return None

def creacionFrameVacio():
    columnas = ['Empleado','Nombre','Dia','Fecha','Ingreso_0','Egreso_0','Ingreso_1','Egreso_1',
                        'Ingreso_2','Egreso_2','Ingreso_3','Egreso_3','Ingreso_4','Egreso_4',
                        ]
    frame = pd.DataFrame(columns=columnas)
    return frame

def insercionBD(managerSQL,frame,query):
    logger.info('Insertando registros en la BD')
    for index, row in frame.iterrows():
        queryInsercion = query.format(row['Empleado'],row['Nombre'], row['Dia'],str(row['Fecha']),
          str(row['Ingreso_0']),str(row['Egreso_0']),
          str(row['Ingreso_1']),str(row['Egreso_1']),
          str(row['Ingreso_2']),str(row['Egreso_2']),
          str(row['Ingreso_3']),str(row['Egreso_3']),
          str(row['Ingreso_4']),str(row['Egreso_4']))
        managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 

def insercionBDLegajos(managerSQL,legajo,nombre,apellido,area,pago,query):
    logger.info('Insertando nuevo operario en la BD')
    queryInsercion = query.format(legajo,nombre,apellido,area,pago)
    managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 

def deleteBDLegajos(managerSQL,legajo,query):
    logger.info('Borrando operario en la BD')
    queryInsercion = query.format(legajo)
    managerSQL.executeQuery(managerSQL.conexion(),queryInsercion) 
        

def frameFichadas():
    
    semaine = {'Lunes': 0,
               'Martes' : 1,
               'Miércoles': 2,
               'Jueves':3,
               'Viernes':4,
               'Sábado':5,
               'Domingo':6,
              }
    
    frame = creacionFrameVacio()
    MedioDia = None   
    
    listaTXT = os.listdir(pathTXT)
    
    if len(listaTXT) == 0:        
        return False
    else:
        archivo = pyip.inputMenu(listaTXT,prompt='Elija uno de los archivos:\n',lettered=True)
    archivo = os.path.join(os.getcwd(),pathTXT,archivo)
    print('\n')
    try:
        with open(archivo) as file:
            for line in file.readlines():
                        if line.startswith('Empleado'):
                            legajo = line.split()[1].replace('.',"")
                            nombre = ''
                            for x in range(3,7):
                                if line.split()[x] == 'Tarjeta':
                                    break
                                else:
                                    nombre += line.split()[x]+' '
                            nombre = nombre.upper()
                        for jour in semaine.keys():        
                            if line.startswith(jour):
                                frame = ingreso_egreso(line,frame,legajo,nombre)
    except:
        with open(archivo,encoding="utf-8") as file:
                    for line in file.readlines():
                        if line.startswith('Empleado'):
                            legajo = line.split()[1].replace('.',"")
                            nombre = ''
                            for x in range(3,7):
                                if line.split()[x] == 'Tarjeta':
                                    break
                                else:
                                    nombre += line.split()[x]+' '
                            nombre = nombre.upper()
                        for jour in semaine.keys():        
                            if line.startswith(jour):
                                frame = ingreso_egreso(line,frame,legajo,nombre)

    return frame

def logicaRotativos(frame,inyeccion=False,fechaInicio=None,fechaFin=None):
    """
    Esta funcion se encarga de limpiar los registros de los dataframes individuales de cada
    empleado rotativo, aqui dentro esta toda la logica de limpieza de esos registros.

    Parameters
    ----------
    frame : DataFrame
        Dataframe con los registros de un empleado Rotativo.

    fechaInicio : datetime.date
        Fecha de inicio para el analisis.
    fechaFin : TYPE
        DESCRIPTION.

    Returns
    -------
    frame : DataFrame
        DataFrame ya corregido. Los lugares vacios corresponden a faltas en los fichajes.

    """

    mascara = (frame['Fecha'] >= fechaInicio) & (frame['Fecha'] <= fechaFin) #mascara para filtrar el frame en funcion de la fecha de inicio y fin
    frameOriginal = frame
    frameEnAnalisis = frame.loc[mascara].copy()
    limpiador = Analizador(frameOriginal=frameOriginal,frameEnAnalisis=frameEnAnalisis)
    newFrame = limpiador.limpiador(inyeccion=inyeccion)
    
        
    return newFrame 

def frameAnalisisIndividual(frame,fechaInicio,fechaFin):
    """
    Esta funcion se encarga de tomar el frame completo de todas las horas e ir armando 1 frame por cada empleado
    y si esta dentro de los rotativos encara con una logica distinta.    
    En caso de ser normal y no fichar salida o ingreso queda el espacio vacio.    
    Por ultima genera un dataframe totalizando todo
    
    Returns
    -------
    Dataframe

    """
    
    manager = ManagerSQL()
    sql_conection = manager.conexion()
    consultaEmpleados = pd.read_sql(queryConsultaEmpleados,sql_conection)

    
    legajosSinInyeccion = consultaEmpleados.loc[(consultaEmpleados['AREA'] != 'INYECCION') &(consultaEmpleados['AREA'] != 'SOPLADO')] 
    legajosSinInyeccion = legajosSinInyeccion['LEG'].unique()
    legajosSinInyeccion = [int(x) for x in legajosSinInyeccion]#los pasa de numpy.int64 a int

    legajosFrame = frame['Empleado'].unique()

    legajos = frame['Empleado'].unique()
    
    fechaInicio = fechaInicio
    fechaFin = fechaFin
    
    lista =[]
    frameAnalisis = creacionFrameVacio()
    for legajo in legajos:
        newFrame = frame[frame['Empleado']==legajo]#legajo es un STR
        if int(legajo) in legajosSinInyeccion:
            mascara = (frame['Fecha'] >= fechaInicio) & (frame['Fecha'] <= fechaFin) #mascara para filtrar el frame en funcion de la fecha de inicio y fin
            newFrame = newFrame.loc[mascara].copy()
        else:
            continue
            
        
        frameAnalisis = frameAnalisis.append(newFrame)
    
    return frameAnalisis

def limpiezaDeRegistros(frame,fechaInicio,fechaFin):
    
    manager = ManagerSQL()
    sql_conection = manager.conexion()
    consultaEmpleados = pd.read_sql(queryConsultaEmpleados,sql_conection)

    
    legajosSinInyeccion = consultaEmpleados.loc[(consultaEmpleados['AREA'] != 'INYECCION') &(consultaEmpleados['AREA'] != 'SOPLADO')] 
    legajosSinInyeccion = legajosSinInyeccion['LEG'].unique()
    legajosSinInyeccion = [int(x) for x in legajosSinInyeccion]#los pasa de numpy.int64 a int

    legajosFrame = frame['Empleado'].unique()

    frameAnalisis = creacionFrameVacio()
    calculador = CalculadorHoras()
    
    for legajo in legajosFrame:
        if int(legajo) in legajosSinInyeccion:
            newFrame = frame[frame['Empleado']==legajo].copy()            
            frameCalculado = calculador.horasTrabajadas(newFrame)        
            frameAnalisis = frameAnalisis.append(frameCalculado)
        
    frameAnalisis = frameAnalisis.reset_index(drop=True) 
    informeNoFichadas(frameAnalisis,fechaInicio,fechaFin,mediosDias=[],feriados=[])#Crea el informe de no fichadas
    campo = 'H.Norm' #campo sobre el cual se filtra para ver las filas que tienen errores en los registros. Es siempre el mismo
    len_noMarca = len(frameAnalisis[frameAnalisis[campo] == 0])
    
    if len_noMarca == 0:
        print('TODOS los registros se encuentran completos, se procede a escribir la Base de datos.')
        insercionBD(manager,frameAnalisis,insertRegistros)
    else:
        nombre = nombreExcelTemporal.format(str(fechaInicio).replace('/','-'),str(fechaFin).replace('/','-'))
        nombre = os.path.join(os.getcwd(),pathExcelTemporal,nombre)
        frameAnalisis.to_excel(nombre,index=False)
          
    
def actualizacionRegistros():
    
    listaExcels = os.listdir(pathExcelTemporal)
    listaCodigosClean = [x for x in listaExcels if 'Excel' in x]
    if len(listaCodigosClean) == 1:
        archivo = listaCodigosClean[0]
        msg = 'Se procede a utilizar el unico archivo en la carpeta\n{}'.format(str(listaCodigosClean[0]))
        print(msg)
    else: 
        archivo = pyip.inputMenu(listaCodigosClean,prompt='Elija uno de los archivos:\n',lettered=True)
    print('\n')
    manager = ManagerSQL()
    calculador = CalculadorHoras()
    
    # nombre = nombreExcelTemporal.format(str(fechaInicio).replace('/','-'),str(fechaFin).replace('/','-'))
    nombre = os.path.join(os.getcwd(),pathExcelTemporal,archivo)
    
    frameAnalisis = pd.read_excel(nombre)
    frameAnalisis['Fecha'] = pd.to_datetime(frameAnalisis['Fecha']).dt.date #transforma 2020-10-01 00:00:00 a 2020-10-01
    
    frameCorregido = ingresoNoFichadas(frameAnalisis)#Corrige los registros que estan en cero    
    insercionBD(manager,frameCorregido,insertRegistros)#Escribe la BD    


def validador(eleccion):
    for letter in eleccion:
        if letter in ["\\","/",":","*","?","<",">","|"]:
            raise Exception('\nExisten caracteres que no pueden estar incluidos en el nombre\n\n')

def hojaTotalizadora(frame,fechaInicio,fechaFin,feriados):
    
    fechaInicio = pd.to_datetime(fechaInicio).date()
    fechaFin = pd.to_datetime(fechaFin).date()
    diasLaborales = list(pd.bdate_range(fechaInicio,fechaFin))

    
    diasLaborales = len(list(pd.bdate_range(fechaInicio,fechaFin)))
    diasLaborales = diasLaborales - len(feriados)
    frame = frame.groupby(["Legajo","Nombre"])
    frameLegajo = frame.sum()
    frameLegajo.reset_index(inplace=True)
    frameDiasLaborales = frame[['Dia']].count()
    frameConcatenado = pd.merge(frameLegajo,frameDiasLaborales,on='Legajo')
    frameConcatenado = frameConcatenado.rename(columns={'Dia':'Dias Trabajados'})
    #frameConcatenado['Faltas'] = diasLaborales - frameConcatenado['Dias Trabajados']
    
    return frameConcatenado
    
    
    
    
def seleccionInformes(fechaInicio,fechaFin,mediosDias=[],feriados=[]):    
    
    
    columnas = ['fecha','ingreso0','egreso0',
                'ingreso1','egreso1',
                'ingreso2','egreso2',
                'ingreso3','egreso3',
                'ingreso4','egreso4']
    columnasReemplazo = {'legajo':'Legajo','nombre':'Nombre',
                         'dia':'Dia','fecha':'Fecha',
                         'ingreso0':'Ingreso_0','egreso0':'Egreso_0',
                         'ingreso1':'Ingreso_1','egreso1':'Egreso_1',
                         'ingreso2':'Ingreso_2','egreso2':'Egreso_2',
                         'ingreso3':'Ingreso_3','egreso3':'Egreso_3',
                         'ingreso4':'Ingreso_4','egreso4':'Egreso_4'}
    calculador = CalculadorHoras() 
    manager = ManagerSQL()
    informes = ['Todos los legajos','Algunos legajos']
    respuesta = pyip.inputMenu(informes,prompt='Seleccione alguno de los informes disponibles: \n',
                               lettered=True)
    print('\n')
    if respuesta == 'Todos los legajos':        
        query = selectAll.format(fechaInicio,fechaFin)
        frameCorregido = pd.read_sql(query,manager.conexion(),parse_dates=columnas)
        frameCorregido['fecha'] = pd.to_datetime(frameCorregido['fecha']).dt.date
        frameCorregido.rename(columns = columnasReemplazo,inplace=True)
    
    elif respuesta == 'Algunos legajos':
        legajosTupla =[]
        legajo = None
        while legajo != '':
            legajo = pyip.inputInt(prompt='Ingrese el numero de legajo:\n\t',blank=True)
            legajosTupla.append(str(legajo))
        legajosTupla = tuple(legajosTupla)
        query = selectSome.format(fechaInicio,fechaFin,legajosTupla)
        frameCorregido = pd.read_sql(query,manager.conexion(),parse_dates=columnas)
        frameCorregido['fecha'] = pd.to_datetime(frameCorregido['fecha']).dt.date
        frameCorregido.rename(columns = columnasReemplazo,inplace=True)
        
    frameCorregido.drop(['id'],axis=1,inplace=True)  
    frameCorregido = calculador.horasTrabajadas(frameCorregido)
    frameCorregido = calculador.restaRetrasosTardanzas(frameCorregido,mediosDias = mediosDias)
    frameExtras = calculador.horasExtrasTrabajadas(frameCorregido,feriados=feriados,mediosDias = mediosDias)
    
    if frameCorregido.empty:
        print('No hay registros sobre los cuales trabajar\n')
    else:
        archivo = pyip.inputCustom(prompt='Ingrese el nombre que desea ponerle al EXCEL: \n',
                                customValidationFunc=validador)
        print('\n')
        nombre = os.path.join(os.getcwd(),pathExcelInforme,archivo)
        nombre = nombre+ '.xlsx'
        frameExtras = frameExtras.sort_values(by=['Legajo','Fecha'])
        frameExtras.to_excel(nombre,sheet_name='Registros',index=False)
        
        book = load_workbook(nombre)
        writer = pd.ExcelWriter(nombre, engine = 'openpyxl')
        writer.book = book      
        
        frameTotalizado = hojaTotalizadora(frameExtras, fechaInicio, fechaFin,feriados)
        frameTotalizado.to_excel(writer, sheet_name = 'Totalizado',index=False)
        writer.save()
        writer.close()     
    
    
    return frameCorregido

def informeFaltasTardanzas(frame,fechaInicio,fechaFin,medioDias=[],feriados=[]):
    
    try:    
        fechaInicio = pd.to_datetime(fechaInicio).date()
        fechaFin = pd.to_datetime(fechaFin).date()
        diasLaborales = list(pd.bdate_range(fechaInicio,fechaFin))
        diasLaborales = [x.date() for x in diasLaborales]

        faltasWord = '\tFalta registrada el dia {}.\n'
        tardanzasWord = '\tTardanza de {} minutos registrada el dia {} ({}).\n'
        retirosWord = '\tRetiro anticipado de {} minutos registrado el dia {} ({}).\n'
        
        legajosFrame = frame['Legajo'].unique()

        if len(feriados) >= 1 :
                for feria in feriados:
                    diasLaborales.remove(feria)
        empleados = {}
        for legajo in legajosFrame:
            empleados[str(legajo)]={}
            empleados[str(legajo)]['Tardanzas'] = {}
            empleados[str(legajo)]['Retiros'] = {}
            empleados[str(legajo)]['Nombre'] = ''

        for legajo in legajosFrame:    
            faltas = []
            newFrame = frame[frame['Legajo']==int(legajo)].copy()
            diasTrabajados = list(newFrame['Fecha'])
            #print('FeriadosTipo: ',type(feriados[0]),feriados[0],'  ',len(feriados))
            #print('DiasTipo: ',type(diasLaborales[0]),diasLaborales[0],'  ',len(diasLaborales))
            
            for dia in diasLaborales:
                if dia not in diasTrabajados:
                    faltas.append(dia)
            empleados[str(legajo)]['Faltas'] = faltas
            
            for x in range(len(newFrame)):
                legajo = newFrame.iloc[x,0]
                nombre = newFrame.iloc[x,1]
                dia = newFrame.iloc[x,2]
                fecha = newFrame.iloc[x,3]
                horaIngreso = pd.to_datetime(('{} 08:00').format(fecha))
                horaSalida = pd.to_datetime(('{} 16:48').format(fecha))
                horaSalidaMedioDia = pd.to_datetime(('{} 12:30').format(fecha))
                cero = pd.to_datetime(('{} 00:00').format(fecha))
                tardanza = 0
                retiro= 0
                
                for idx in range(5,13,2):
                    if newFrame.iloc[x,idx] == cero:
                        salida = newFrame.iloc[x,idx -2]
                        break
                
                if newFrame.iloc[x,4] > horaIngreso:
                    tardanza = round((((newFrame.iloc[x,4] - horaIngreso).seconds)/60),2)               
                
                if fecha in medioDias:
                    if salida < horaSalidaMedioDia:
                        print(salida,'   ',horaSalidaMedioDia)
                        retiro = round((((horaSalidaMedioDia - salida).seconds)/60),2)                        
                elif fecha not in medioDias:
                    if salida < horaSalida and dia != 'Sábado':
                        retiro = round((((horaSalida - salida).seconds)/60),2)  

                if tardanza != 0:                    
                    empleados[str(legajo)]['Tardanzas'][str(fecha)] = (dia,tardanza)
                if retiro != 0:
                    empleados[str(legajo)]['Retiros'][str(fecha)] = (dia,retiro)
                
                empleados[str(legajo)]['Nombre']= nombre
        
        doc = docx.Document()
        doc.add_heading(('Faltas, tardanzas y retiros entre {} y {}').format(fechaInicio,fechaFin), 0)
        for key in empleados.keys():
            
            doc.add_heading(('Informe sobre {}:').format(empleados[key]['Nombre']),level=1)
            
            
            doc.add_heading(('Faltas:'),level=2)
            primerParrafo = doc.add_paragraph()
            for falta in empleados[key]['Faltas']:#es una lista el resultado 
                primerParrafo.add_run((faltasWord).format(falta))
                
            
            doc.add_heading(('Tardanzas:'),level=2)
            segundoParrafo = doc.add_paragraph()
            for llave,valor in empleados[key]['Tardanzas'].items():
                segundoParrafo.add_run((tardanzasWord).format(valor[1],llave,valor[0]))
            
            
            doc.add_heading(('Retiros anticipados:'),level=2)
            tercerParrafoInterno = doc.add_paragraph()
            for llave,valor in empleados[key]['Retiros'].items():
                tercerParrafoInterno.add_run((retirosWord).format(valor[1],llave,valor[0]))
        
        archivo = pyip.inputCustom(prompt='Ingrese el nombre que desea ponerle al informe PDF: \n',
                            customValidationFunc=validador)

        nombre = os.path.join(os.getcwd(),pathInformesFaltasTardanzas,archivo)
        word = nombre+ '.docx'
        pdf = nombre + '.pdf'
        
        pathToWord = word
        pathToPDF = pdf
        doc.save(pathToWord)
        logger.info('Creacion del Word de manera correcta')
        sleep(0.5)
        wdFormatPDF = 17 # Word's numeric code for PDFs.
        wordObj = win32com.client.Dispatch('Word.Application')
        docObj = wordObj.Documents.Open(pathToWord)
        docObj.SaveAs(pathToPDF, FileFormat=wdFormatPDF)
        docObj.Close()
        wordObj.Quit()
        logger.info('Creacion del PDF de manera correcta')
        os.remove(pathToWord)
    except AttributeError:
        print('Existio un problema en la creacion del word/pdf.')
        logger.error("excepcion desconocida: %s", traceback.format_exc())

def datosOperario(areas,formaDePago):
    
    legajo = pyip.inputInt(prompt = 'Ingrese el LEGAJO del empleado',min=0)
    print('\n')
    nombre = pyip.inputStr(prompt = 'Ingrese el NOMBRE del empleado').upper()
    print('\n')
    apellido = pyip.inputStr(prompt = 'Ingrese el APELLIDO del empleado').upper()
    print('\n')
    area = pyip.inputMenu(areas,prompt='Ingrese una de las AREAS posibles\n',lettered=True).upper()
    print('\n')
    pago = pyip.inputMenu(formaDePago,prompt='Ingrese el tipo de pago\n',lettered=True)
    print('\n')
    
    return legajo,nombre,apellido,area,pago

def repreguntar():
    decision = pyip.inputYesNo(prompt='¿Los datos ingresados son correctos? (SI/NO)  ',yesVal='SI',noVal='NO')
    print('\n')    
    if decision == 'SI':
        return True
    else:
        return False

    
class Motor:
    def mainLoop(self):
        print('-'*100)
        print(' '*50,'Bienvenido',' '*50)
        print('-'*100)
    
    
        tareas = ['Ordenado de registros','Creación de informes','Gestion de Base de datos','Salir']
        tareasOrdenado = ['Limpieza de registros','Actualización de registros','Volver','Salir']
        tareasInformes = ['Ingreso de fechas','Volver','Salir']
        tareasBD = ['Insertar registro','Actualizar registro','Eliminar registro','Descargar','Volver','Salir']
        
        continuar = True
        
        
        while continuar:
            respuesta = pyip.inputMenu(tareas,prompt='¿Que desea hacer?\n',lettered=True)
            print('\n')
            continuarOrdenado = True
            continuarCreacion = True
            continuarGestionBD = True
            
            if respuesta == 'Ordenado de registros':
                while continuarOrdenado:
                    
                    ordenadoRespuesta = pyip.inputMenu(tareasOrdenado,prompt='\n¿Que desea hacer?\n',lettered=True)
                    print('\n')
                
                    if ordenadoRespuesta == 'Limpieza de registros':
                        fechaInicio,fechaFin,feriados,mediosDias = fechasDeCalculo()
                        frame = frameFichadas()
                        
                        if frame.empty:
                            print('No existen archivos que limpiar\n')
                        else:
                            legajos = frameAnalisisIndividual(frame,fechaInicio,fechaFin)
                            limpiezaDeRegistros(legajos, fechaInicio, fechaFin)               
                    elif ordenadoRespuesta == 'Actualización de registros':
                        actualizacionRegistros()                
                    elif ordenadoRespuesta == 'Volver':
                        print('Volviendo al PRIMER MENU')
                        continuarOrdenado = False                    
                    elif ordenadoRespuesta == 'Salir':
                        continuarOrdenado = False
                        continuar = False
            
            
            elif respuesta == 'Creación de informes':
                while continuarCreacion:
                    informesRespuesta = pyip.inputMenu(tareasInformes,prompt='\n¿Que desea hacer?\n',lettered=True)
                    print('\n')
                    if informesRespuesta == 'Ingreso de fechas':
                        fechaInicio,fechaFin,feriados,mediosDias = fechasDeCalculo()
                        frameCorregido = seleccionInformes(fechaInicio, fechaFin,feriados = feriados,mediosDias= mediosDias)
                        if frameCorregido.empty:
                            pass 
                        else:
                            informeFaltasTardanzas(frameCorregido,fechaInicio,fechaFin,
                                    feriados=feriados,medioDias = mediosDias)
                        continuarCreacion = True                
                    
                    elif informesRespuesta == 'Volver':
                        print('Volviendo al PRIMER MENU')
                        continuarCreacion = False
                        
                    elif informesRespuesta == 'Salir':
                        continuarCreacion = False
                        continuar = False
    
            
            
            elif respuesta == 'Gestion de Base de datos':
                while continuarGestionBD:
                    baseDeDatosRespuesta = pyip.inputMenu(tareasBD,prompt='\n¿Que desea hacer?\n',lettered=True)
                    print('\n')
                    ['Insertar registro','Actualizar registro','Eliminar registro','Descargar','Volver','Salir']
                    if baseDeDatosRespuesta == 'Insertar registro':
                        decision = False

                        while not decision:

                            legajo,nombre,apellido,area,pago = datosOperario(areas, formaDePago)
                            print(legajo,nombre,apellido,area,pago)
                            decision = repreguntar()

                        managerSQL = ManagerSQL()
                        insercionBDLegajos(managerSQL, legajo, nombre, apellido, area, pago, insertEmpleado)
                        continuarGestionBD = True                
                    
                    elif baseDeDatosRespuesta == 'Actualizar registro':
                        print('SIN IMPLEMENTAR AUN')
                        print('Volviendo al PRIMER MENU')
                        continuarGestionBD = True
                    
                    elif baseDeDatosRespuesta == 'Eliminar registro':
                        decision = False
                        while not decision:
                            legajo = pyip.inputInt(prompt='Ingrese el LEGAJO del empleado a eliminar',min=0)
                            decision = repreguntar()
                        print('\n')
                        managerSQL = ManagerSQL()
                        deleteBDLegajos(managerSQL, legajo, deleteEmpleado)     
                        continuarGestionBD = True
                        
                    elif baseDeDatosRespuesta == 'Descargar':
                        manager = ManagerSQL()
                        sql_conection = manager.conexion()
                        consultaEmpleados = pd.read_sql(queryConsultaEmpleados,sql_conection)
                        archivo = pyip.inputCustom(prompt='Ingrese el nombre que desea ponerle al EXCEL: \n',
                                customValidationFunc=validador)
                        print('\n')        
                        nombre = os.path.join(os.getcwd(),pathExcelInforme,archivo)
                        nombre = nombre+ '.xlsx'
                        consultaEmpleados = consultaEmpleados.sort_values(by=['LEG'])
                        consultaEmpleados.to_excel(nombre,sheet_name='Registros',index=False)
                        
                        continuarGestionBD = True
                    
                    elif baseDeDatosRespuesta == 'Volver':
                        print('Volviendo al PRIMER MENU')
                        continuarGestionBD = False
                        
                    elif baseDeDatosRespuesta == 'Salir':
                        continuarGestionBD = False
                        continuar = False
    
            
            elif respuesta == 'Salir':
                continuar= False


        
    
    
    
if __name__ == '__main__':
    try:
        motor = Motor()
        motor.mainLoop()
        sys.exit()
    except Exception:
        logger.error("excepcion desconocida: %s", traceback.format_exc())
        sleep(5)
